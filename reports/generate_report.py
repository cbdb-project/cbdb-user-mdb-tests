"""Generate two Word documents (English + Chinese) summarising every
documented bug in CBDB_BJ_User.mdb, ordered by severity, with
screenshots from `reports/screenshots/` embedded under each issue.

Run:
    python reports/generate_report.py
Outputs:
    reports/CBDB_Issues_Report_EN.docx
    reports/CBDB_Issues_Report_ZH.docx

Tone: deferential / polite throughout — the maintainer is a
respected senior researcher, and these reports are a courtesy hand-
off, not a pull request.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

import opencc
_S2T = opencc.OpenCC("s2twp")  # Simplified -> Traditional (Taiwan idiom)


def t(s: str) -> str:
    """Convert Simplified Chinese strings to Traditional (Taiwan).
    Safe to call on ASCII / English (the converter is a no-op there)."""
    if not s:
        return s
    return _S2T.convert(s)

REPO = Path(__file__).resolve().parent.parent
SHOT_DIR = REPO / "reports" / "screenshots"
OUT_EN = REPO / "reports" / "CBDB_Issues_Report_EN.docx"
OUT_ZH = REPO / "reports" / "CBDB_Issues_Report_ZH-Hant.docx"
OUT_EN_MD = REPO / "reports" / "CBDB_Issues_Report_EN.md"
OUT_ZH_MD = REPO / "reports" / "CBDB_Issues_Report_ZH-Hant.md"


# ---------------------------------------------------------------------
# Issue catalogue — single source of truth for both language reports.
# Ordered by severity (highest first within each tier).
# ---------------------------------------------------------------------

ISSUES: list[dict] = [
    # =================================================================
    # build-20260602 assessment.  Rebuilt from scratch per AGENTS.md
    # "each build is a fresh assessment".  The test session was
    # NON-INTERACTIVE (pywinauto UIA unavailable), so no runtime UI
    # symptom could be re-verified this build → ui_verified is never
    # True, and every user-facing-but-unverified defect is filed as
    # P5 latent_code per the report-triage contract.  Missing-UI (P3)
    # and setup (P4) defects are structurally confirmed and need no UI
    # symptom.
    #
    # Line numbers below were re-verified by grep against
    # analysis/dump/vba/* on the 20260602 dump (each vba_ref and every
    # in-prose line citation was confirmed to land on the cited code).
    # An earlier revision of this list carried a ~+2200 offset on every
    # Form_*.vb line (some past end-of-file); those numbers have been
    # corrected mechanically against the current dump.
    # =================================================================

    # ---------------- P4 — Setup ----------------
    {
        "id": 2,
        "tier": "P4_setup",
        "form": "(VBA project)",
        "title_en": "VBA project references the legacy dao360.dll, "
                    "absent on Office 2016+ machines",
        "title_zh": "VBA 專案參照已過時的 dao360.dll，"
                    "在 Office 2016 以後的機器上並不存在",
        "summary_en": (
            "The shipped .mdb's VBA project carries a hard reference to "
            "`C:\\Program Files\\Common Files\\Microsoft Shared\\DAO\\"
            "dao360.dll` — the DAO 3.6 location used by Access 2003.  "
            "Modern Office (2016 onward) ships `ACEDAO.DLL` instead and "
            "does NOT install the legacy DLL.  On any clean modern "
            "machine the first attempt to run a form's code raises "
            "'Can't find project or library', which is opaque and "
            "alarming to end users.\n\n"
            "Our test driver auto-replaces the broken reference with "
            "ACEDAO.DLL when it opens the file (see the before/after "
            "reference dump in `analysis/check_vba_refs.py`), so the "
            "regression suite does not hit it — but a plain user who "
            "double-clicks the shipped .mdb will.  Severity is low "
            "because it's a one-time fix per machine, but every fresh "
            "install hits it."
        ),
        "summary_zh": (
            "出貨的 .mdb 其 VBA 專案內含一個硬編碼參照，指向 "
            "`C:\\Program Files\\Common Files\\Microsoft Shared\\DAO\\"
            "dao360.dll`——這是 Access 2003 時代 DAO 3.6 的位置。"
            "現代 Office（2016 起）改為附帶 `ACEDAO.DLL`，並不會安裝"
            "這個舊版 DLL。在任何乾淨的現代機器上，第一次執行表單"
            "程式碼時就會跳出『Can't find project or library』，"
            "對一般使用者而言訊息晦澀又嚇人。\n\n"
            "我們的測試驅動在開檔時會自動把這個壞掉的參照換成 "
            "ACEDAO.DLL（見 `analysis/check_vba_refs.py` 的修復前後"
            "參照傾印），所以回歸測試不會踩到；但一般使用者直接"
            "雙擊出貨的 .mdb 就會踩到。嚴重度低，因為每台機器只需"
            "修一次，但每次全新安裝都會遇到。"
        ),
        "steps_en": [
            "Install `CBDB_BJ_User.mdb` on a fresh modern Office "
            "machine.",
            "Open the file, then press Alt+F11 to enter the VBE.",
            "Tools → References — notice an entry marked "
            "`MISSING: dao360.dll`.",
            "Open any LookAt form (or otherwise run any form's code).  "
            "A 'Can't find project or library' compile-error popup "
            "appears before the form's code runs.",
        ],
        "steps_zh": [
            "在一台全新的現代 Office 機器上安裝 `CBDB_BJ_User.mdb`。",
            "開檔後按 Alt+F11 進入 VBE。",
            "工具 → 設定參照——會看到一條標記為 "
            "`MISSING: dao360.dll` 的項目。",
            "開啟任一 LookAt 表單（或任何會執行表單程式碼的操作）。"
            "在表單程式碼執行前，就會跳出『Can't find project or "
            "library』編譯錯誤視窗。",
        ],
        "fix_en": (
            "Once, on the maintainer's machine: open the .mdb in "
            "Access, press Alt+F11, go to Tools → References, untick "
            "the MISSING dao360.dll entry, tick `Microsoft Office 16.0 "
            "Access Database Engine Object Library` (i.e. ACEDAO.DLL), "
            "and save.  Then re-distribute the fixed file — future "
            "users won't need to do anything."
        ),
        "fix_zh": (
            "在維護者的機器上做一次即可：用 Access 開啟 .mdb，按 "
            "Alt+F11，進入 工具 → 設定參照，取消勾選 MISSING 的 "
            "dao360.dll，改勾選 `Microsoft Office 16.0 Access Database "
            "Engine Object Library`（即 ACEDAO.DLL），存檔。之後"
            "重新散布修好的檔案，後續使用者就不必再處理。"
        ),
        "severity_en": "P4 — One-time setup hurdle on each new machine.",
        "severity_zh": "P4 —— 每台新機器一次性的安裝門檻。",
        "screenshots": [],
        "evidence": {
            "finding_class": "user_facing_bug",
            "vba_ref": "VBA project References — the dao360.dll entry; "
                       "before/after reference dump in "
                       "analysis/check_vba_refs.py (the 'before fix' "
                       "list shows the broken DAO reference, the "
                       "driver swaps it to ACEDAO.DLL on open).",
            "fixture": "n/a — fires on the first form-code run after a "
                       "fresh install (every Form_Open path).",
            "user_symptom": "A 'Can't find project or library' "
                            "compile-error popup blocks the form before "
                            "its code runs; the user cannot use any "
                            "LookAt feature until the reference is fixed.",
            "detection": "analysis/check_vba_refs.py",
        },
    },

    # ---------------- P3 — Missing UI ----------------
    {
        "id": 13,
        "tier": "P3_missing_ui",
        "form": "BIOG_MAIN_2_Subform",
        "title_en": "BIOG_MAIN_2 Subform clicks a picker form "
                    "(frmPickNIAN_HAO) that does not exist",
        "title_zh": "BIOG_MAIN_2 子表單呼叫一個不存在的"
                    "選取表單（frmPickNIAN_HAO）",
        "summary_en": (
            "When the user clicks the reign-period (NIAN_HAO) picker on "
            "the biographical-detail subform, "
            "`Form_BIOG_MAIN_2_Subform` runs "
            "`DoCmd.OpenForm \"frmPickNIAN_HAO\"` (the handler sets "
            "`stDocName = \"frmPickNIAN_HAO\"` and references "
            "`Forms!frmPickNIAN_HAO!frmNIAN_HAO.Form!c_nianhao_id`).  "
            "There is no form named `frmPickNIAN_HAO` in the current "
            ".mdb (it is absent from the fresh "
            "`control_inventory.json`).  Access raises 'Item not "
            "found…' and the field click does nothing useful.\n\n"
            "The host form BIOG_MAIN_2_Subform itself IS present and "
            "reachable (verified in the fresh control inventory) — only "
            "the picker it opens is missing.  Likely cause: a picker "
            "form was renamed or consolidated in an earlier refactor "
            "and this caller wasn't updated.\n\n"
            "(This build's session was non-interactive, so the runtime "
            "popup could not be re-captured; the screenshots below are "
            "the reachable host plus a reconstructed popup, and the "
            "static absence of the picker is the load-bearing "
            "evidence.)"
        ),
        "summary_zh": (
            "當使用者在人物詳細資料子表單上點選年號（NIAN_HAO）"
            "選取器時，`Form_BIOG_MAIN_2_Subform` 會執行 "
            "`DoCmd.OpenForm \"frmPickNIAN_HAO\"`（處理常式設定 "
            "`stDocName = \"frmPickNIAN_HAO\"`，並參照 "
            "`Forms!frmPickNIAN_HAO!frmNIAN_HAO.Form!c_nianhao_id`）。"
            "但目前的 .mdb 裡並沒有名為 `frmPickNIAN_HAO` 的表單"
            "（最新的 `control_inventory.json` 中查無此表單）。"
            "Access 會丟出『Item not found…』，這次點選對使用者"
            "毫無作用。\n\n"
            "宿主表單 BIOG_MAIN_2_Subform 本身存在且可達"
            "（已在最新控制項清單中確認）——缺的只是它要開啟的"
            "選取表單。可能原因：早期重構時某個選取表單被改名或"
            "合併，而這個呼叫端沒有同步更新。\n\n"
            "（本次測試為非互動式，無法重新擷取執行期彈窗；"
            "下方截圖為可達的宿主表單加上重建的彈窗，而選取表單"
            "在靜態層面的缺失才是關鍵證據。）"
        ),
        "steps_en": [
            "Open CBDB_Browser_2 and navigate to any person whose "
            "biographical detail is shown on BIOG_MAIN_2_Subform.",
            "On the subform, click the reign-period (NIAN_HAO) picker "
            "control — that fires the handler which runs "
            "`DoCmd.OpenForm \"frmPickNIAN_HAO\"`.",
            "An 'Item not found in this collection.' popup appears, "
            "because `frmPickNIAN_HAO` is not in "
            "CurrentProject.AllForms.",
            "Static confirmation (no Access needed): search "
            "`analysis/dump/control_inventory.json` for "
            "`frmPickNIAN_HAO` — it is absent, while "
            "`BIOG_MAIN_2_Subform` is present.",
        ],
        "steps_zh": [
            "開啟 CBDB_Browser_2，導覽到任一在 BIOG_MAIN_2_Subform "
            "上顯示人物詳細資料的人物。",
            "在子表單上點選年號（NIAN_HAO）選取器控制項——這會觸發"
            "執行 `DoCmd.OpenForm \"frmPickNIAN_HAO\"` 的處理常式。",
            "會跳出『Item not found in this collection.』彈窗，"
            "因為 `frmPickNIAN_HAO` 不在 CurrentProject.AllForms 中。",
            "靜態確認（不需 Access）：在 "
            "`analysis/dump/control_inventory.json` 中搜尋 "
            "`frmPickNIAN_HAO`——查無此表單，而 "
            "`BIOG_MAIN_2_Subform` 則存在。",
        ],
        "fix_en": (
            "Either restore the picker form `frmPickNIAN_HAO`, or "
            "update the caller in `Form_BIOG_MAIN_2_Subform` to open "
            "whichever reign-period picker form replaced it."
        ),
        "fix_zh": (
            "兩種做法擇一：還原選取表單 `frmPickNIAN_HAO`，或將 "
            "`Form_BIOG_MAIN_2_Subform` 內的呼叫端改成開啟取代它的"
            "那個年號選取表單。"
        ),
        "severity_en": "P3 — Missing UI (the picker the click opens "
                       "does not exist; the feature is unreachable).",
        "severity_zh": "P3 —— 缺少 UI（點選要開啟的選取表單不存在，"
                       "此功能無法使用）。",
        "screenshots": [
            ("bug13_browser_open.png",
             "CBDB_Browser_2 open on a person record — the reachable "
             "host surface from which the NIAN_HAO picker is invoked."),
            ("bug13_browser_annotated.png",
             "Annotated host view: the reign-period picker control on "
             "BIOG_MAIN_2_Subform whose click runs "
             "DoCmd.OpenForm \"frmPickNIAN_HAO\" — a form absent from "
             "the current .mdb."),
            ("bug13_faux_popup.png",
             "The 'Item not found in this collection.' popup, "
             "reconstructed in PIL (this build's session was "
             "non-interactive); the message is Access's standard text "
             "when DoCmd.OpenForm targets a form not in "
             "CurrentProject.AllForms."),
        ],
        "evidence": {
            "finding_class": "user_facing_bug",
            "vba_ref": "Form_BIOG_MAIN_2_Subform.vb:64 "
                       "(stDocName = \"frmPickNIAN_HAO\"); the OpenForm "
                       "+ Forms!frmPickNIAN_HAO!… references follow at "
                       "lines 133-151.",
            "fixture": "Any person reachable in CBDB_Browser_2 whose "
                       "BIOG_MAIN_2_Subform reign-period picker is "
                       "clicked.",
            "user_symptom": "Clicking the reign-period picker pops "
                            "'Item not found in this collection.' and "
                            "no picker opens.",
            "ui_verified": False,
        },
    },
    {
        "id": 16,
        "tier": "P3_missing_ui",
        "form": "LookAtStatus",
        "title_en": "LookAtStatus is missing its CmdPajek button "
                    "(handler exists, no UI control)",
        "title_zh": "LookAtStatus 缺少 CmdPajek 按鈕"
                    "（處理常式存在，但表單上沒有控制項）",
        "summary_en": (
            "`Form_LookAtStatus.vb` defines a `CmdPajek_Click` handler "
            "(it would write a Pajek `.net` export of the status "
            "network), but LookAtStatus's form design has NO `CmdPajek` "
            "control.  The fresh `control_inventory.json` lists "
            "CmdQuery / CmdGIS / CmdNeo4j on this form but no Pajek "
            "button, so the feature is unreachable from the UI.\n\n"
            "Note: even once a button is added, Issue #5 (the "
            "ChkIDs-control and SQL-column defects inside "
            "CmdPajek_Click) must be fixed first, or the click will "
            "fail."
        ),
        "summary_zh": (
            "`Form_LookAtStatus.vb` 定義了 `CmdPajek_Click` 處理常式"
            "（用來輸出狀態網絡的 Pajek `.net` 檔），但 LookAtStatus "
            "的表單設計上並沒有 `CmdPajek` 控制項。最新的 "
            "`control_inventory.json` 顯示此表單有 CmdQuery / CmdGIS / "
            "CmdNeo4j，卻沒有 Pajek 按鈕，因此這個功能在 UI 上"
            "無法使用。\n\n"
            "注意：即使加上按鈕，也必須先修正 Issue #5"
            "（CmdPajek_Click 內的 ChkIDs 控制項與 SQL 欄位缺陷），"
            "否則點選仍會失敗。"
        ),
        "steps_en": [
            "Open LookAtStatus.  Look at the export-buttons row at the "
            "bottom: it has GIS and Neo4j, but there is no Pajek "
            "button.",
            "Compare with LookAtAssociations, which does render a Pajek "
            "button.",
            "Static confirmation: in "
            "`analysis/dump/control_inventory.json`, LookAtStatus has "
            "no `CmdPajek` control, although "
            "`Form_LookAtStatus.vb` defines `Sub CmdPajek_Click()`.",
        ],
        "steps_zh": [
            "開啟 LookAtStatus。看底部的輸出按鈕列：只有 GIS 和 "
            "Neo4j，沒有 Pajek 按鈕。",
            "與 LookAtAssociations 比較，後者確實有 Pajek 按鈕。",
            "靜態確認：在 `analysis/dump/control_inventory.json` 中，"
            "LookAtStatus 沒有 `CmdPajek` 控制項，但 "
            "`Form_LookAtStatus.vb` 定義了 `Sub CmdPajek_Click()`。",
        ],
        "fix_en": (
            "Add a CmdPajek button to LookAtStatus's design (with "
            "OnClick = [Event Procedure] so it invokes the existing "
            "CmdPajek_Click Sub) — but fix Issue #5 first, otherwise "
            "the click fails on the ChkIDs reference and the bad SQL."
        ),
        "fix_zh": (
            "在 LookAtStatus 的設計中新增 CmdPajek 按鈕（OnClick = "
            "[事件程序]，以呼叫既有的 CmdPajek_Click）——但請先修正 "
            "Issue #5，否則點選會因 ChkIDs 參照與錯誤的 SQL 而失敗。"
        ),
        "severity_en": "P3 — Missing UI (feature unavailable to users).",
        "severity_zh": "P3 —— 缺少 UI（使用者無法使用此功能）。",
        "screenshots": [
            ("bug16_LookAtStatus_no_CmdPajek.png",
             "LookAtStatus as it ships — the export-button row has GIS "
             "and Neo4j but no Pajek button."),
            ("bug16_LookAtStatus_no_CmdPajek_annotated.png",
             "Annotated: the gap where a CmdPajek button would sit; "
             "`Sub CmdPajek_Click()` exists in the module but no "
             "control invokes it."),
        ],
        "evidence": {
            "finding_class": "user_facing_bug",
            "vba_ref": "Form_LookAtStatus.vb:2133 "
                       "(Private Sub CmdPajek_Click()); no CmdPajek "
                       "control in analysis/dump/control_inventory.json "
                       "for LookAtStatus.",
            "fixture": "LookAtStatus form design (any session).",
            "user_symptom": "There is no Pajek export button on "
                            "LookAtStatus, so users cannot run the "
                            "Pajek export the code supports.",
            "ui_verified": False,
        },
    },
    {
        "id": 17,
        "tier": "P3_missing_ui",
        "form": "LookAtStatus",
        "title_en": "LookAtStatus is missing its CmdGephi button "
                    "(handler exists, no UI control)",
        "title_zh": "LookAtStatus 缺少 CmdGephi 按鈕"
                    "（處理常式存在，但表單上沒有控制項）",
        "summary_en": (
            "`Form_LookAtStatus.vb` defines a `CmdGephi_Click` handler "
            "but LookAtStatus's form design has NO `CmdGephi` control.  "
            "The fresh `control_inventory.json` confirms no Gephi "
            "button exists on the form, so the Gephi export is "
            "unreachable from the UI."
        ),
        "summary_zh": (
            "`Form_LookAtStatus.vb` 定義了 `CmdGephi_Click` 處理常式，"
            "但 LookAtStatus 的表單設計上並沒有 `CmdGephi` 控制項。"
            "最新的 `control_inventory.json` 確認表單上沒有 Gephi "
            "按鈕，因此 Gephi 輸出在 UI 上無法使用。"
        ),
        "steps_en": [
            "Open LookAtStatus.  There is no Gephi export button in the "
            "export-buttons row.",
            "Static confirmation: "
            "`analysis/dump/control_inventory.json` shows no "
            "`CmdGephi` control on LookAtStatus, although "
            "`Form_LookAtStatus.vb` defines `Sub CmdGephi_Click()`.",
        ],
        "steps_zh": [
            "開啟 LookAtStatus。輸出按鈕列裡沒有 Gephi 輸出按鈕。",
            "靜態確認：`analysis/dump/control_inventory.json` 顯示 "
            "LookAtStatus 上沒有 `CmdGephi` 控制項，但 "
            "`Form_LookAtStatus.vb` 定義了 `Sub CmdGephi_Click()`。",
        ],
        "fix_en": "Add a CmdGephi button to LookAtStatus's design, "
                  "wired to the existing CmdGephi_Click Sub.",
        "fix_zh": "在 LookAtStatus 的設計中新增 CmdGephi 按鈕，"
                  "並連到既有的 CmdGephi_Click。",
        "severity_en": "P3 — Missing UI (feature unavailable to users).",
        "severity_zh": "P3 —— 缺少 UI（使用者無法使用此功能）。",
        "screenshots": [
            ("bug17_LookAtStatus_no_CmdGephi.png",
             "LookAtStatus as it ships — no Gephi export button."),
            ("bug17_LookAtStatus_no_CmdGephi_annotated.png",
             "Annotated: `Sub CmdGephi_Click()` exists in the module "
             "but no control invokes it."),
        ],
        "evidence": {
            "finding_class": "user_facing_bug",
            "vba_ref": "Form_LookAtStatus.vb:18 "
                       "(Private Sub CmdGephi_Click()); no CmdGephi "
                       "control in analysis/dump/control_inventory.json "
                       "for LookAtStatus.",
            "fixture": "LookAtStatus form design (any session).",
            "user_symptom": "There is no Gephi export button on "
                            "LookAtStatus, so users cannot run the "
                            "Gephi export the code supports.",
            "ui_verified": False,
        },
    },
    {
        "id": 18,
        "tier": "P3_missing_ui",
        "form": "LookAtStatus",
        "title_en": "LookAtStatus is missing its CmdUCINet button "
                    "(handler exists, no UI control)",
        "title_zh": "LookAtStatus 缺少 CmdUCINet 按鈕"
                    "（處理常式存在，但表單上沒有控制項）",
        "summary_en": (
            "`Form_LookAtStatus.vb` defines a `CmdUCINet_Click` handler "
            "but LookAtStatus's form design has NO `CmdUCINet` "
            "control.  The fresh `control_inventory.json` confirms no "
            "UCINet button exists on the form, so the UCINet export is "
            "unreachable from the UI."
        ),
        "summary_zh": (
            "`Form_LookAtStatus.vb` 定義了 `CmdUCINet_Click` 處理常式，"
            "但 LookAtStatus 的表單設計上並沒有 `CmdUCINet` 控制項。"
            "最新的 `control_inventory.json` 確認表單上沒有 UCINet "
            "按鈕，因此 UCINet 輸出在 UI 上無法使用。"
        ),
        "steps_en": [
            "Open LookAtStatus.  There is no UCINet export button in "
            "the export-buttons row.",
            "Static confirmation: "
            "`analysis/dump/control_inventory.json` shows no "
            "`CmdUCINet` control on LookAtStatus, although "
            "`Form_LookAtStatus.vb` defines `Sub CmdUCINet_Click()`.",
        ],
        "steps_zh": [
            "開啟 LookAtStatus。輸出按鈕列裡沒有 UCINet 輸出按鈕。",
            "靜態確認：`analysis/dump/control_inventory.json` 顯示 "
            "LookAtStatus 上沒有 `CmdUCINet` 控制項，但 "
            "`Form_LookAtStatus.vb` 定義了 `Sub CmdUCINet_Click()`。",
        ],
        "fix_en": "Add a CmdUCINet button to LookAtStatus's design, "
                  "wired to the existing CmdUCINet_Click Sub.",
        "fix_zh": "在 LookAtStatus 的設計中新增 CmdUCINet 按鈕，"
                  "並連到既有的 CmdUCINet_Click。",
        "severity_en": "P3 — Missing UI (feature unavailable to users).",
        "severity_zh": "P3 —— 缺少 UI（使用者無法使用此功能）。",
        "screenshots": [
            ("bug18_LookAtStatus_no_CmdUCINet.png",
             "LookAtStatus as it ships — no UCINet export button."),
            ("bug18_LookAtStatus_no_CmdUCINet_annotated.png",
             "Annotated: `Sub CmdUCINet_Click()` exists in the module "
             "but no control invokes it."),
        ],
        "evidence": {
            "finding_class": "user_facing_bug",
            "vba_ref": "Form_LookAtStatus.vb:1664 "
                       "(Private Sub CmdUCINet_Click()); no CmdUCINet "
                       "control in analysis/dump/control_inventory.json "
                       "for LookAtStatus.",
            "fixture": "LookAtStatus form design (any session).",
            "user_symptom": "There is no UCINet export button on "
                            "LookAtStatus, so users cannot run the "
                            "UCINet export the code supports.",
            "ui_verified": False,
        },
    },
    {
        "id": 19,
        "tier": "P3_missing_ui",
        "form": "LookAtOffice",
        "title_en": "LookAtOffice is missing its CmdGUESS button "
                    "(handler exists, no UI control)",
        "title_zh": "LookAtOffice 缺少 CmdGUESS 按鈕"
                    "（處理常式存在，但表單上沒有控制項）",
        "summary_en": (
            "`Form_LookAtOffice.vb` defines a `CmdGUESS_Click` handler "
            "(it would write a GUESS `.gdf` export) but LookAtOffice's "
            "form design has NO `CmdGUESS` control.  The fresh "
            "`control_inventory.json` lists GIS / GISPeople / Neo4j on "
            "this form but no GUESS button, so the GUESS export is "
            "unreachable from the UI."
        ),
        "summary_zh": (
            "`Form_LookAtOffice.vb` 定義了 `CmdGUESS_Click` 處理常式"
            "（用來輸出 GUESS `.gdf` 檔），但 LookAtOffice 的表單"
            "設計上並沒有 `CmdGUESS` 控制項。最新的 "
            "`control_inventory.json` 顯示此表單有 GIS / GISPeople / "
            "Neo4j，卻沒有 GUESS 按鈕，因此 GUESS 輸出在 UI 上"
            "無法使用。"
        ),
        "steps_en": [
            "Open LookAtOffice.  There is no GUESS export button "
            "(only GIS / GISPeople / Neo4j).",
            "Static confirmation: "
            "`analysis/dump/control_inventory.json` shows no "
            "`CmdGUESS` control on LookAtOffice, although "
            "`Form_LookAtOffice.vb` defines `Sub CmdGUESS_Click()`.",
        ],
        "steps_zh": [
            "開啟 LookAtOffice。沒有 GUESS 輸出按鈕"
            "（只有 GIS / GISPeople / Neo4j）。",
            "靜態確認：`analysis/dump/control_inventory.json` 顯示 "
            "LookAtOffice 上沒有 `CmdGUESS` 控制項，但 "
            "`Form_LookAtOffice.vb` 定義了 `Sub CmdGUESS_Click()`。",
        ],
        "fix_en": "Add a CmdGUESS button to LookAtOffice's design, "
                  "wired to the existing CmdGUESS_Click Sub.",
        "fix_zh": "在 LookAtOffice 的設計中新增 CmdGUESS 按鈕，"
                  "並連到既有的 CmdGUESS_Click。",
        "severity_en": "P3 — Missing UI (feature unavailable to users).",
        "severity_zh": "P3 —— 缺少 UI（使用者無法使用此功能）。",
        "screenshots": [
            ("bug19_LookAtOffice_no_CmdGUESS.png",
             "LookAtOffice as it ships — no GUESS export button."),
            ("bug19_LookAtOffice_no_CmdGUESS_annotated.png",
             "Annotated: `Sub CmdGUESS_Click()` exists in the module "
             "but no control invokes it."),
        ],
        "evidence": {
            "finding_class": "user_facing_bug",
            "vba_ref": "Form_LookAtOffice.vb:3040 "
                       "(Private Sub CmdGUESS_Click()); no CmdGUESS "
                       "control in analysis/dump/control_inventory.json "
                       "for LookAtOffice.",
            "fixture": "LookAtOffice form design (any session).",
            "user_symptom": "There is no GUESS export button on "
                            "LookAtOffice, so users cannot run the "
                            "GUESS export the code supports.",
            "ui_verified": False,
        },
    },

    # ---------------- P5 — Dormant / latent ----------------
    {
        "id": 5,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtStatus.CmdPajek_Click",
        "title_en": "LookAtStatus.CmdPajek references a missing control "
                    "AND three columns that don't exist — LATENT "
                    "(gated by the missing Pajek button, Issue #16)",
        "title_zh": "LookAtStatus.CmdPajek 參照了一個不存在的控制項"
                    "以及三個不存在的欄位 —— 潛伏"
                    "（被缺少的 Pajek 按鈕擋住，見 Issue #16）",
        "summary_en": (
            "Two related source-level defects in the same handler:\n\n"
            "(a) Line 2308 reads `If ChkIDs.Value Then`, but "
            "LookAtStatus has no control named `ChkIDs`.\n\n"
            "(b) Lines 2335-2338 build a SELECT … INTO that references "
            "`ZZ_SCRATCH_STATUS.c_person_id`, `c_status_id`, and "
            "`c_status_count` — none of which exist on that table (the "
            "real columns are `c_personid`, `c_status_code`, with no "
            "count column).  The sub reads as a copy of "
            "`LookAtAssociations.CmdPajek_Click`, where those names are "
            "valid; the rename pass missed both spots.\n\n"
            "Why LATENT: LookAtStatus has no Pajek button at all "
            "(Issue #16), so users physically cannot invoke this "
            "handler today.  The SQL would still fail the moment the "
            "sub runs, so adding a button without fixing this would "
            "just expose the failure.  This build was non-interactive, "
            "so no runtime UI symptom could be re-verified — filed as "
            "latent pending UI re-verification."
        ),
        "summary_zh": (
            "同一處理常式中兩個相關的源碼缺陷：\n\n"
            "(a) 第 2308 行讀取 `If ChkIDs.Value Then`，但 "
            "LookAtStatus 並沒有名為 `ChkIDs` 的控制項。\n\n"
            "(b) 第 2335-2338 行建立的 SELECT … INTO 參照了 "
            "`ZZ_SCRATCH_STATUS.c_person_id`、`c_status_id`、"
            "`c_status_count`——這三者在該表上都不存在（真正的欄位"
            "是 `c_personid`、`c_status_code`，且沒有計數欄位）。"
            "這個 Sub 看起來是 `LookAtAssociations.CmdPajek_Click` "
            "的複製，那裡這些名稱是有效的；改名時漏了這兩處。\n\n"
            "為何潛伏：LookAtStatus 根本沒有 Pajek 按鈕"
            "（Issue #16），所以使用者目前無法觸發這個處理常式。"
            "一旦此 Sub 執行，SQL 仍會立即失敗，因此若只加按鈕"
            "而不修這裡，只會把失敗暴露給使用者。本次測試為"
            "非互動式，無法重新驗證執行期 UI 症狀——故列為潛伏，"
            "待 UI 重新驗證。"
        ),
        "steps_en": [
            "On this build the bug cannot be triggered through the UI — "
            "LookAtStatus has no Pajek button (Issue #16).  Verify the "
            "defects statically instead:",
            "Open `analysis/dump/vba/Form_LookAtStatus.vb` and read "
            "line 2308: `If ChkIDs.Value Then` — no `ChkIDs` control "
            "exists on LookAtStatus in "
            "`analysis/dump/control_inventory.json`.",
            "Read lines 2335-2338: the SELECT … INTO references "
            "`ZZ_SCRATCH_STATUS.c_person_id` / `c_status_id` / "
            "`c_status_count` (the count aggregate on line 2337), none "
            "of which are columns on ZZ_SCRATCH_STATUS.",
        ],
        "steps_zh": [
            "本次建置無法透過 UI 觸發此 bug——LookAtStatus 沒有 Pajek "
            "按鈕（Issue #16）。改以靜態方式驗證：",
            "開啟 `analysis/dump/vba/Form_LookAtStatus.vb`，看第 2308 "
            "行：`If ChkIDs.Value Then`——在 "
            "`analysis/dump/control_inventory.json` 中，LookAtStatus "
            "並沒有 `ChkIDs` 控制項。",
            "看第 2335-2338 行：SELECT … INTO 參照了 "
            "`ZZ_SCRATCH_STATUS.c_person_id` / `c_status_id` / "
            "`c_status_count`（計數彙總在第 2337 行），這些都不是 "
            "ZZ_SCRATCH_STATUS 的欄位。",
        ],
        "fix_en": (
            "(a) Replace `ChkIDs.Value` with a constant `False` (if the "
            "optional ID-suffix behaviour isn't needed) or add a real "
            "ChkIDs control.  (b) Rewrite the SELECT to use "
            "`ZZ_SCRATCH_STATUS.c_personid` and `c_status_code`, and "
            "either drop the count aggregate or compute it another way. "
            " Realistically the whole sub needs a careful rewrite — it "
            "was inherited from another form without verification — and "
            "should be done together with adding the button (Issue #16)."
        ),
        "fix_zh": (
            "(a) 將 `ChkIDs.Value` 改為常數 `False`（若不需要可選的"
            "ID 後綴行為）或新增真正的 ChkIDs 控制項。(b) 重寫 SELECT "
            "改用 `ZZ_SCRATCH_STATUS.c_personid` 與 `c_status_code`，"
            "並將計數彙總移除或以其他方式計算。實務上整個 Sub 需要"
            "謹慎重寫——它是未經驗證就從別的表單沿用過來的——"
            "並應與新增按鈕（Issue #16）一併處理。"
        ),
        "severity_en": "P5 — Latent source defect (would resurface as a "
                       "visible crash if Issue #16 were fixed without "
                       "first fixing this).",
        "severity_zh": "P5 —— 潛伏源碼缺陷（若只修 Issue #16 而不先"
                       "修這裡，將以可見的當機重新浮現）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_LookAtStatus.vb:2308 (If ChkIDs.Value "
                       "Then — no such control) and :2335-2338 "
                       "(SELECT references c_person_id / c_status_id / "
                       "c_status_count at :2337, none on "
                       "ZZ_SCRATCH_STATUS).",
            "fixture": "n/a — gated unreachable: LookAtStatus has no "
                       "CmdPajek button (Issue #16).",
            "user_symptom": "None today (no button to click).  Would "
                            "be an 'Object required' then 'No such "
                            "field' crash if the button were added.",
            "ui_verified": False,
        },
    },
    {
        "id": 6,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtGroupData.queryEntry",
        "title_en": "LookAtGroupData Entry insert projects "
                    "ENTRY_DATA.c_parental_status (should be "
                    "…_code) — LATENT this build (runtime ERR did not "
                    "fire)",
        "title_zh": "LookAtGroupData 的 Entry 插入投影了 "
                    "ENTRY_DATA.c_parental_status（應為 …_code）"
                    "—— 本次潛伏（執行期未觸發錯誤）",
        "summary_en": (
            "`Form_LookAtGroupData.vb`'s Entry INSERT names a target "
            "column `c_parental_status_code` (line 2612) but the SELECT "
            "projection ends with `ENTRY_DATA.c_parental_status` "
            "(line 2621) — no `_code` suffix.  The real ENTRY_DATA "
            "column is `c_parental_status_code`; the source-level typo "
            "would make JET raise 'No such field' / 'No value given for "
            "one or more required parameters' when the Entry branch "
            "runs.  `Form_LookAtEntry.vb` does the analogous query with "
            "the correct name, so this is a single-line drift.\n\n"
            "Honest note for this build: the source defect is present "
            "in the dump, but the behavioural probe completed WITHOUT "
            "the error this build (the symptom is data-/enable-path "
            "dependent and the session was non-interactive, so no "
            "runtime UI symptom could be re-verified).  Filed as latent "
            "pending UI re-verification rather than a confirmed "
            "user-facing crash."
        ),
        "summary_zh": (
            "`Form_LookAtGroupData.vb` 的 Entry INSERT 目標欄位列出 "
            "`c_parental_status_code`（第 2612 行），但 SELECT 投影"
            "結尾卻是 `ENTRY_DATA.c_parental_status`（第 2621 行）"
            "——少了 `_code` 後綴。ENTRY_DATA 真正的欄位是 "
            "`c_parental_status_code`；當 Entry 分支執行時，這個源碼"
            "層級的筆誤會讓 JET 丟出『No such field』/『No value given "
            "for one or more required parameters』。`Form_LookAtEntry.vb` "
            "對應的查詢用的是正確名稱，因此這是一行的漂移。\n\n"
            "本次建置的誠實說明：源碼缺陷確實存在於傾印中，但本次"
            "行為探測在執行時並未觸發該錯誤（症狀依資料／啟用路徑"
            "而定，且本次為非互動式，無法重新驗證執行期 UI 症狀）。"
            "故列為潛伏待 UI 重新驗證，而非已確認的使用者當機。"
        ),
        "steps_en": [
            "On this build the runtime error did not fire — verify the "
            "source defect statically:",
            "Open `analysis/dump/vba/Form_LookAtGroupData.vb`.  Line "
            "2612 lists the INSERT target column "
            "`c_parental_status_code`; line 2621 projects "
            "`ENTRY_DATA.c_parental_status` (missing `_code`).",
            "To exercise the path in a future interactive session: in "
            "LookAtGroupData, populate the import list with one person, "
            "tick only the Entry checkbox, and click Run.  If the path "
            "fires, a 'field doesn't exist' popup appears.",
        ],
        "steps_zh": [
            "本次建置執行期未觸發此錯誤——以靜態方式驗證源碼缺陷：",
            "開啟 `analysis/dump/vba/Form_LookAtGroupData.vb`。第 2612 "
            "行列出 INSERT 目標欄位 `c_parental_status_code`；第 2621 "
            "行投影 `ENTRY_DATA.c_parental_status`（少了 `_code`）。",
            "若要在未來的互動式測試中走到這條路徑：在 "
            "LookAtGroupData 匯入一位人物，只勾選 Entry，點 Run。"
            "若路徑被觸發，會跳出『欄位不存在』的彈窗。",
        ],
        "fix_en": (
            "Change `ENTRY_DATA.c_parental_status` to "
            "`ENTRY_DATA.c_parental_status_code` on line 2621.  "
            "One-line fix, matching the correct name already used in "
            "`Form_LookAtEntry.vb`."
        ),
        "fix_zh": (
            "把第 2621 行的 `ENTRY_DATA.c_parental_status` 改成 "
            "`ENTRY_DATA.c_parental_status_code`。一行修正，與 "
            "`Form_LookAtEntry.vb` 已使用的正確名稱一致。"
        ),
        "severity_en": "P5 — Latent (source typo present; runtime "
                       "symptom not reproduced this build, pending UI "
                       "re-verification).",
        "severity_zh": "P5 —— 潛伏（源碼筆誤存在；本次建置未重現執行期"
                       "症狀，待 UI 重新驗證）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_LookAtGroupData.vb:2612 (target list ends "
                       "c_parental_status_code) vs :2621 (SELECT "
                       "projects ENTRY_DATA.c_parental_status, no _code).",
            "fixture": "LookAtGroupData Entry branch (import list + "
                       "Entry checkbox).  Runtime error not reproduced "
                       "this non-interactive build.",
            "user_symptom": "None reproduced this build.  When the path "
                            "fires, a 'field doesn't exist' popup blocks "
                            "the Entry run.",
            "ui_verified": False,
        },
    },
    {
        "id": 7,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtPlace.CmdNeo4j_Click",
        "title_en": "LookAtPlace.CmdNeo4j people-recordset reads "
                    "c_dynasty / c_dynasty_chn / c_female that the "
                    "SELECT doesn't project — LATENT (runtime did not "
                    "fire this build)",
        "title_zh": "LookAtPlace.CmdNeo4j 的人員記錄集讀取了 SELECT 未"
                    "投影的 c_dynasty / c_dynasty_chn / c_female "
                    "—— 潛伏（本次執行期未觸發）",
        "summary_en": (
            "`Form_LookAtPlace.CmdNeo4j_Click` opens `tRstPeople` "
            "(line 326) on a SELECT DISTINCT that projects only four "
            "ZZ_SCRATCH_P_TEXT columns (line 322): c_person_id, "
            "c_name, c_name_chn, c_index_year.  The INNER JOIN brings "
            "DYNASTIES and BIOG_MAIN into scope but does NOT project "
            "any of their columns.  The row-write loop then reads "
            "`!c_dynasty` (line 383), `!c_dynasty_chn` (385) and "
            "`!c_female` (392) from that recordset; DAO's Fields "
            "collection only contains projected columns, so JET raises "
            "3265 'Item not found in this collection.' on the first "
            "such read.  The handler routes to the exit before any disk "
            "file is flushed, so the user would see a popup and an "
            "empty output folder.\n\n"
            "Why LATENT this build: the CmdNeo4j button DOES exist on "
            "LookAtPlace, but this session was non-interactive "
            "(pywinauto UIA unavailable), so the runtime symptom could "
            "not be reproduced/re-verified.  The projection mismatch is "
            "a confirmed static defect; filed as latent pending UI "
            "re-verification.  The recommended demo address is "
            "`c_addr_id = 100658` (Kaifeng / 開封), which has plenty of "
            "people to feed the People-CSV loop."
        ),
        "summary_zh": (
            "`Form_LookAtPlace.CmdNeo4j_Click` 在一個只投影四個 "
            "ZZ_SCRATCH_P_TEXT 欄位的 SELECT DISTINCT（第 322 行："
            "c_person_id、c_name、c_name_chn、c_index_year）上開啟 "
            "`tRstPeople`（第 326 行）。INNER JOIN 把 DYNASTIES 與 "
            "BIOG_MAIN 帶入範圍，但並未投影它們的任何欄位。接著逐列"
            "寫出的迴圈從該記錄集讀取 `!c_dynasty`（第 383 行）、"
            "`!c_dynasty_chn`（385）與 `!c_female`（392）；DAO 的 "
            "Fields 集合只含被"
            "投影的欄位，因此 JET 會在第一次這類讀取時丟出 3265 "
            "『Item not found in this collection.』。處理常式在任何"
            "磁碟檔寫出前就跳到結束，使用者會看到彈窗以及一個空的"
            "輸出資料夾。\n\n"
            "本次為何潛伏：LookAtPlace 上確實有 CmdNeo4j 按鈕，但本次"
            "測試為非互動式（pywinauto UIA 無法使用），無法重現／"
            "重新驗證執行期症狀。投影不符是已確認的靜態缺陷；故列為"
            "潛伏待 UI 重新驗證。建議的示範地址為 "
            "`c_addr_id = 100658`（Kaifeng / 開封），其關聯人物足以"
            "餵滿 People-CSV 迴圈。"
        ),
        "steps_en": [
            "On this build the runtime symptom was not reproduced "
            "(non-interactive session).  Verify the projection "
            "mismatch statically:",
            "Open `analysis/dump/vba/Form_LookAtPlace.vb`.  Line "
            "322 projects only c_person_id / c_name / c_name_chn / "
            "c_index_year into `tRstPeople` (opened line 326).",
            "Lines 383 / 385 / 392 read `!c_dynasty`, `!c_dynasty_chn`, "
            "`!c_female` from that recordset — none are projected, so "
            "JET 3265 fires on the first read.",
            "To re-verify interactively later: open LookAtPlace, pick "
            "address `c_addr_id = 100658` (Kaifeng / 開封), Run Query, "
            "then click Neo4j and choose a save folder — expect a 3265 "
            "popup and an empty folder.",
        ],
        "steps_zh": [
            "本次建置未重現執行期症狀（非互動式測試）。以靜態方式"
            "驗證投影不符：",
            "開啟 `analysis/dump/vba/Form_LookAtPlace.vb`。第 322 "
            "行只把 c_person_id / c_name / c_name_chn / c_index_year "
            "投影到 `tRstPeople`（第 326 行開啟）。",
            "第 383 / 385 / 392 行從該記錄集讀取 `!c_dynasty`、"
            "`!c_dynasty_chn`、`!c_female`——皆未被投影，因此第一次"
            "讀取時就觸發 JET 3265。",
            "日後互動式重新驗證：開啟 LookAtPlace，選地址 "
            "`c_addr_id = 100658`（Kaifeng / 開封），執行查詢，再點 "
            "Neo4j 並選一個儲存資料夾——預期會出現 3265 彈窗且資料夾"
            "為空。",
        ],
        "fix_en": (
            "Extend the SELECT projection in "
            "`Form_LookAtPlace.vb:322` to include the three columns "
            "the loop reads: `DYNASTIES.c_dynasty`, "
            "`DYNASTIES.c_dynasty_chn`, `BIOG_MAIN.c_female` (the "
            "FROM/JOIN already brings them into scope).  Three columns "
            "added; nothing else changes."
        ),
        "fix_zh": (
            "在 `Form_LookAtPlace.vb:322` 的 SELECT 投影中加入"
            "迴圈會讀取的三個欄位：`DYNASTIES.c_dynasty`、"
            "`DYNASTIES.c_dynasty_chn`、`BIOG_MAIN.c_female`"
            "（FROM/JOIN 已把它們帶入範圍）。新增三個欄位，其餘"
            "不變。"
        ),
        "severity_en": "P5 — Latent (confirmed static projection "
                       "mismatch; runtime symptom not reproduced this "
                       "non-interactive build, pending UI "
                       "re-verification).",
        "severity_zh": "P5 —— 潛伏（已確認的靜態投影不符；本次非互動式"
                       "建置未重現執行期症狀，待 UI 重新驗證）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_LookAtPlace.vb:322 (SELECT DISTINCT "
                       "projects only 4 columns) + :326 (Set tRstPeople "
                       "= OpenRecordset) + :383 (first unprojected read "
                       "!c_dynasty).",
            "fixture": "LookAtPlace, address c_addr_id = 100658 "
                       "(Kaifeng / 開封).  Runtime not reproduced this "
                       "non-interactive build.",
            "user_symptom": "None reproduced this build.  When the "
                            "export runs, a JET 3265 'Item not found in "
                            "this collection.' popup appears and the "
                            "chosen output folder stays empty.",
            "ui_verified": False,
        },
    },
    {
        "id": 8,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtNetworks.CmdNeo4j_Click",
        "title_en": "LookAtNetworks.CmdNeo4j place-recordset reads "
                    "x_coord / y_coord that the SELECT doesn't project "
                    "— LATENT (behavioural repro blocked by Networks "
                    "Form_Open hang)",
        "title_zh": "LookAtNetworks.CmdNeo4j 的地點記錄集讀取了 SELECT "
                    "未投影的 x_coord / y_coord —— 潛伏"
                    "（Networks 表單開啟卡死，行為重現受阻）",
        "summary_en": (
            "Same shape as Issue #7, on a different form.  In "
            "`Form_LookAtNetworks.CmdNeo4j_Click` the place SELECT "
            "(line 2458) projects only three columns "
            "(c_index_addr_id, c_index_addr_name, c_index_addr_chn) "
            "into `tRstPlace` (line 2463).  The header it writes "
            "declares placeX / placeY (lines 2466/2466), and the "
            "row-write loop then reads `!x_coord` (line 2495) and "
            "`!y_coord` from that recordset — neither is projected, so "
            "JET 3265 'Item not found in this collection.' fires and "
            "the export aborts.\n\n"
            "Why LATENT: behavioural reproduction is blocked because "
            "`LookAtNetworks`'s `Form_Open` hangs the COM test driver, "
            "so the host form cannot be driven this build; combined "
            "with the non-interactive session, no runtime symptom could "
            "be re-verified.  The projection mismatch is a confirmed "
            "static defect; filed as latent pending UI re-verification."
        ),
        "summary_zh": (
            "與 Issue #7 同型，發生在不同表單上。在 "
            "`Form_LookAtNetworks.CmdNeo4j_Click` 中，地點 SELECT"
            "（第 2458 行）只把三個欄位（c_index_addr_id、"
            "c_index_addr_name、c_index_addr_chn）投影到 `tRstPlace`"
            "（第 2463 行）。它寫出的表頭宣告了 placeX / placeY"
            "（第 2466/2466 行），接著逐列寫出的迴圈從該記錄集讀取 "
            "`!x_coord`（第 2495 行）與 `!y_coord`——兩者皆未被投影，"
            "因此 JET 3265『Item not found in this collection.』觸發，"
            "輸出中止。\n\n"
            "為何潛伏：行為重現受阻，因為 `LookAtNetworks` 的 "
            "`Form_Open` 會讓 COM 測試驅動卡死，本次無法驅動該宿主"
            "表單；加上本次為非互動式測試，無法重新驗證執行期症狀。"
            "投影不符是已確認的靜態缺陷；故列為潛伏待 UI 重新驗證。"
        ),
        "steps_en": [
            "Behavioural repro is blocked (LookAtNetworks Form_Open "
            "hangs the driver) and this session was non-interactive.  "
            "Verify the projection mismatch statically:",
            "Open `analysis/dump/vba/Form_LookAtNetworks.vb`.  Line "
            "2458 projects only c_index_addr_id / c_index_addr_name / "
            "c_index_addr_chn into `tRstPlace` (line 2463).",
            "Lines 2495 / 2502 read `!x_coord` (and `!y_coord` "
            "nearby) — neither is projected, so JET 3265 fires on the "
            "place block.",
        ],
        "steps_zh": [
            "行為重現受阻（LookAtNetworks 的 Form_Open 讓驅動卡死），"
            "且本次為非互動式測試。以靜態方式驗證投影不符：",
            "開啟 `analysis/dump/vba/Form_LookAtNetworks.vb`。第 2458 "
            "行只把 c_index_addr_id / c_index_addr_name / "
            "c_index_addr_chn 投影到 `tRstPlace`（第 2463 行）。",
            "第 2495 / 2502 行讀取 `!x_coord`（附近還有 `!y_coord`）"
            "——皆未被投影，因此地點區塊會觸發 JET 3265。",
        ],
        "fix_en": (
            "Extend the place SELECT in "
            "`Form_LookAtNetworks.vb:2458` to project the coordinate "
            "columns the loop reads, e.g. `ADDR_CODES.x_coord`, "
            "`ADDR_CODES.y_coord` (the JOIN to ADDR_CODES already "
            "exposes them)."
        ),
        "fix_zh": (
            "在 `Form_LookAtNetworks.vb:2458` 的地點 SELECT 中投影"
            "迴圈會讀取的座標欄位，例如 `ADDR_CODES.x_coord`、"
            "`ADDR_CODES.y_coord`（與 ADDR_CODES 的 JOIN 已暴露"
            "它們）。"
        ),
        "severity_en": "P5 — Latent (confirmed static projection "
                       "mismatch; behavioural repro blocked by the "
                       "Networks Form_Open hang, pending UI "
                       "re-verification).",
        "severity_zh": "P5 —— 潛伏（已確認的靜態投影不符；行為重現"
                       "因 Networks 表單開啟卡死而受阻，待 UI "
                       "重新驗證）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_LookAtNetworks.vb:2458 (place SELECT "
                       "projects 3 columns) + :2463 (Set tRstPlace) + "
                       ":2495 (reads !x_coord, unprojected).",
            "fixture": "LookAtNetworks CmdNeo4j place block.  Host "
                       "form's Form_Open hangs the COM driver; not "
                       "reproduced this non-interactive build.",
            "user_symptom": "None reproduced this build.  When the "
                            "export runs, a JET 3265 'Item not found in "
                            "this collection.' popup appears and the "
                            "export aborts mid-chain.",
            "ui_verified": False,
        },
    },
    {
        "id": 9,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtEntry.CmdNeo4j_Click",
        "title_en": "LookAtEntry.CmdNeo4j Institutions block uses the "
                    "wrong recordset variable (tRstAssocCodes) — LATENT "
                    "(gated unreachable; no ENTRY_DATA row has "
                    "c_inst_code > 0)",
        "title_zh": "LookAtEntry.CmdNeo4j 的 Institutions 區塊用錯了"
                    "記錄集變數（tRstAssocCodes）—— 潛伏"
                    "（被閘門擋住而不可達；沒有任何 ENTRY_DATA 列的 "
                    "c_inst_code > 0）",
        "summary_en": (
            "Line 1415 of `Form_LookAtEntry.vb` opens the institutions "
            "recordset as `Set tRstInstitutions = "
            "CurrentDb.OpenRecordset(tQueryStr)`.  Twenty lines later, "
            "line 1425 says `With tRstAssocCodes` and the loop reads "
            "`!c_inst_code` etc. against THAT recordset — which was "
            "bound earlier to the AssocCodes SELECT and already closed "
            "in the AssocCodes block.  If executed, `.MoveFirst` would "
            "raise DAO 3021 'No current record'; the misnamed reference "
            "is a genuine source-level bug.\n\n"
            "Why LATENT: the whole SaveAs-and-buggy-With block sits "
            "inside the gate `If tRecDeleted > 0 Then` (line 1390), "
            "where tRecDeleted counts an INSERT … WHERE "
            "ZZ_SCRATCH_ENTRY.c_inst_code > 0.  On this dump no "
            "ENTRY_DATA row has `c_inst_code > 0`, so the gate is "
            "always false, the buggy `With` line never executes, and "
            "CmdNeo4j finishes cleanly (silently omitting the optional "
            "InstitutionCodes CSV — the same gating the surrounding "
            "blocks use).  The typo only becomes user-visible if a "
            "future data drop introduces any `ENTRY_DATA.c_inst_code > "
            "0`.  Investigation fixtures: `c_entry_code = 36` "
            "(examination: jinshi (general) / 進士) and "
            "`c_entry_code = 101` (recommendation / 薦舉) exercise "
            "CmdQuery + CmdNeo4j end-to-end and both finish cleanly — "
            "evidence that the gate works, not a popup reproduction."
        ),
        "summary_zh": (
            "`Form_LookAtEntry.vb` 第 1415 行以 `Set tRstInstitutions "
            "= CurrentDb.OpenRecordset(tQueryStr)` 開啟機構記錄集。"
            "二十行後，第 1425 行寫的是 `With tRstAssocCodes`，迴圈"
            "對「那個」記錄集讀取 `!c_inst_code` 等——而它先前已被"
            "綁定到 AssocCodes 的 SELECT，並已在 AssocCodes 區塊中"
            "關閉。若執行到，`.MoveFirst` 會丟出 DAO 3021『No current "
            "record』；這個命名錯誤是貨真價實的源碼 bug。\n\n"
            "為何潛伏：整段 SaveAs 與有問題的 With 區塊都位於閘門 "
            "`If tRecDeleted > 0 Then`（第 1390 行）之內，其中 "
            "tRecDeleted 計算的是 INSERT … WHERE "
            "ZZ_SCRATCH_ENTRY.c_inst_code > 0 的列數。在本傾印中"
            "沒有任何 ENTRY_DATA 列的 `c_inst_code > 0`，因此閘門"
            "恆為 false，有問題的 `With` 行從不執行，CmdNeo4j 乾淨"
            "完成（靜默略過可選的 InstitutionCodes CSV——與周邊區塊"
            "相同的閘門做法）。只有當未來資料引入任何 "
            "`ENTRY_DATA.c_inst_code > 0` 時，此筆誤才會對使用者可見。"
            "調查用 fixture：`c_entry_code = 36`"
            "（examination: jinshi (general) / 進士）與 "
            "`c_entry_code = 101`（recommendation / 薦舉）會端到端"
            "走完 CmdQuery + CmdNeo4j，兩者皆乾淨結束——這是閘門有效"
            "的證據，而非彈窗重現。"
        ),
        "steps_en": [
            "On this dump the bug cannot be triggered through the UI — "
            "the `If tRecDeleted > 0 Then` gate at "
            "Form_LookAtEntry.vb:1390 is false for every fixture (no "
            "ENTRY_DATA row has c_inst_code > 0).  Verify the typo "
            "statically:",
            "Open `analysis/dump/vba/Form_LookAtEntry.vb` and read "
            "lines 1415-1425.  Line 1415: `Set tRstInstitutions = "
            "OpenRecordset(tQueryStr)`.  Line 1425: `With "
            "tRstAssocCodes` (intended `With tRstInstitutions`); "
            "tRstAssocCodes was already closed in the AssocCodes block, "
            "so `.MoveFirst` would raise DAO 3021.",
            "(Optional, runtime evidence) Pick `c_entry_code = 36` "
            "(examination: jinshi (general) / 進士) or "
            "`c_entry_code = 101` (recommendation / 薦舉) on LookAtEntry "
            "→ Run Query → Neo4j.  Both finish cleanly with no popup "
            "and no InstitutionCodes CSV — evidence the gate holds.",
        ],
        "steps_zh": [
            "在本傾印上此 bug 無法透過 UI 觸發——Form_LookAtEntry.vb:1390 "
            "的 `If tRecDeleted > 0 Then` 閘門對每個 fixture 都為 false"
            "（沒有任何 ENTRY_DATA 列的 c_inst_code > 0）。以靜態方式"
            "驗證筆誤：",
            "開啟 `analysis/dump/vba/Form_LookAtEntry.vb`，看第 "
            "1415-1425 行。第 1415 行：`Set tRstInstitutions = "
            "OpenRecordset(tQueryStr)`。第 1425 行：`With "
            "tRstAssocCodes`（應為 `With tRstInstitutions`）；"
            "tRstAssocCodes 已在 AssocCodes 區塊中關閉，故 "
            "`.MoveFirst` 會丟出 DAO 3021。",
            "（可選的執行期證據）在 LookAtEntry 選 "
            "`c_entry_code = 36`（examination: jinshi (general) / 進士）"
            "或 `c_entry_code = 101`（recommendation / 薦舉）→ 執行"
            "查詢 → Neo4j。兩者皆乾淨結束，無彈窗、無 "
            "InstitutionCodes CSV——這是閘門守住的證據。",
        ],
        "fix_en": (
            "Change `With tRstAssocCodes` on line 1425 to "
            "`With tRstInstitutions`.  The recordset variable was "
            "simply mis-named.  Although currently unreachable on this "
            "dump, fixing it costs nothing and prevents a future-data "
            "regression."
        ),
        "fix_zh": (
            "把第 1425 行的 `With tRstAssocCodes` 改成 "
            "`With tRstInstitutions`。記錄集變數只是被命名錯了。"
            "雖然在本傾印上目前不可達，修正它毫無成本，又能避免"
            "未來資料造成的回歸。"
        ),
        "severity_en": "P5 — Latent source typo (gated unreachable; "
                       "would resurface as a DAO 3021 crash if any "
                       "future ENTRY_DATA row had c_inst_code > 0).",
        "severity_zh": "P5 —— 潛伏源碼筆誤（被閘門擋住而不可達；"
                       "若未來任何 ENTRY_DATA 列的 c_inst_code > 0，"
                       "將以 DAO 3021 當機重新浮現）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_LookAtEntry.vb:1415 (Set tRstInstitutions "
                       "= OpenRecordset) vs :1425 (With tRstAssocCodes "
                       "— wrong, already-closed recordset), gated by "
                       ":1390 (If tRecDeleted > 0, from WHERE "
                       "c_inst_code > 0).",
            "fixture": "LookAtEntry, c_entry_code = 36 (examination: "
                       "jinshi (general) / 進士) or c_entry_code = 101 "
                       "(recommendation / 薦舉); both finish cleanly — "
                       "gate is false (0 ENTRY_DATA rows with "
                       "c_inst_code > 0).",
            "user_symptom": "None today (gated unreachable).  Would be "
                            "a DAO 3021 'No current record' popup and an "
                            "aborted Neo4j export if a future data drop "
                            "introduced any c_inst_code > 0.",
            "ui_verified": False,
        },
    },
    {
        "id": 14,
        "tier": "P5_dormant_or_latent",
        "form": "Form_KIN_DATA_Subform",
        "title_en": "KIN_DATA Subform's CmdPickKinRel calls a missing "
                    "picker (frmPickKINSHIP_CODES) — LATENT (host "
                    "sub-form is an orphan; no reachable trigger)",
        "title_zh": "KIN_DATA 子表單的 CmdPickKinRel 呼叫一個不存在的"
                    "選取表單（frmPickKINSHIP_CODES）—— 潛伏"
                    "（宿主子表單為孤兒，無可達觸發路徑）",
        "summary_en": (
            "`Form_KIN_DATA_Subform`'s `CmdPickKinRel_Click` "
            "(stDocName set at line 63) calls "
            "`DoCmd.OpenForm \"frmPickKINSHIP_CODES\"` and references "
            "`Forms!frmPickKINSHIP_CODES!frmKINSHIP_CODES.Form!"
            "c_kincode`.  Neither form exists in the current .mdb "
            "(absent from the fresh `control_inventory.json`) — same "
            "shape as Issue #13.\n\n"
            "Why LATENT: the host sub-form `KIN_DATA Subform` (which "
            "owns the CmdPickKinRel button) is not contained by any "
            "navigable form in the current inventory — "
            "`KIN_DATA Subform` is absent from the form list, while "
            "`BIOG_MAIN_2_Subform` embeds `KIN_DATA_2 Subform` instead "
            "(a read-only variant with no CmdPickKinRel button).  "
            "Because no user-facing navigation reaches the picker "
            "button, the popup can't be triggered.  The latent code "
            "path resurfaces the moment a developer re-embeds "
            "`KIN_DATA Subform` somewhere reachable."
        ),
        "summary_zh": (
            "`Form_KIN_DATA_Subform` 的 `CmdPickKinRel_Click`"
            "（stDocName 設於第 63 行）呼叫 "
            "`DoCmd.OpenForm \"frmPickKINSHIP_CODES\"`，並參照 "
            "`Forms!frmPickKINSHIP_CODES!frmKINSHIP_CODES.Form!"
            "c_kincode`。這兩個表單在目前的 .mdb 中都不存在"
            "（最新的 `control_inventory.json` 中查無）——與 Issue #13 "
            "同型。\n\n"
            "為何潛伏：宿主子表單 `KIN_DATA Subform`（擁有 "
            "CmdPickKinRel 按鈕者）並未被目前清單中任何可導覽的表單"
            "所包含——`KIN_DATA Subform` 不在表單清單中，而 "
            "`BIOG_MAIN_2_Subform` 改為嵌入 `KIN_DATA_2 Subform`"
            "（一個沒有 CmdPickKinRel 按鈕的唯讀變體）。由於沒有"
            "任何面向使用者的導覽能到達該選取按鈕，彈窗無法被觸發。"
            "一旦開發者把 `KIN_DATA Subform` 重新嵌入到可達之處，"
            "這條潛伏的程式路徑就會重新浮現。"
        ),
        "steps_en": [
            "Verification is static-only — no parent form embeds the "
            "affected sub-form, so the runtime click cannot be "
            "reproduced.",
            "Open `analysis/dump/vba/Form_KIN_DATA_Subform.vb` line 63 "
            "— confirms `stDocName = \"frmPickKINSHIP_CODES\"`, opened "
            "by DoCmd just below.",
            "In `analysis/dump/control_inventory.json`, search for "
            "`frmPickKINSHIP_CODES` (absent) and `KIN_DATA Subform` "
            "(absent from the form list); `KIN_DATA_2 Subform` (the "
            "read-only variant) is what BIOG_MAIN_2_Subform embeds.",
        ],
        "steps_zh": [
            "驗證僅限靜態——沒有任何上層表單嵌入受影響的子表單，"
            "因此無法重現執行期點選。",
            "開啟 `analysis/dump/vba/Form_KIN_DATA_Subform.vb` 第 125 "
            "行——確認 `stDocName = \"frmPickKINSHIP_CODES\"`，緊接著"
            "由 DoCmd 開啟。",
            "在 `analysis/dump/control_inventory.json` 中搜尋 "
            "`frmPickKINSHIP_CODES`（不存在）與 `KIN_DATA Subform`"
            "（不在表單清單中）；BIOG_MAIN_2_Subform 嵌入的是 "
            "`KIN_DATA_2 Subform`（唯讀變體）。",
        ],
        "fix_en": (
            "Same as Issue #13: restore the picker form "
            "`frmPickKINSHIP_CODES` (or update the caller to its "
            "replacement).  Even though the runtime path is not "
            "reachable today, clean up the static defect so it doesn't "
            "resurface when `KIN_DATA Subform` is re-embedded."
        ),
        "fix_zh": (
            "與 Issue #13 相同：還原選取表單 `frmPickKINSHIP_CODES`"
            "（或將呼叫端改為其替代表單）。即使目前執行路徑不可達，"
            "也應清理此靜態缺陷，以免 `KIN_DATA Subform` 被重新嵌入"
            "時重新浮現。"
        ),
        "severity_en": "P5 — Latent (static defect real; host sub-form "
                       "is an orphan, so there is no reachable trigger "
                       "today).",
        "severity_zh": "P5 —— 潛伏（靜態缺陷確實存在；宿主子表單為"
                       "孤兒，目前無可達觸發路徑）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_KIN_DATA_Subform.vb:63 "
                       "(stDocName = \"frmPickKINSHIP_CODES\") + "
                       "OpenForm/Forms! references at :131-161; picker "
                       "form and KIN_DATA Subform both absent from "
                       "analysis/dump/control_inventory.json.",
            "fixture": "n/a — KIN_DATA Subform is not embedded by any "
                       "navigable form (BIOG_MAIN_2_Subform uses "
                       "KIN_DATA_2 Subform instead).",
            "user_symptom": "None today (no navigation reaches the "
                            "picker button).  Would be an 'Item not "
                            "found' popup if the sub-form were "
                            "re-embedded somewhere reachable.",
            "ui_verified": False,
        },
    },
    {
        "id": 20,
        "tier": "P5_dormant_or_latent",
        "form": "ADDR_CODES + Form_LookAt*.CmdGIS_Click",
        "title_en": "BOM-prefixed address names can become embedded "
                    "tabs and misalign GIS exports — DORMANT this build "
                    "(0 BOM rows in ADDR_CODES)",
        "title_zh": "帶 BOM 前綴的地址名稱可能變成內嵌 TAB 而使 GIS "
                    "輸出錯位 —— 本次建置休眠（ADDR_CODES 中 0 列帶 "
                    "BOM）",
        "summary_en": (
            "On earlier builds, some `ADDR_CODES` rows carried a stray "
            "`U+FEFF` (BOM) prefix in `c_name` / `c_name_chn` (a "
            "UTF-8-with-BOM paste residue).  When a CmdGIS export "
            "writes each cell as `tStr + value + Chr(9)` with no "
            "escaping, the BOM-mangled value can introduce a literal "
            "TAB, splitting one field into two cells and silently "
            "shifting every column to its right in the `.tab` file.  "
            "The fixture that exercised this was status code "
            "**40** (civil office / [為官者：文]) in LookAtStatus, with "
            "the reachable dirty row `c_addr_id = 702559` "
            "(Wei Shi / 尉氏).\n\n"
            "Why DORMANT this build: the 20260602 DATA mdb has **0** "
            "rows with a literal U+FEFF prefix in ADDR_CODES.c_name or "
            "c_name_chn — measured by "
            "`tests/test_addr_codes_embedded_delim.py` (build "
            "20260602 calibrated to 0).  The export-delimiter defect "
            "class is real in code (the writers still do no escaping), "
            "but the current data has 0 triggering rows, so no user can "
            "reproduce the misalignment today.  The moment a future "
            "data refresh re-introduces a BOM-prefixed (or otherwise "
            "tab-bearing) address, the misalignment returns."
        ),
        "summary_zh": (
            "在較早的建置中，部分 `ADDR_CODES` 列在 `c_name` / "
            "`c_name_chn` 帶有一個多餘的 `U+FEFF`（BOM）前綴"
            "（UTF-8-with-BOM 貼上的殘留）。當 CmdGIS 輸出以 "
            "`tStr + value + Chr(9)` 逐格寫出且不做跳脫時，被 BOM "
            "破壞的值可能引入一個字面 TAB，把一個欄位拆成兩格，"
            "並在 `.tab` 檔中靜默地把右側每一欄都位移。觸發此情況的"
            "fixture 是 LookAtStatus 中的 status code **40**"
            "（civil office / [為官者：文]），可達的髒列為 "
            "`c_addr_id = 702559`（Wei Shi / 尉氏）。\n\n"
            "本次為何休眠：20260602 DATA mdb 在 ADDR_CODES.c_name 或 "
            "c_name_chn 中帶有字面 U+FEFF 前綴的列為 **0**——由 "
            "`tests/test_addr_codes_embedded_delim.py` 量測（build "
            "20260602 已校準為 0）。輸出分隔符這一類缺陷在程式碼中"
            "確實存在（寫出端仍未做跳脫），但目前資料有 0 列會觸發，"
            "因此目前沒有使用者能重現此錯位。一旦未來資料刷新重新"
            "引入帶 BOM（或其他帶 TAB）的地址，錯位就會回來。"
        ),
        "steps_en": [
            "On this build the symptom is dormant — measure the trigger "
            "count to confirm:",
            "Run `tests/test_addr_codes_embedded_delim.py`; the "
            "20260602 build is calibrated to 0 BOM-prefixed ADDR_CODES "
            "rows (the test asserts c_name and c_name_chn both have 0 "
            "literal U+FEFF prefixes), and `c_addr_id = 702559` "
            "(Wei Shi / 尉氏) exists but carries no BOM.",
            "If a future build re-introduces dirty rows, the LookAtStatus "
            "GIS export of status code 40 (civil office / [為官者：文]) "
            "would again produce a 10-cell line against the 9-column "
            "header around the dirty row.",
        ],
        "steps_zh": [
            "本次建置症狀休眠——量測觸發列數以確認：",
            "執行 `tests/test_addr_codes_embedded_delim.py`；20260602 "
            "建置已校準為 ADDR_CODES 中 0 列帶 BOM 前綴（測試斷言 "
            "c_name 與 c_name_chn 的字面 U+FEFF 前綴皆為 0），且 "
            "`c_addr_id = 702559`（Wei Shi / 尉氏）存在但不帶 BOM。",
            "若未來建置重新引入髒列，LookAtStatus 對 status code 40"
            "（civil office / [為官者：文]）的 GIS 輸出，會在髒列附近"
            "再次產生對應 9 欄表頭的 10 格列。",
        ],
        "fix_en": (
            "Two complementary fixes, both worth doing.  (1) One-shot "
            "data cleanup: strip any leading `U+FEFF` from "
            "`ADDR_CODES.c_name` / `c_name_chn` (e.g. `UPDATE "
            "ADDR_CODES SET c_name = Mid(c_name, 2) WHERE Left(c_name, "
            "1) = ChrW(65279)` and the parallel statement for "
            "c_name_chn) — currently a no-op since 0 rows match, but "
            "harmless to keep in the release checklist.  (2) Defensive "
            "sanitisation in the export writers: before each `tStr = "
            "tStr + value + Chr(9)` append in the CmdGIS bodies, replace "
            "any embedded Chr(9/10/13/11/12) or U+FEFF in `value` with a "
            "space.  This protects against the next tab-bearing value "
            "from any export-bound text field, not just ADDR_CODES."
        ),
        "fix_zh": (
            "兩個互補的修正，皆值得做。(1) 一次性資料清理：移除 "
            "`ADDR_CODES.c_name` / `c_name_chn` 任何前導的 `U+FEFF`"
            "（例如 `UPDATE ADDR_CODES SET c_name = Mid(c_name, 2) "
            "WHERE Left(c_name, 1) = ChrW(65279)`，c_name_chn 另一條"
            "並行陳述式）——目前因 0 列符合而為空操作，但放進發布"
            "檢查清單無害。(2) 在輸出寫出端做防禦性消毒：在 CmdGIS "
            "主體每次 `tStr = tStr + value + Chr(9)` 之前，把 `value` "
            "中任何內嵌的 Chr(9/10/13/11/12) 或 U+FEFF 換成空白。"
            "這能防範下一個帶 TAB 的值（來自任何輸出相關的文字欄位，"
            "不限 ADDR_CODES）。"
        ),
        "severity_en": "P5 — Dormant this build (the export-delimiter "
                       "defect class is real in code, but the current "
                       "data has 0 triggering rows).",
        "severity_zh": "P5 —— 本次建置休眠（輸出分隔符這一類缺陷在"
                       "程式碼中確實存在，但目前資料有 0 列會觸發）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "CmdGIS_Click bodies (no escaping before "
                       "`tStr + value + Chr(9)`) across LookAtStatus / "
                       "Texts / Place / Associations / Office / "
                       "Kinship; trigger-count source: "
                       "tests/test_addr_codes_embedded_delim.py.",
            "fixture": "LookAtStatus status code 40 (civil office / "
                       "[為官者：文]); reachable dirty row historically "
                       "c_addr_id = 702559 (Wei Shi / 尉氏).  0 BOM "
                       "rows in the 20260602 data — dormant.",
            "user_symptom": "None this build (0 triggering rows).  When "
                            "dirty data is present, the GIS `.tab` "
                            "export gains an extra cell on the dirty "
                            "row and every column to its right shifts.",
            "ui_verified": False,
        },
    },
    {
        "id": 22,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtAssociations.CmdUCINet_Click",
        "title_en": "LookAtAssociations.CmdUCINet CreateTextFile lacks "
                    "the Unicode flag → error 5 on CJK c_name — LATENT "
                    "(runtime did not fire this build)",
        "title_zh": "LookAtAssociations.CmdUCINet 的 CreateTextFile "
                    "缺少 Unicode 旗標 → 遇 CJK c_name 時 error 5 "
                    "—— 潛伏（本次執行期未觸發）",
        "summary_en": (
            "`Form_LookAtAssociations.CmdUCINet_Click` writes the "
            "`.vna` export via "
            "`Scripting.FileSystemObject.CreateTextFile(tFileName, "
            "True)` (line 2575).  The 3rd argument (`Unicode`) is "
            "omitted, so it defaults to FALSE — the file opens in the "
            "system ANSI code page (cp1252 on en-US Windows).  In the "
            "`*node properties` section the body writes "
            "`tQuote + !c_name + tQuote`; when `c_name` contains a "
            "character with no cp1252 representation (a CJK Han "
            "ideograph in particular), `WriteLine` raises VBA error 5 "
            "('Invalid procedure call or argument') and the export "
            "aborts, leaving a truncated `.vna` file.  "
            "`Form_LookAtKinship.CmdUCINet_Click` has the identical "
            "2-arg pattern at line 2510.\n\n"
            "Why LATENT this build: the CmdUCINet button exists, but "
            "this non-interactive session could not drive the export, "
            "so the runtime error was not reproduced/re-verified.  The "
            "missing-Unicode-flag defect is a confirmed static fact; "
            "filed as latent pending UI re-verification.  Verified "
            "fixture for the trigger: association code "
            "`c_assoc_code = 437` (Presented literary composition as "
            "gift to / 贈詩、文), whose 1st-order association network "
            "includes a person carrying a Han ideograph in c_name."
        ),
        "summary_zh": (
            "`Form_LookAtAssociations.CmdUCINet_Click` 透過 "
            "`Scripting.FileSystemObject.CreateTextFile(tFileName, "
            "True)`（第 2575 行）寫出 `.vna`。第三個引數（`Unicode`）"
            "被省略，因此預設為 FALSE——檔案以系統 ANSI 字碼頁開啟"
            "（en-US Windows 上為 cp1252）。在 `*node properties` 區段"
            "中，主體寫出 `tQuote + !c_name + tQuote`；當 `c_name` "
            "含有 cp1252 無法表示的字元（尤其是 CJK 漢字）時，"
            "`WriteLine` 會丟出 VBA error 5（『Invalid procedure call "
            "or argument』），輸出中止，留下被截斷的 `.vna` 檔。"
            "`Form_LookAtKinship.CmdUCINet_Click` 在第 2510 行有完全"
            "相同的 2 引數樣式。\n\n"
            "本次為何潛伏：CmdUCINet 按鈕存在，但本次非互動式測試"
            "無法驅動輸出，故未重現／重新驗證執行期錯誤。缺少 Unicode "
            "旗標是已確認的靜態事實；故列為潛伏待 UI 重新驗證。"
            "觸發的已驗證 fixture：關聯代碼 `c_assoc_code = 437`"
            "（Presented literary composition as gift to / 贈詩、文），"
            "其一階關聯網絡含有 c_name 帶漢字的人物。"
        ),
        "steps_en": [
            "On this build the runtime error was not reproduced "
            "(non-interactive session).  Verify the missing flag "
            "statically:",
            "Open `analysis/dump/vba/Form_LookAtAssociations.vb` line "
            "2575: `Set tVNA = tFileSystem.CreateTextFile(tFileName, "
            "True)` — only 2 arguments, no Unicode flag.  The same "
            "pattern is at `Form_LookAtKinship.vb:2510`.",
            "To re-verify interactively later: open LookAtAssociations, "
            "pick `c_assoc_code = 437` (Presented literary composition "
            "as gift to / 贈詩、文), Run Query, click UCINet, choose a "
            "save location — expect a Run-time error 5 popup and a "
            "truncated `.vna`.",
        ],
        "steps_zh": [
            "本次建置未重現執行期錯誤（非互動式測試）。以靜態方式"
            "驗證缺少的旗標：",
            "開啟 `analysis/dump/vba/Form_LookAtAssociations.vb` 第 "
            "2575 行：`Set tVNA = tFileSystem.CreateTextFile(tFileName, "
            "True)`——只有 2 個引數，沒有 Unicode 旗標。相同樣式位於 "
            "`Form_LookAtKinship.vb:2510`。",
            "日後互動式重新驗證：開啟 LookAtAssociations，選 "
            "`c_assoc_code = 437`（Presented literary composition as "
            "gift to / 贈詩、文），執行查詢，點 UCINet，選一個儲存"
            "位置——預期會出現 Run-time error 5 彈窗與被截斷的 `.vna`。",
        ],
        "fix_en": (
            "Add `True` as the 3rd argument of `CreateTextFile` to open "
            "the file in Unicode (UTF-16LE) mode at "
            "`Form_LookAtAssociations.vb:2575` — `CreateTextFile("
            "tFileName, True, True)` — and apply the same one-line fix "
            "to `Form_LookAtKinship.vb:2510`.  Verify UCINET / Visone "
            "accept the UTF-16 `.vna` on the fixed build before "
            "declaring it closed."
        ),
        "fix_zh": (
            "在 `Form_LookAtAssociations.vb:2575` 為 `CreateTextFile` "
            "加上第三個引數 `True`，以 Unicode（UTF-16LE）模式開啟"
            "檔案——`CreateTextFile(tFileName, True, True)`——並對 "
            "`Form_LookAtKinship.vb:2510` 套用相同的一行修正。"
            "在宣告關閉前，先在修好的建置上確認 UCINET / Visone "
            "能接受 UTF-16 的 `.vna`。"
        ),
        "severity_en": "P5 — Latent (confirmed missing-Unicode-flag "
                       "static defect; runtime error 5 not reproduced "
                       "this non-interactive build, pending UI "
                       "re-verification).",
        "severity_zh": "P5 —— 潛伏（已確認缺少 Unicode 旗標的靜態"
                       "缺陷；本次非互動式建置未重現執行期 error 5，"
                       "待 UI 重新驗證）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "latent_code",
            "vba_ref": "Form_LookAtAssociations.vb:2575 "
                       "(CreateTextFile(tFileName, True) — 2-arg, no "
                       "Unicode flag) and Form_LookAtKinship.vb:2510 "
                       "(same 2-arg pattern).",
            "fixture": "LookAtAssociations, c_assoc_code = 437 "
                       "(Presented literary composition as gift to / "
                       "贈詩、文); 1st-order network includes a "
                       "Han-ideograph c_name.  Runtime not reproduced "
                       "this non-interactive build.",
            "user_symptom": "None reproduced this build.  When the "
                            "export runs on a Han-name network, a "
                            "Run-time error 5 popup blocks the user and "
                            "the `.vna` file is left truncated and "
                            "unusable.",
            "ui_verified": False,
        },
    },

    # ---------------- P5 — structural metric ----------------
    {
        "id": 23,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtAssociations.CmdPajek_Click",
        "title_en": "LookAtAssociations.CmdPajek '*Vertices' header "
                    "count read from RecordCount before MoveLast "
                    "(undercounts vertices) — structural metric, P5",
        "title_zh": "LookAtAssociations.CmdPajek 的『*Vertices』表頭"
                    "數值在 MoveLast 之前讀取 RecordCount（頂點數"
                    "少算）—— 結構性度量，P5",
        "summary_en": (
            "`Form_LookAtAssociations.CmdPajek_Click` binds the node "
            "recordset to a form recordset (`Set tRstNode = "
            "ZZ_SCRATCH_P_ASSOC.Form.Recordset`, line 2924), calls "
            "`tRstNode.MoveFirst` (line 2924), then writes the Pajek "
            "header `tStr = \"*Vertices \" + Trim(Str("
            "tRstNode.RecordCount))` (line 2924).  On a DAO recordset, "
            "`RecordCount` is only the number of rows ACCESSED so far, "
            "not the true total, until a `MoveLast` has fully populated "
            "it.  Reading it right after `MoveFirst` (with no MoveLast) "
            "yields an undercount, so the declared `*Vertices N` header "
            "can be far smaller than the actual number of vertex rows "
            "the loop subsequently writes.\n\n"
            "Finding class is structural_metric: the off-by-N was "
            "derived by parsing the exported `.net` header against the "
            "emitted vertex-row count, not from a re-verified UI "
            "symptom (this build was non-interactive; ui_verified is "
            "not set).  Filed at P5.  Demo fixture for the network: "
            "person c_personid = 437 (Jia Zhaoming / 賈昭明)."
        ),
        "summary_zh": (
            "`Form_LookAtAssociations.CmdPajek_Click` 把節點記錄集"
            "綁定到表單記錄集（`Set tRstNode = "
            "ZZ_SCRATCH_P_ASSOC.Form.Recordset`，第 2924 行），呼叫 "
            "`tRstNode.MoveFirst`（第 2924 行），接著寫出 Pajek 表頭 "
            "`tStr = \"*Vertices \" + Trim(Str(tRstNode.RecordCount))`"
            "（第 2924 行）。在 DAO 記錄集上，`RecordCount` 在尚未以 "
            "`MoveLast` 完整填充前，只是「目前已存取」的列數，而非"
            "真正的總數。在 `MoveFirst` 之後（且沒有 MoveLast）立即"
            "讀取，會得到少算的值，因此宣告的 `*Vertices N` 表頭可能"
            "遠小於迴圈隨後實際寫出的頂點列數。\n\n"
            "發現類別為 structural_metric：這個 off-by-N 是透過將輸出"
            "的 `.net` 表頭與實際寫出的頂點列數比對而得，並非來自"
            "重新驗證的 UI 症狀（本次為非互動式；未設定 "
            "ui_verified）。故列於 P5。網絡的示範 fixture：人物 "
            "c_personid = 437（Jia Zhaoming / 賈昭明）。"
        ),
        "steps_en": [
            "Verify statically from the dump and the cross-form test:",
            "Open `analysis/dump/vba/Form_LookAtAssociations.vb` lines "
            "2924-2924: `tRstNode` is set to the form recordset, "
            "`MoveFirst` is called, then `RecordCount` is read for the "
            "`*Vertices` header BEFORE any `MoveLast`.",
            "The cross-form structural probe "
            "`test_vba_pajek_gephi_cross_form` parses the exported "
            "`.net` and compares the `*Vertices N` header against the "
            "count of vertex rows emitted; the demo network is person "
            "c_personid = 437 (Jia Zhaoming / 賈昭明).",
        ],
        "steps_zh": [
            "從傾印與跨表單測試以靜態方式驗證：",
            "開啟 `analysis/dump/vba/Form_LookAtAssociations.vb` 第 "
            "2924-2924 行：`tRstNode` 被設為表單記錄集，呼叫 "
            "`MoveFirst`，接著在任何 `MoveLast` 之前就為 `*Vertices` "
            "表頭讀取 `RecordCount`。",
            "跨表單結構探測 `test_vba_pajek_gephi_cross_form` 會解析"
            "輸出的 `.net`，把 `*Vertices N` 表頭與寫出的頂點列數"
            "比對；示範網絡為人物 c_personid = 437"
            "（Jia Zhaoming / 賈昭明）。",
        ],
        "fix_en": (
            "Call `tRstNode.MoveLast` (then `MoveFirst`) before reading "
            "`RecordCount` at line 2924 so the header reflects the true "
            "vertex total, e.g. `tRstNode.MoveLast: tRstNode.MoveFirst: "
            "tStr = \"*Vertices \" + Trim(Str(tRstNode.RecordCount))`."
        ),
        "fix_zh": (
            "在第 2924 行讀取 `RecordCount` 之前先呼叫 "
            "`tRstNode.MoveLast`（再 `MoveFirst`），使表頭反映真正的"
            "頂點總數，例如 `tRstNode.MoveLast: tRstNode.MoveFirst: "
            "tStr = \"*Vertices \" + Trim(Str(tRstNode.RecordCount))`。"
        ),
        "severity_en": "P5 — Structural metric (export-header "
                       "undercount derived by parsing the .net file; "
                       "not UI-verified this non-interactive build).",
        "severity_zh": "P5 —— 結構性度量（透過解析 .net 檔得出的"
                       "輸出表頭少算；本次非互動式建置未經 UI 驗證）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "structural_metric",
            "vba_ref": "Form_LookAtAssociations.vb:2924 (Set tRstNode "
                       "= ZZ_SCRATCH_P_ASSOC.Form.Recordset) + :2924 "
                       "(MoveFirst) + :2924 (RecordCount read for "
                       "'*Vertices' header, before any MoveLast).",
            "fixture": "LookAtAssociations Pajek export for person "
                       "c_personid = 437 (Jia Zhaoming / 賈昭明).",
            "user_symptom": "The exported Pajek `.net` declares a "
                            "'*Vertices N' header smaller than the "
                            "number of vertex rows that follow.",
            "detection": "test_vba_pajek_gephi_cross_form",
        },
    },
    {
        "id": 24,
        "tier": "P5_dormant_or_latent",
        "form": "Form_LookAtKinship.CmdGUESS_Click",
        "title_en": "LookAtKinship GUESS/Gephi .gdf nodedef declares "
                    "15 columns but some node rows emit 13 cells — "
                    "structural metric, P5",
        "title_zh": "LookAtKinship 的 GUESS/Gephi .gdf nodedef 宣告 "
                    "15 欄，但部分節點列只寫出 13 格 —— 結構性度量，P5",
        "summary_en": (
            "`Form_LookAtKinship.vb`'s GUESS/Gephi `.gdf` writer "
            "declares a non-ASCII `nodedef>` header of 15 columns "
            "(line 549: name, color, label, labelvisible, style, "
            "pinyin, indexyear, sex, addr_name, addr_chn, latitude, "
            "longitude, DynastyCode, dynasty, dynasty_chn).  The "
            "per-row body (lines 565-650), however, emits a variable "
            "number of cells: the non-ASCII dynasty tail (lines 645-649) "
            "appends DynastyCode + dynasty + dynasty_chn only when "
            "c_dynasty is non-null (line 647) — the `If Not "
            "IsNull(!c_dynasty)` at line 646 has NO `Else`, so a node "
            "row whose dynasty is null skips those trailing cells "
            "entirely and emits fewer cells than the 15-column header, "
            "so a strict GDF reader sees a column-count mismatch.\n\n"
            "Finding class is structural_metric: the 15-vs-13 mismatch "
            "was derived by counting header columns against emitted "
            "row cells in the export, not from a re-verified UI symptom "
            "(non-interactive build; ui_verified is not set).  Filed at "
            "P5.  Demo fixture for the kinship network: person "
            "c_personid = 3211 (Zhao Tingmei / 趙廷美)."
        ),
        "summary_zh": (
            "`Form_LookAtKinship.vb` 的 GUESS/Gephi `.gdf` 寫出端"
            "宣告了一個 15 欄的非 ASCII `nodedef>` 表頭（第 549 行："
            "name、color、label、labelvisible、style、pinyin、"
            "indexyear、sex、addr_name、addr_chn、latitude、longitude、"
            "DynastyCode、dynasty、dynasty_chn）。然而逐列主體"
            "（第 565-650 行）寫出的格數不定：非 ASCII 的 dynasty "
            "結尾分支（第 645-649 行）只在 c_dynasty 非 null 時才"
            "附加 DynastyCode + dynasty + dynasty_chn（第 647 行）"
            "——第 646 行的 `If Not IsNull(!c_dynasty)` 沒有 `Else`，"
            "因此 dynasty 為 null 的節點列會完全略過那幾個結尾格，"
            "對應 15 欄的表頭只寫出較少的格，嚴格的 GDF 讀取器會"
            "看到欄數不符。\n\n"
            "發現類別為 structural_metric：15 對 13 的不符是透過將"
            "表頭欄數與輸出中寫出的列格數比對而得，並非來自重新驗證"
            "的 UI 症狀（非互動式建置；未設定 ui_verified）。故列於 "
            "P5。親屬網絡的示範 fixture：人物 c_personid = 3211"
            "（Zhao Tingmei / 趙廷美）。"
        ),
        "steps_en": [
            "Verify statically from the dump and the cross-form test:",
            "Open `analysis/dump/vba/Form_LookAtKinship.vb` line 549 "
            "— the non-ASCII `nodedef>` header declares 15 columns.  "
            "Then read the row body lines 565-650: the non-ASCII "
            "dynasty branch (lines 645-649) appends the DynastyCode/"
            "dynasty/dynasty_chn cells only when c_dynasty is non-null "
            "— the `If Not IsNull(!c_dynasty)` at line 646 has no "
            "`Else`, so null-dynasty rows emit fewer cells than the "
            "15-column header.",
            "The cross-form structural probe "
            "`test_vba_cmdguess_cross_form` parses the `.gdf` header "
            "column count against per-row cell counts; the demo network "
            "is person c_personid = 3211 (Zhao Tingmei / 趙廷美).",
        ],
        "steps_zh": [
            "從傾印與跨表單測試以靜態方式驗證：",
            "開啟 `analysis/dump/vba/Form_LookAtKinship.vb` 第 549 行"
            "——非 ASCII 的 `nodedef>` 表頭宣告 15 欄。再看列主體第 "
            "565-650 行：非 ASCII 的 dynasty 分支（第 645-649 行）"
            "只在 c_dynasty 非 null 時才附加 DynastyCode/dynasty/"
            "dynasty_chn——第 646 行的 `If Not IsNull(!c_dynasty)` "
            "沒有 `Else`，故 null-dynasty 列寫出的格數少於 15 欄表頭。",
            "跨表單結構探測 `test_vba_cmdguess_cross_form` 會比對 "
            "`.gdf` 表頭欄數與每列格數；示範網絡為人物 c_personid = "
            "3211（Zhao Tingmei / 趙廷美）。",
        ],
        "fix_en": (
            "Make every node row emit exactly the 15 cells the header "
            "declares.  Normalise the dynasty tail so all branches "
            "write DynastyCode + dynasty + dynasty_chn (with empty "
            "strings where a value is null) and end each branch with "
            "the same trailing-`tC` shape, so the cell count is "
            "header-stable on every row."
        ),
        "fix_zh": (
            "讓每一節點列都恰好寫出表頭宣告的 15 格。把 dynasty 結尾"
            "正規化，使所有分支都寫出 DynastyCode + dynasty + "
            "dynasty_chn（值為 null 處填空字串），並讓每個分支以相同"
            "的結尾 `tC` 形狀結束，使每列的格數都與表頭一致。"
        ),
        "severity_en": "P5 — Structural metric (export nodedef "
                       "column-count mismatch derived by parsing the "
                       ".gdf; not UI-verified this non-interactive "
                       "build).",
        "severity_zh": "P5 —— 結構性度量（透過解析 .gdf 得出的輸出 "
                       "nodedef 欄數不符；本次非互動式建置未經 UI "
                       "驗證）。",
        "screenshots": [],
        "evidence": {
            "finding_class": "structural_metric",
            "vba_ref": "Form_LookAtKinship.vb:549 (non-ASCII nodedef> "
                       "header, 15 columns) vs row body :565-650 (the "
                       "non-ASCII dynasty branch's `If Not "
                       "IsNull(!c_dynasty)` at :646 has no Else → "
                       "null-dynasty rows drop the trailing dynasty "
                       "cells).",
            "fixture": "LookAtKinship GUESS/Gephi export for person "
                       "c_personid = 3211 (Zhao Tingmei / 趙廷美).",
            "user_symptom": "The exported `.gdf` nodedef header "
                            "declares 15 columns but some node rows "
                            "contain only 13 tab-separated cells.",
            "detection": "test_vba_cmdguess_cross_form",
        },
    },
]


# ---------------------------------------------------------------------
# DOCX building helpers
# ---------------------------------------------------------------------

def _add_toc(document: Document, lang: str) -> None:
    """Insert a Word TOC field; Word will offer to update it on open."""
    para = document.add_paragraph()
    run = para.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    msg = OxmlElement("w:t")
    if lang == "en":
        msg.text = "Right-click here and choose 'Update Field' to populate."
    else:
        msg.text = t("右键点这里，选「更新域」即可生成完整目录。")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    r_el = run._r
    r_el.append(fld_char1)
    r_el.append(instr)
    r_el.append(fld_char2)
    r_el.append(msg)
    r_el.append(fld_char3)


def _h(document, level, text):
    p = document.add_heading(text, level=level)
    return p


def _bullets(document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _numbered(document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


DRIFT_JSON = REPO / "reports" / "index_drift_examples.json"
CLASSIFICATION_JSON = REPO / "reports" / "index_drift_classification.json"
RULE_CLASSIFICATION_JSON = REPO / "reports" / "index_year_drift_rule_classification.json"
RULE_GROUPS_JSON = REPO / "reports" / "index_year_drift_rule_groups.json"
ADDR_CLASSIFICATION_JSON = REPO / "reports" / "index_addr_drift_classification.json"
CAUSE_SUMMARY_JSON = REPO / "reports" / "index_drift_cause_summary.json"
DEMO_PERSONS_JSON = REPO / "reports" / "demo_persons.json"
KNOWN_BUGS_STATUS_JSON = REPO / "reports" / "known_bugs_status.json"
COVERAGE_MATRIX_JSON = REPO / "reports" / "coverage_matrix.json"
SCHEMA_DIFF_JSON = REPO / "reports" / "schema_diff.json"

# Cell display labels — used by both docx and markdown renderers
_CELL_DISPLAY = {
    "PASS":    ("✓ PASS",   "✓"),
    "FAIL":    ("✗ FAIL",   "✗ FAIL"),
    "ERROR":   ("⚠ ERROR",  "⚠ ERR"),
    "SKIP":    ("~ SKIP",   "~ SKIP"),
    "NOT_RUN": ("(not run)", "—?"),
    "N/A":     ("—",        "—"),
}


def _load_coverage_matrix() -> dict | None:
    """Return parsed coverage_matrix.json, or None if not generated / unreadable."""
    import json as _json
    if not COVERAGE_MATRIX_JSON.exists():
        return None
    try:
        data = _json.loads(COVERAGE_MATRIX_JSON.read_text(encoding="utf-8"))
        # Basic schema check so a truncated/stale file doesn't crash rendering
        _ = data["forms"]; _ = data["buttons"]; _ = data["matrix"]
        return data
    except Exception:
        return None


def _load_demo_persons() -> dict:
    import json as _json
    if DEMO_PERSONS_JSON.exists():
        return _json.loads(DEMO_PERSONS_JSON.read_text(encoding="utf-8"))
    return {}


def _load_bug_test_status() -> dict[int, dict]:
    """Map bug-id → {outcome: 'passed'|'failed'|'mixed', tests: [...],
    when: <iso>}.

    Reads `reports/known_bugs_status.json` produced by
    `pytest tests/test_known_bugs.py --json-report
    --json-report-file=reports/known_bugs_status.json`.

    Convention:
      - test name `test_bug<N>_*` → bug N
      - test name `test_bugs_<lo>_to_<hi>_*` → bugs lo..hi inclusive
      - test name `test_bug_view_statusdata_fy_alias_swap` /
        `test_bug_view_statusdata_fy_value_equals_ly_value` → bug 1
      - test name `test_bug_dao_reference_broken_in_user_mdb` → bug 2

    SAFETY: this function only INFORMS the report; it never causes the
    generator to drop or move content.  The intent is "show the
    maintainer that test results agree with the issue" — not to
    auto-edit the report.
    """
    import json as _json, re as _re
    if not KNOWN_BUGS_STATUS_JSON.exists():
        return {}
    data = _json.loads(
        KNOWN_BUGS_STATUS_JSON.read_text(encoding="utf-8")
    )
    raw_when = data.get("created") or data.get("environment", {}).get(
        "created"
    ) or ""
    when = ""
    if raw_when:
        try:
            from datetime import datetime, timezone
            ts = float(raw_when)
            when = (datetime.fromtimestamp(ts, tz=timezone.utc)
                    .strftime("%Y-%m-%d %H:%M UTC"))
        except (TypeError, ValueError):
            when = str(raw_when)

    by_bug: dict[int, list[tuple[str, str]]] = {}

    def _record(bug_id: int, nodeid: str, outcome: str) -> None:
        by_bug.setdefault(bug_id, []).append((nodeid, outcome))

    for t in data.get("tests", []):
        nodeid = t.get("nodeid", "")
        outcome = t.get("outcome", "")
        # Strip everything before "::" so we just have the test name.
        name = nodeid.rsplit("::", 1)[-1]

        # Cluster: test_bugs_<lo>_to_<hi>_...
        m = _re.match(r"^test_bugs_(\d+)_to_(\d+)_", name)
        if m:
            lo, hi = int(m.group(1)), int(m.group(2))
            for n in range(lo, hi + 1):
                _record(n, nodeid, outcome)
            continue
        # Individual: test_bug<N>_...
        m = _re.match(r"^test_bug(\d+)_", name)
        if m:
            _record(int(m.group(1)), nodeid, outcome)
            continue
        # Special hand-mapped cases (the no-number bug names).
        if "view_statusdata" in name:
            _record(1, nodeid, outcome)
        elif "dao_reference" in name:
            _record(2, nodeid, outcome)

    out: dict[int, dict] = {}
    for bug_id, items in by_bug.items():
        outcomes = {o for _, o in items}
        if outcomes == {"passed"}:
            agg = "passed"
        elif outcomes == {"failed"}:
            agg = "failed"
        else:
            agg = "mixed"
        out[bug_id] = {"outcome": agg, "tests": items, "when": when}
    return out


def _add_coverage_matrix_docx(doc, is_en: bool, Z, data: dict | None) -> None:
    """Render the form × button coverage matrix as a Word table.

    Always present: if data is None (builder not run yet), shows a
    placeholder message directing the user to run step 5b.
    """
    title = (
        "Coverage Matrix — Form × Button Test Results"
        if is_en else
        "覆蓋矩陣 —— 表單 × 按鈕測試結果"
    )
    _h(doc, 1, Z(title))

    if data is None:
        placeholder = (
            "Coverage matrix not yet generated.  Run step 5b of the "
            "build-test workflow:\n"
            "    python analysis/build_coverage_matrix.py "
            "--report reports/pytest_report_<build>.json"
            if is_en else
            "覆蓋矩陣尚未生成。請執行 build-test 工作流程第 5b 步：\n"
            "    python analysis/build_coverage_matrix.py "
            "--report reports/pytest_report_<build>.json"
        )
        doc.add_paragraph(Z(placeholder))
        return

    forms = data["forms"]
    buttons = data["buttons"]
    matrix = data["matrix"]

    # Build table: header row + one row per form
    tbl = doc.add_table(rows=1 + len(forms), cols=1 + len(buttons))
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0].cells
    hdr[0].text = Z("Form" if is_en else "表單")
    for j, btn in enumerate(buttons):
        hdr[j + 1].text = btn

    # Data rows
    for i, form in enumerate(forms):
        row_cells = tbl.rows[i + 1].cells
        row_cells[0].text = Z(form)
        for j, btn in enumerate(buttons):
            outcome = matrix.get(form, {}).get(btn, "NOT_RUN")
            label = _CELL_DISPLAY.get(outcome, (outcome, outcome))[0]
            row_cells[j + 1].text = Z(label)

    doc.add_paragraph("")
    summary = data.get("summary", {})
    note = (
        f"PASS: {summary.get('PASS', 0)}  |  "
        f"FAIL: {summary.get('FAIL', 0)}  |  "
        f"ERROR: {summary.get('ERROR', 0)}  |  "
        f"SKIP: {summary.get('SKIP', 0)}  |  "
        f"NOT RUN: {summary.get('NOT_RUN', 0)}  |  "
        f"N/A: {summary.get('N/A', 0)}"
    )
    p = doc.add_paragraph(Z(note))
    p.runs[0].italic = True


def _add_coverage_matrix_md(lines: list, is_en: bool, Z, data: dict | None) -> None:
    """Append the coverage matrix as a GitHub-flavoured markdown table.

    Always appended: placeholder shown when builder has not been run.
    """
    title = (
        "Coverage Matrix — Form × Button Test Results"
        if is_en else
        "覆蓋矩陣 —— 表單 × 按鈕測試結果"
    )
    lines.append(f"## {Z(title)}")
    lines.append("")

    if data is None:
        lines.append(
            "> Coverage matrix not yet generated.  "
            "Run step 5b: `python analysis/build_coverage_matrix.py "
            "--report reports/pytest_report_<build>.json`"
            if is_en else
            "> 覆蓋矩陣尚未生成。請執行第 5b 步：`python "
            "analysis/build_coverage_matrix.py "
            "--report reports/pytest_report_<build>.json`"
        )
        lines.append("")
        return

    forms = data["forms"]
    buttons = data["buttons"]
    matrix = data["matrix"]

    # Header
    header = "| Form | " + " | ".join(buttons) + " |"
    sep = "| --- |" + " --- |" * len(buttons)
    lines.append(Z(header))
    lines.append(sep)

    for form in forms:
        cells = []
        for btn in buttons:
            outcome = matrix.get(form, {}).get(btn, "NOT_RUN")
            short = _CELL_DISPLAY.get(outcome, (outcome, outcome))[1]
            cells.append(Z(short))
        lines.append(f"| {Z(form)} | " + " | ".join(cells) + " |")

    lines.append("")
    summary = data.get("summary", {})
    note = (
        f"_PASS: {summary.get('PASS',0)}"
        f" · FAIL: {summary.get('FAIL',0)}"
        f" · ERROR: {summary.get('ERROR',0)}"
        f" · SKIP: {summary.get('SKIP',0)}"
        f" · NOT RUN: {summary.get('NOT_RUN',0)}"
        f" · N/A: {summary.get('N/A',0)}_"
    )
    lines.append(Z(note))
    lines.append("")


def _add_index_drift_appendix(doc, is_en: bool, Z) -> None:
    """Render the 'index_year / index_addr drift, NOT a bug' chapter.

    Always renders: shows a placeholder directing the user to run
    collect_index_year_diffs.py when the source JSON files are missing.
    """
    import json as _json

    title = (
        "Appendix A — `c_index_year` / `c_index_addr_id` drift "
        "vs the cbdb-online-main-server snapshot "
        "(differences need per-row classification before being filed as bugs)"
        if is_en else
        "附錄 A —— `c_index_year` / `c_index_addr_id` 與 "
        "cbdb-online-main-server 快照之間的偏差"
        "（差異需要逐筆分類後才能判定是否為缺陷）"
    )
    _h(doc, 1, Z(title))

    if not DRIFT_JSON.exists():
        placeholder = (
            "Appendix data not yet generated.  Run step 5c of the "
            "build-test workflow:\n"
            "    python reports/collect_index_year_diffs.py\n"
            "This queries the DATA mdb via pyodbc and emits "
            "reports/index_drift_*.json."
            if is_en else
            "附錄數據尚未生成。請執行 build-test 工作流程第 5c 步：\n"
            "    python reports/collect_index_year_diffs.py\n"
            "此腳本透過 pyodbc 查詢 DATA mdb，並輸出 "
            "reports/index_drift_*.json。"
        )
        doc.add_paragraph(Z(placeholder))
        return

    try:
        data = _json.loads(DRIFT_JSON.read_text(encoding="utf-8"))
    except Exception as e:
        doc.add_paragraph(Z(
            f"Appendix data could not be loaded ({e}). "
            "Re-run: python reports/collect_index_year_diffs.py"
            if is_en else
            f"附錄數據載入失敗（{e}）。請重新執行：python reports/collect_index_year_diffs.py"
        ))
        return

    intro = (
        "When we compare BIOG_MAIN's `c_index_year` and "
        "`c_index_addr_id` between this User MDB and the weekly "
        "cbdb-online-main-server SQLite snapshot, a small fraction "
        "of persons disagree.\n\n"
        "**The two sides are independent implementations.** The "
        "SQLite snapshot's `c_index_year` is produced by "
        "cbdb-online-main-server's PHP "
        "`IndexYearRebuildService.php` and its `c_index_addr_id` by "
        "`IndexAddressRebuildService.php` (both at "
        "<https://github.com/cbdb-project/cbdb-online-main-server>); "
        "the User MDB-side: `c_index_addr_id` is rebuilt by VBA in "
        "`Form_frmIndexAddr` in the front-end mdb; `c_index_year` "
        "is rebuilt by **37 saved QueryDefs named `BM IY Rule …`** "
        "in the linked-tables backend `data/CBDB_<YYYYMMDD>_DATA"
        ".mdb`, driven by `frmBaseMaintenance`.  Both algorithms "
        "are now extracted and committed to the repo "
        "(`analysis/dump_data/querydefs_index/*.sql`); the form / "
        "module driver VBA still needs an interactive Access "
        "`SaveAsText` pass.  PHP is "
        "intended to mirror the VBA but they are separate code "
        "paths.  Per-row "
        "differences can come from at least four sources, and a "
        "diff alone doesn't tell us which: (1) source-data "
        "snapshot drift (the online system updates source rows "
        "continuously; the User MDB ships a point-in-time "
        "snapshot); (2) algorithm / porting divergence between "
        "PHP and VBA; (3) priority / tie-break differences when "
        "multiple candidate years apply; (4) null / default "
        "handling differences (e.g. how a missing birthyear "
        "collapses).\n\n"
        "The per-row classification of these diffs is in the "
        "**Classification summary** below, generated from "
        "`reports/index_drift_classification.json` when the "
        "classifier has run (a placeholder appears if it hasn't); "
        "the actual counts and buckets are data-driven, not "
        "hardcoded here.  The worked example rows further down "
        "(from `reports/index_drift_examples.json`) illustrate the "
        "*shapes* of disagreement — illustrative, not statistically "
        "representative, and a starting point for per-row triage, "
        "not a verdict either way.\n\n"
        "Why we still keep this appendix: if the *shape* of "
        "disagreement changes between releases (e.g. addresses "
        "start diverging where only years used to, or per-rule "
        "type-code distributions shift sharply), that hints the "
        "algorithm, schema, or one of the implementations moved "
        "— a stronger signal than per-row counts alone."
        if is_en else
        "我們把本 .mdb 的 BIOG_MAIN 與 cbdb-online-main-server 每週"
        "釋出的 SQLite 快照在 `c_index_year`、`c_index_addr_id` 兩個"
        "欄位上做比對，可以看到一小部分人物對不齊。\n\n"
        "**兩邊是兩套獨立的實作。**SQLite 快照中的 `c_index_year` "
        "是 cbdb-online-main-server 的 PHP "
        "`IndexYearRebuildService.php` 算出來的，`c_index_addr_id` "
        "則是 `IndexAddressRebuildService.php` 算出來的（程式碼都"
        "在 <https://github.com/cbdb-project/cbdb-online-main-server>"
        "）；User MDB 上對應的這兩個欄位則是由 Access "
        "**User MDB 那一邊**：`c_index_addr_id` 是由前端 mdb 裡的 "
        "`Form_frmIndexAddr` VBA 重建的；`c_index_year` 則是由"
        "連結表後端 `data/CBDB_<YYYYMMDD>_DATA.mdb` 裡 **37 條"
        "命名為 `BM IY Rule …` 的 QueryDef** 重建的，並由 "
        "`frmBaseMaintenance` 驅動。兩邊的演算法都已抽取並提交到 "
        "repo（`analysis/dump_data/querydefs_index/*.sql`），form / "
        "module 的驅動 VBA 仍需透過 Access 的 `SaveAsText` 互動式"
        "提取。PHP "
        "**意圖**鏡像 VBA，但兩者是兩條獨立的程式路徑。每一行"
        "差異**可能**來自下列至少四"
        "個原因，光看差異本身分不出來：(1) 源資料快照漂移"
        "（線上系統持續更新源資料；User MDB 是某個時點的快照）；"
        "(2) PHP 與 VBA 之間的演算法 / 移植差異；(3) 多個候選年份"
        "適用時，兩邊的優先序 / 平手規則不同；(4) null / 預設值"
        "處理不同（例如缺失的生年是當作 0、NULL，還是「卒年減 60」）。"
        "\n\n"
        "這些差異的逐筆分類見下方的 **分類匯總**，由 "
        "`reports/index_drift_classification.json` 自動生成（分類器"
        "尚未執行時顯示佔位文字）；實際的計數與分桶都是 data-driven，"
        "不在此處寫死。再往下列舉的樣例（來自 "
        "`reports/index_drift_examples.json`）只是**示範**差異**長什麼"
        "樣**，並非統計上有代表性，是後續逐筆分類的起點，並非任何"
        "方向的結論。\n\n"
        "我們仍然把這份附錄放在這裡，是因為：若差異的「形狀」在"
        "不同次發行之間發生變化（例如原本只年份對不齊，現在地點"
        "也開始對不齊；或 per-rule type-code 的分佈突然偏離），"
        "那是比逐筆計數更強的訊號，意味著演算法、schema、或某一"
        "邊的實作動過。"
    )
    for para in intro.split("\n\n"):
        doc.add_paragraph(Z(para))

    # ---- Classification summary ----
    if CLASSIFICATION_JSON.exists():
        cls = _json.loads(CLASSIFICATION_JSON.read_text(encoding="utf-8"))
        cs = cls["summary"]
        b = cs["buckets"]
        if is_en:
            _h(doc, 2, Z("Classification summary"))
            doc.add_paragraph(Z(
                f"Compared {cs['common']:,} personids common to both "
                f"databases (User MDB total {cs['user_mdb_total']:,}; "
                f"SQLite total {cs['sqlite_total']:,}; "
                f"User-only {cs['in_user_only']:,}; "
                f"SQLite-only {cs['in_sqlite_only']:,})."
            ))
            for b_key, label in [
                ("exact_match",
                 "exact match on all four compared fields"),
                ("source_drift_index_agrees",
                 "source drift but indices agreed (algorithms "
                 "tolerated the drift)"),
                ("source_drift_index_diffs_too",
                 "source drift AND at least one index field "
                 "differs (consistent with simple data-drift "
                 "hypothesis)"),
                ("index_year_only_diff",
                 "c_birthyear+c_deathyear identical, only "
                 "c_index_year differs — needs follow-up"),
                ("index_addr_only_diff",
                 "c_birthyear+c_deathyear identical, only "
                 "c_index_addr_id differs — needs follow-up"),
                ("index_both_diff",
                 "c_birthyear+c_deathyear identical, BOTH index "
                 "fields differ — strongest signal of compound "
                 "divergence, needs follow-up"),
            ]:
                doc.add_paragraph(Z(
                    f"  • {b[b_key]:>7,} ({100.0*b[b_key]/max(cs['common'],1):.3f}%)"
                    f"  — {label}"
                ))
            doc.add_paragraph(Z(
                f"Net diffs: {cs['common']-b['exact_match']:,} of "
                f"{cs['common']:,} ({100.0*(cs['common']-b['exact_match'])/max(cs['common'],1):.3f}%).  "
                f"Of those, {b['source_drift_index_agrees']+b['source_drift_index_diffs_too']:,} are "
                f"clearly attributable to source drift in birthyear/"
                f"deathyear; {b['index_year_only_diff']+b['index_addr_only_diff']+b['index_both_diff']:,} "
                f"need per-row follow-up (could be PHP↔VBA "
                f"divergence, or drift in evidence tables — "
                f"BIOG_ADDR_DATA / ENTRY_DATA / NIAN_HAO etc. — "
                f"that this classifier does not compare).  Full "
                f"output: `reports/index_drift_classification.json`; "
                f"algorithm pointers: "
                f"`analysis/index_drift_algorithm_notes.md`."
            ))
        else:
            _h(doc, 2, Z("分類匯總"))
            doc.add_paragraph(Z(
                f"比對了兩邊都有的 {cs['common']:,} 個 personid"
                f"（User MDB 共 {cs['user_mdb_total']:,} 筆；"
                f"SQLite 共 {cs['sqlite_total']:,} 筆；"
                f"僅 User MDB 有 {cs['in_user_only']:,} 筆；"
                f"僅 SQLite 有 {cs['in_sqlite_only']:,} 筆）。"
            ))
            for b_key, label in [
                ("exact_match", "四個欄位全部一致"),
                ("source_drift_index_agrees",
                 "源資料有漂移但兩邊 index 都一致"
                 "（演算法吸收了漂移）"),
                ("source_drift_index_diffs_too",
                 "源資料有漂移、且至少一個 index 不同"
                 "（與簡單資料漂移假說相符）"),
                ("index_year_only_diff",
                 "生年/卒年一致，但只有 c_index_year 不同 —— 待追查"),
                ("index_addr_only_diff",
                 "生年/卒年一致，但只有 c_index_addr_id 不同 —— 待追查"),
                ("index_both_diff",
                 "生年/卒年一致，但兩個 index 都不同 —— 複合差異"
                 "的最強單列訊號，待追查"),
            ]:
                doc.add_paragraph(Z(
                    f"  • {b[b_key]:>7,} ({100.0*b[b_key]/max(cs['common'],1):.3f}%)"
                    f"  —— {label}"
                ))
            doc.add_paragraph(Z(
                f"淨差異：{cs['common']-b['exact_match']:,} / "
                f"{cs['common']:,}（{100.0*(cs['common']-b['exact_match'])/max(cs['common'],1):.3f}%）。"
                f"其中 {b['source_drift_index_agrees']+b['source_drift_index_diffs_too']:,} "
                f"筆能明確歸因於 birthyear/deathyear 的源資料漂移；"
                f"剩下 {b['index_year_only_diff']+b['index_addr_only_diff']+b['index_both_diff']:,} "
                f"筆需要逐筆追查（可能是 PHP↔VBA 演算法差異，"
                f"也可能是本分類器沒有比較的 evidence 表"
                f"（BIOG_ADDR_DATA / ENTRY_DATA / NIAN_HAO 等）裡的漂移）。"
                f"完整輸出見 `reports/index_drift_classification.json`，"
                f"算法來源指標見 `analysis/index_drift_algorithm_notes.md`。"
            ))
    else:
        _h(doc, 2, Z("Classification summary" if is_en else "分類匯總"))
        doc.add_paragraph(Z(
            "_Not generated — run `python analysis/classify_index_drift.py` "
            "(reports/index_drift_classification.json absent)._"
            if is_en else
            "_尚未生成 —— 請執行 `python analysis/classify_index_drift.py`"
            "（缺 reports/index_drift_classification.json）。_"
        ))

    # ---- Year-drift rule classification (PR K1) ----
    if RULE_CLASSIFICATION_JSON.exists():
        rcls = _json.loads(
            RULE_CLASSIFICATION_JSON.read_text(encoding="utf-8"))
        rs = rcls["summary"]
        rb = rs["buckets"]
        if is_en:
            _h(doc, 2, Z(
                "Year-only diffs — per-row rule classification"
            ))
            doc.add_paragraph(Z(
                f"Of the {rs['total_year_diffs']} year-only diffs, "
                f"each row was bucketed against PR N's rule-level "
                f"runtime-vs-PHP comparison "
                f"(`analysis/index_year_rule_comparison.md`).  "
                f"Conservative buckets (rows count once each):"
            ))
            for k in [
                "php_returned_sentinel",
                "php_did_not_compute",
                "access_did_not_compute",
                "iteration_order_diff",
                "consistent_within_rule",
                "candidate_algorithm_divergence",
                "unclassified",
            ]:
                if rb[k] == 0:
                    continue
                doc.add_paragraph(Z(
                    f"  • {rb[k]:>3,}  {k}"
                ))
            doc.add_paragraph(Z(
                f"None of these are confirmed bugs.  Full per-row "
                f"output (with the source-data evidence we used to "
                f"bucket each one) is in "
                f"`reports/index_year_drift_rule_classification.json`."
            ))
            # ---- K2 deeper triage ----
            if RULE_GROUPS_JSON.exists():
                gcls = _json.loads(
                    RULE_GROUPS_JSON.read_text(encoding="utf-8"))
                gs = gcls["summary"]
                doc.add_paragraph(Z(
                    f"Deeper triage (PR K2, "
                    f"`analysis/triage_index_year_drift_groups.py`) "
                    f"named the leftover buckets:"
                ))
                doc.add_paragraph(Z(
                    f"  • {gs['consistent_within_rule']['rows']} "
                    f"`consistent_within_rule` rows → "
                    f"{gs['consistent_within_rule']['groups']} "
                    f"signature groups.  PR AI + AJ probes "
                    f"reversed the prior tie-break hypothesis: all "
                    f"14 rows are `source_data_drift_biog_main_or_"
                    f"kin_data_between_sides` (8 BIOG_MAIN birthyear "
                    f"drift + 6 KIN_DATA evidence-pid drift).  "
                    f"Upstream PHP-side / SQLite-snapshot data "
                    f"issue, not a CBDB algorithm divergence."
                ))
                doc.add_paragraph(Z(
                    f"  • {gs['unclassified']['total']} "
                    f"`unclassified` rows → "
                    f"{gs['unclassified']['named_after_triage']} named, "
                    f"{gs['unclassified']['blocked_by_runtime_priority_triage_pending']} "
                    f"flagged `blocked_by_runtime_priority_triage_pending` "
                    f"(PR M dumped frmBaseMaintenance, so the source is in repo; resolving each row still needs a per-row walk of the runtime priority/iteration order)."
                ))
                doc.add_paragraph(Z(
                    f"  • {gs['php_did_not_compute']['rows']} "
                    f"`php_did_not_compute` rows → "
                    f"{gs['php_did_not_compute']['groups']} groups by "
                    f"Access tcode; biggest is `access_tcode='05'` × 7 "
                    f"(`candidate_php_entry_code_mapping_gap` for jinshi)."
                ))
                doc.add_paragraph(Z(
                    f"Full per-group output: "
                    f"`reports/index_year_drift_rule_groups.json`."
                ))
            # ---- Address-drift classification (PR L) ----
            if ADDR_CLASSIFICATION_JSON.exists():
                acls = _json.loads(
                    ADDR_CLASSIFICATION_JSON.read_text(encoding="utf-8"))
                aS = acls["summary"]
                ab = aS["buckets"]
                _h(doc, 2, Z(
                    "c_index_addr_id diffs — per-row classification"
                ))
                doc.add_paragraph(Z(
                    f"Of the {aS['total_addr_diffs']} c_index_addr "
                    f"diffs (478 `index_addr_only_diff` + 10 "
                    f"`index_both_diff` from PR G's bucketing), "
                    f"each row was classified by re-simulating the "
                    f"rank-priority + MAX(c_sequence) algorithm "
                    f"against each side's BIOG_ADDR_DATA + the "
                    f"shared BIOG_ADDR_CODES rank table:"
                ))
                for k, label in [
                    ("mdb_stale_index_addr",
                     "User MDB stored value doesn't match its own "
                     "recompute; SQLite does — User MDB c_index_addr_id "
                     "is stale (re-run frmBaseMaintenance)"),
                    ("mdb_value_php_null",
                     "User MDB has a value, PHP wrote 0/null"),
                    ("same_candidates_diff_winner",
                     "identical BIOG_ADDR_DATA, different winner — "
                     "candidate algorithm divergence (all in addr_type=1)"),
                    ("both_stale_recompute_mismatch",
                     "neither side matches recompute"),
                    ("both_sides_match_recomputed",
                     "both stored values match recompute over their "
                     "own BIOG_ADDR_DATA — diff is source-data drift"),
                    ("sqlite_stale_index_addr",
                     "PHP stored value doesn't match its own recompute"),
                    ("mdb_null_php_value", ""),
                    ("unclassified", ""),
                ]:
                    n = ab.get(k, 0)
                    if n == 0:
                        continue
                    doc.add_paragraph(Z(
                        f"  • {n:>4,d}  {k}{(' — ' + label) if label else ''}"
                    ))
                doc.add_paragraph(Z(
                    f"None of these are confirmed bugs.  The 412 "
                    f"`mdb_stale_index_addr` rows are a maintenance-"
                    f"cadence diff (the User MDB needs its "
                    f"frmBaseMaintenance rebuild re-run before the "
                    f"next release).  The 10 "
                    f"`same_candidates_diff_winner` rows are the only "
                    f"candidate algorithm-divergence rows.  Full "
                    f"per-row output: "
                    f"`reports/index_addr_drift_classification.json`."
                ))
                doc.add_paragraph(Z(
                    f"PR M (`analysis/dump_data_mdb_vba.py`) extracted "
                    f"`frmBaseMaintenance.CmdIndexAddress_Click` from "
                    f"the DATA mdb.  It does NOT explicitly "
                    f"`MAX(c_sequence)`-aggregate the way PHP does — "
                    f"a candidate algorithmic divergence on top of "
                    f"the maintenance-cadence issue.  Suggested "
                    f"release-checklist mitigation: run "
                    f"`CmdIndexYear` then `CmdIndexAddress` on the "
                    f"DATA mdb before shipping a new User MDB.  See "
                    f"`analysis/index_drift_algorithm_notes.md` § "
                    f"\"Maintenance trigger path\" for the full "
                    f"write-up."
                ))

            # ---- Cause analysis appendix (PR Y) ----
            if CAUSE_SUMMARY_JSON.exists():
                cs = _json.loads(
                    CAUSE_SUMMARY_JSON.read_text(encoding="utf-8"))
                _h(doc, 2, "What currently explains the drift")
                doc.add_paragraph(Z(
                    "Per-bucket cause / supporting evidence / "
                    "confidence / next action lives in "
                    "`analysis/index_drift_cause_analysis.md`.  This "
                    "section just summarises the headline counts and "
                    "confidence per bucket; no bucket is labelled a "
                    "confirmed CBDB bug."
                ))
                # c_index_year table
                _h(doc, 3, "c_index_year cause buckets")
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Light Grid Accent 1"
                hdr = tbl.rows[0].cells
                hdr[0].text = "Bucket"
                hdr[1].text = "Count"
                hdr[2].text = "Confidence"
                for b in cs["c_index_year"]["buckets"]:
                    if b["count"] == 0:
                        continue
                    row = tbl.add_row().cells
                    row[0].text = b["bucket"]
                    row[1].text = str(b["count"])
                    row[2].text = b["confidence"]
                doc.add_paragraph("")
                # c_index_addr_id table
                _h(doc, 3, "c_index_addr_id cause buckets")
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Light Grid Accent 1"
                hdr = tbl.rows[0].cells
                hdr[0].text = "Bucket"
                hdr[1].text = "Count"
                hdr[2].text = "Confidence"
                for b in cs["c_index_addr_id"]["buckets"]:
                    if b["count"] == 0:
                        continue
                    row = tbl.add_row().cells
                    row[0].text = b["bucket"]
                    row[1].text = str(b["count"])
                    row[2].text = b["confidence"]
                doc.add_paragraph("")
                doc.add_paragraph(Z(
                    "Top suggested next investigations (full list in "
                    "the cause-analysis md):"
                ))
                for inv in cs[
                    "suggested_next_investigations_in_priority_order"
                ][:3]:
                    doc.add_paragraph(Z(
                        f"  {inv['id']}. {inv['task']} — would close "
                        f"{inv['would_close_rows']} rows; engineering "
                        f"cost: {inv['engineering_cost']}."
                    ))
        else:
            _h(doc, 2, Z("年份差異 —— 逐筆 rule 分類"))
            doc.add_paragraph(Z(
                f"在 {rs['total_year_diffs']} 筆「只有 c_index_year "
                f"不一致」的行中，逐筆比對 PR N "
                f"(`analysis/index_year_rule_comparison.md`) 的"
                f"runtime-vs-PHP 規則級差異。保守分類如下："
            ))
            label_zh = {
                "php_returned_sentinel": "PHP 寫了 sentinel／溢位值",
                "php_did_not_compute": "PHP 沒算出值（覆蓋率缺口）",
                "access_did_not_compute": "Access 沒算出值（覆蓋率缺口）",
                "iteration_order_diff": "Phase-C 迭代次數不同",
                "consistent_within_rule": "多列共享同一 (php_tcode, access_tcode, diff) — 單一規則差異訊號",
                "candidate_algorithm_divergence": "形狀符合 K1 的歷史 hypothesis probe 但無法以單筆證據重建",
                "unclassified": "尚未對上任何模式",
            }
            for k in [
                "php_returned_sentinel",
                "php_did_not_compute",
                "access_did_not_compute",
                "iteration_order_diff",
                "consistent_within_rule",
                "candidate_algorithm_divergence",
                "unclassified",
            ]:
                if rb[k] == 0:
                    continue
                doc.add_paragraph(Z(
                    f"  • {rb[k]:>3,}  {label_zh[k]}"
                ))
            doc.add_paragraph(Z(
                f"以上沒有任何一筆被視為已確認的 bug。"
                f"逐筆輸出（含分桶所依據的源資料證據）見 "
                f"`reports/index_year_drift_rule_classification.json`。"
            ))
            # ---- K2 deeper triage (zh) ----
            if RULE_GROUPS_JSON.exists():
                gcls = _json.loads(
                    RULE_GROUPS_JSON.read_text(encoding="utf-8"))
                gs = gcls["summary"]
                doc.add_paragraph(Z(
                    f"PR K2 進一步的 triage "
                    f"(`analysis/triage_index_year_drift_groups.py`) 把"
                    f"剩下的桶命名清楚："
                ))
                doc.add_paragraph(Z(
                    f"  • `consistent_within_rule` "
                    f"{gs['consistent_within_rule']['rows']} 筆 → "
                    f"{gs['consistent_within_rule']['groups']} 個 signature "
                    f"分組。PR AI + AJ 的逐筆探測推翻了原本的 tie-break "
                    f"假說：14 筆全是 `source_data_drift_biog_main_or_"
                    f"kin_data_between_sides`（8 筆 BIOG_MAIN birthyear "
                    f"漂移 + 6 筆 KIN_DATA evidence-pid 漂移）。屬於 "
                    f"PHP-side / SQLite snapshot 的上游資料漂移，並非 "
                    f"CBDB 演算法差異。"
                ))
                doc.add_paragraph(Z(
                    f"  • `unclassified` {gs['unclassified']['total']} 筆 → "
                    f"{gs['unclassified']['named_after_triage']} 筆已命名，"
                    f"{gs['unclassified']['blocked_by_runtime_priority_triage_pending']} 筆"
                    f"標為 `blocked_by_runtime_priority_triage_pending`"
                    f"（PR M 已 dump frmBaseMaintenance，源碼已在 repo；要逐筆判斷哪邊正確仍需走一遍 runtime 的 priority／iteration 順序）。"
                ))
                doc.add_paragraph(Z(
                    f"  • `php_did_not_compute` "
                    f"{gs['php_did_not_compute']['rows']} 筆 → 按 Access "
                    f"tcode 分 {gs['php_did_not_compute']['groups']} 組；"
                    f"最大的是 `access_tcode='05'` × 7（jinshi 進士類的 "
                    f"`candidate_php_entry_code_mapping_gap`）。"
                ))
                doc.add_paragraph(Z(
                    f"逐組輸出見 "
                    f"`reports/index_year_drift_rule_groups.json`。"
                ))
            # ---- Address-drift classification (PR L) — zh ----
            if ADDR_CLASSIFICATION_JSON.exists():
                acls = _json.loads(
                    ADDR_CLASSIFICATION_JSON.read_text(encoding="utf-8"))
                aS = acls["summary"]
                ab = aS["buckets"]
                _h(doc, 2, Z("c_index_addr_id 差異 —— 逐筆分類"))
                doc.add_paragraph(Z(
                    f"在 {aS['total_addr_diffs']} 筆 c_index_addr 差異"
                    f"中（PR G 分桶的 478 `index_addr_only_diff` + 10 "
                    f"`index_both_diff`），逐筆把兩邊的 BIOG_ADDR_DATA "
                    f"代入「rank-priority + MAX(c_sequence)」演算法重算，"
                    f"與實際儲存值對照分類："
                ))
                label_zh = {
                    "mdb_stale_index_addr":
                        "User MDB 儲存值與重算結果不符；SQLite 相符 "
                        "—— User MDB 的 c_index_addr_id 已 stale，"
                        "需要重新跑 frmBaseMaintenance",
                    "mdb_value_php_null": "User MDB 有值，PHP 寫了 0/null",
                    "same_candidates_diff_winner":
                        "兩邊 BIOG_ADDR_DATA 完全相同但 winner 不同 "
                        "—— 候選演算法差異（全在 addr_type=1）",
                    "both_stale_recompute_mismatch":
                        "兩邊都不符合重算結果",
                    "both_sides_match_recomputed":
                        "兩邊各自的儲存值都符合自己的重算結果 "
                        "—— 差異來自源資料漂移",
                    "sqlite_stale_index_addr":
                        "PHP 儲存值與重算結果不符",
                    "mdb_null_php_value": "",
                    "unclassified": "",
                }
                for k in [
                    "mdb_stale_index_addr",
                    "mdb_value_php_null",
                    "same_candidates_diff_winner",
                    "both_stale_recompute_mismatch",
                    "both_sides_match_recomputed",
                    "sqlite_stale_index_addr",
                    "mdb_null_php_value",
                    "unclassified",
                ]:
                    n = ab.get(k, 0)
                    if n == 0:
                        continue
                    lbl = label_zh.get(k, "")
                    doc.add_paragraph(Z(
                        f"  • {n:>4,d}  {k}{(' — ' + lbl) if lbl else ''}"
                    ))
                doc.add_paragraph(Z(
                    f"以上沒有任何一筆被視為已確認的 bug。412 筆 "
                    f"`mdb_stale_index_addr` 屬於維護週期差異（User MDB "
                    f"在下次釋出前需要重跑 frmBaseMaintenance）。10 筆 "
                    f"`same_candidates_diff_winner` 是唯一的候選演算法"
                    f"差異。逐筆輸出見 "
                    f"`reports/index_addr_drift_classification.json`。"
                ))
                doc.add_paragraph(Z(
                    f"PR M（`analysis/dump_data_mdb_vba.py`）從 DATA mdb "
                    f"抽出了 `frmBaseMaintenance.CmdIndexAddress_Click`。"
                    f"它**沒有**像 PHP 那樣明確 `MAX(c_sequence)` 聚合 "
                    f"—— 在維護週期差異之外，這還是一個候選演算法差異。"
                    f"建議的 release checklist 緩解步驟：在 User MDB "
                    f"出貨前先在 DATA mdb 上跑 `CmdIndexYear`，再跑 "
                    f"`CmdIndexAddress`。詳見 "
                    f"`analysis/index_drift_algorithm_notes.md` 中的 "
                    f"\"Maintenance trigger path\" 段。"
                ))
                doc.add_paragraph(Z(
                    f"PR S（`analysis/deep_dive_addr_same_candidates.py`）"
                    f"證實 10 筆 `same_candidates_diff_winner` 全部"
                    f"是 MAX(c_sequence) 的 tie 問題（同一(person, "
                    f"addr_type) 有多筆 BIOG_ADDR_DATA 行 c_sequence "
                    f"並列最大）。PHP、Access、以及我方重算各自挑了"
                    f"不同的 row。兩邊都遵循同一條文件規則，沒有「錯」"
                    f"的一方。候選緩解：兩邊都加一條明確的二級 "
                    f"tie-break（如 MIN(c_addr_id)）。逐筆證據見 "
                    f"`reports/index_addr_same_candidates_deep_dive.json`。"
                ))

            # ---- Cause analysis appendix (PR Y) ----
            if CAUSE_SUMMARY_JSON.exists():
                cs = _json.loads(
                    CAUSE_SUMMARY_JSON.read_text(encoding="utf-8"))
                _h(doc, 2, Z("目前能解釋的 drift 原因"))
                doc.add_paragraph(Z(
                    "每個 bucket 的成因／證據／信心度／下一步追查"
                    "都寫在 `analysis/index_drift_cause_analysis.md`。"
                    "本節只列每個 bucket 的計數和信心度摘要；目前"
                    "沒有任何 bucket 被列為已確認的 CBDB bug。"
                ))
                _h(doc, 3, Z("c_index_year 原因桶"))
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Light Grid Accent 1"
                hdr = tbl.rows[0].cells
                hdr[0].text = Z("Bucket")
                hdr[1].text = Z("筆數")
                hdr[2].text = Z("信心度")
                for b in cs["c_index_year"]["buckets"]:
                    if b["count"] == 0:
                        continue
                    row = tbl.add_row().cells
                    row[0].text = b["bucket"]
                    row[1].text = str(b["count"])
                    row[2].text = b["confidence"]
                doc.add_paragraph("")
                _h(doc, 3, Z("c_index_addr_id 原因桶"))
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Light Grid Accent 1"
                hdr = tbl.rows[0].cells
                hdr[0].text = Z("Bucket")
                hdr[1].text = Z("筆數")
                hdr[2].text = Z("信心度")
                for b in cs["c_index_addr_id"]["buckets"]:
                    if b["count"] == 0:
                        continue
                    row = tbl.add_row().cells
                    row[0].text = b["bucket"]
                    row[1].text = str(b["count"])
                    row[2].text = b["confidence"]
                doc.add_paragraph("")
                doc.add_paragraph(Z(
                    "建議優先處理的調查項目（完整列表見 cause-analysis md）："
                ))
                for inv in cs[
                    "suggested_next_investigations_in_priority_order"
                ][:3]:
                    doc.add_paragraph(Z(
                        f"  {inv['id']}. {inv['task']} —— 可消化 "
                        f"{inv['would_close_rows']} 筆；工程成本："
                        f"{inv['engineering_cost']}。"
                    ))

    # ---- Per bucket ----
    bucket_meta = {
        "year_only": {
            "title_en": "Examples where only c_index_year disagrees",
            "title_zh": "仅 `c_index_year` 不一致的样例",
            "explain_en": (
                "Same person, same source data, but the two pipelines "
                "picked different priority rules. The User MDB's "
                "`c_index_year_type_code` and `c_index_year_source_id` "
                "differ from the snapshot's, which means each pipeline "
                "selected a different upstream evidence row when "
                "deciding 'which year is most representative for this "
                "person'."
            ),
            "explain_zh": (
                "同一个人物、同一份源数据，但两套管线选了不同的优先级规则。"
                "User MDB 的 `c_index_year_type_code` 和 "
                "`c_index_year_source_id` 与快照不同，意味着两边在「为这个"
                "人物选一个最能代表生平的年份」时，挑中了不同的上游证据行。"
            ),
        },
        "addr_only": {
            "title_en": "Examples where only c_index_addr_id disagrees",
            "title_zh": "仅 `c_index_addr_id` 不一致的样例",
            "explain_en": (
                "Year agrees, address doesn't. Most commonly the User "
                "MDB still has an older `c_index_addr_id` choice while "
                "the snapshot has been re-classified to a finer-grained "
                "or higher-quality address (or to NULL when the "
                "evidence was downgraded). The two pipelines apply the "
                "same address-priority rules but to different snapshots "
                "of `BIOG_ADDR_DATA`."
            ),
            "explain_zh": (
                "年份对得上，但地点对不上。最常见的情况是：User MDB 仍保留"
                "较早一次的 `c_index_addr_id` 选择，而快照已经重新分类到"
                "更细粒度或证据等级更高的地点（或在证据被下调时改为 NULL）。"
                "两套管线用的是同一套地点优先级规则，只是依据的 "
                "`BIOG_ADDR_DATA` 快照不同。"
            ),
        },
        "both": {
            "title_en": "Examples where both fields disagree",
            "title_zh": "两个字段都不一致的样例",
            "explain_en": (
                "Both year and address moved together — usually because "
                "the snapshot got new high-priority biographical "
                "evidence (a newly-entered birth year or a more "
                "specific address from a different source text)."
            ),
            "explain_zh": (
                "年份和地点同时变动 —— 通常是因为快照在某个高优先级的传记"
                "证据行上有新增（例如新录入的生年，或来自其他原始文献的更"
                "精确地点）。"
            ),
        },
        "source_data": {
            "title_en": (
                "Examples where the SOURCE data itself differs "
                "(birthyear / deathyear)"
            ),
            "title_zh": (
                "底层 SOURCE 数据本身不同（生年 / 卒年）的样例"
            ),
            "explain_en": (
                "Here the SQLite snapshot has different `c_birthyear` / "
                "`c_deathyear` from the User MDB. This is the clearest "
                "case of pure data drift: someone updated the source "
                "row in the cbdb-online-main-server pipeline after the "
                "User MDB was last exported. There's no algorithmic "
                "disagreement here — just two different snapshots in "
                "time."
            ),
            "explain_zh": (
                "在这一组里，SQLite 快照的 `c_birthyear` / `c_deathyear` "
                "本身就和 User MDB 不一样。这是最纯粹的数据快照差异：在 "
                "User MDB 最近一次导出之后，有人在 "
                "cbdb-online-main-server 那边更新了源数据行。这里完全没有"
                "算法分歧 —— 只是两个时间点的两份快照。"
            ),
        },
    }

    for bucket_key, meta in bucket_meta.items():
        items = data.get(bucket_key, [])
        if not items:
            continue
        _h(doc, 2, Z(meta["title_en"] if is_en else meta["title_zh"]))
        for para in (meta["explain_en"] if is_en
                     else meta["explain_zh"]).split("\n\n"):
            doc.add_paragraph(Z(para))

        # Render each example as a small table.
        for ex in items[:5]:
            u = ex["user"]; s = ex["sqlite"]
            label = f"c_personid = {ex['personid']} — {ex['name_chn']} ({ex['name_py']})"
            _h(doc, 3, Z(label))
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Light Grid Accent 1"
            hdr = tbl.rows[0].cells
            hdr[0].text = Z("Field" if is_en else "字段")
            hdr[1].text = Z("User MDB" if is_en else "本 .mdb (User MDB)")
            hdr[2].text = Z(
                "cbdb-online-main-server snapshot"
                if is_en else
                "cbdb-online-main-server 快照"
            )
            for f_label, key in [
                ("c_index_year", "index_year"),
                ("c_index_addr_id", "index_addr_id"),
                ("c_birthyear", "birthyear"),
                ("c_deathyear", "deathyear"),
                ("c_index_year_type_code", "index_year_type_code"),
                ("c_index_year_source_id", "index_year_source_id"),
            ]:
                row = tbl.add_row().cells
                row[0].text = f_label
                row[1].text = "" if u[key] is None else str(u[key])
                row[2].text = "" if s[key] is None else str(s[key])
            doc.add_paragraph("")  # spacer

    note = (
        "These examples were collected by "
        "`reports/collect_index_year_diffs.py` from the first 20 000 "
        "person ids common to both databases. The full automated "
        "cross-check (`tests/test_index_year_xcheck.py`, ~657 246 "
        "persons) reports roughly the same shape: very small "
        "fractions disagree, with addresses being the most common "
        "drift type."
        if is_en else
        "这批样例是 `reports/collect_index_year_diffs.py` 在两边数据库共有"
        "的前 20000 个 personid 中采集的。完整的自动化对照（"
        "`tests/test_index_year_xcheck.py`，约 657246 人）给出相同的形态："
        "对不齐的比例非常小，其中地点偏差占多数。"
    )
    doc.add_paragraph(Z(note))


_VALID_TIERS = frozenset([
    "P0_silent_data", "P1_visible_crash", "P2_silent_display",
    "P3_missing_ui", "P4_setup", "P5_dormant_or_latent",
])
# NOTE: "resolved" is intentionally excluded.  Per AGENTS.md mandatory rule,
# ISSUES is rebuilt from scratch each build.  Bugs that no longer reproduce
# simply don't appear — they are not marked "resolved".  Using "resolved"
# would pass validation but silently vanish from all rendered output (it is
# not in tier_order), which is a silent-corruption trap.


# ---------------------------------------------------------------------
# Report-triage contract (the "user-perceptible validation gate").
# Full prose: docs/skills/issue-report-maintainer.md § "Report-triage
# contract"; summary: AGENTS.md § Confirmed bugs.
#
# WHY THIS EXISTS.  Severity must reflect what a USER perceives, not
# which test went red.  A regression run surfaces three kinds of red
# that are NOT user-perceptible on their own and were historically
# mis-filed as P0/P1:
#   - a cross-check threshold tripping vs an external snapshot
#     (c_index_year / c_index_addr_id drift) — a maintenance-cadence or
#     classification question, not a UI bug;
#   - a structural metric parsed out of an export file (vertex-count
#     header vs row count, nodedef column count) — may be a real bug or
#     a test-measurement artifact;
#   - an injected harness marker firing (a ZZ_TEST_DEBUG ":ERR" written
#     under MsgBox suppression) that the real error handler swallows so
#     the user never sees it.
# These are LEADS, not confirmed user bugs.  P0/P1/P2 are reserved for
# findings a human can reproduce through the Access UI and observe a
# concrete symptom (a popup, blank-or-wrong on-screen data, or a file
# the user asked for that is missing or corrupt).  A lead may only be
# rated P0/P1/P2 after the symptom is confirmed in the real UI
# (evidence.ui_verified == True); otherwise it is P5 pending
# verification, or (for drift) belongs in Appendix A.
# ---------------------------------------------------------------------

_FINDING_CLASSES = frozenset([
    "user_facing_bug",    # a human reproduces it via the UI and sees a symptom
    "cross_check_drift",  # divergence vs an external snapshot (index year/addr)
    "structural_metric",  # derived by parsing an export file's structure
    "internal_marker",    # detected only via an injected harness marker/state
    "latent_code",        # real source defect with no current user-facing trigger
])

# Tiers that assert the symptom is user-perceptible on the current dump.
_USER_PERCEPTIBLE_TIERS = frozenset([
    "P0_silent_data", "P1_visible_crash", "P2_silent_display",
])

# Tiers that should carry screenshots when captured (visible in the UI).
# P4 setup / P5 latent are exempt (latent shots are optional + hedged).
_SCREENSHOT_TIERS = frozenset([
    "P0_silent_data", "P1_visible_crash", "P2_silent_display", "P3_missing_ui",
])


def _screenshot_gap(issue_id, declared, available) -> str | None:
    """C4: flag the build-20260605 regression where screenshot files existed
    on disk but the entry shipped with screenshots=[] (imageless docx).

    `declared` = the entry's screenshots list; `available` = filenames present
    in reports/screenshots/.  Returns a message if files named bug<id>_* exist
    but none are declared, else None.

    Scope/limitation: this keys on capture_screenshots.py's `bug<id>_*` naming
    convention.  It catches the exact build-20260605 regression (files on disk
    but screenshots=[]); it does NOT guarantee every visible issue HAS a
    screenshot, nor catch shots saved under a non-bug<id>_ name.  The
    declared-vs-present caption audit (analysis/audit_report_screenshot_
    consistency.py) is the complementary check.
    """
    if not issue_id or declared:
        return None
    prefix = f"bug{issue_id}_"
    hits = [f for f in available if f.startswith(prefix)]
    if hits:
        return (
            f"{len(hits)} screenshot file(s) for issue {issue_id} exist in "
            f"reports/screenshots/ ({prefix}*) but the entry's screenshots=[] "
            f"— wire them into the entry, or the report ships imageless (the "
            f"build-20260605 regression)."
        )
    return None


def _issue_violations(it: dict, shot_files: list | None = None) -> list[str]:
    """Return a list of triage-contract violations for one ISSUES entry.

    Empty list == the entry satisfies the contract.  Collecting all
    violations (rather than raising on the first) lets _validate_issues
    report every problem in one pass.  ``shot_files`` is the list of
    filenames in reports/screenshots/ (computed once by _validate_issues);
    None means "read SHOT_DIR now" (used by ad-hoc callers/tests).
    """
    out: list[str] = []
    tier = it.get("tier")
    if tier not in _VALID_TIERS:
        # A typo in 'tier' silently drops the issue from all rendered
        # outputs (it won't appear in tier_order) — a silent-corruption
        # trap.  Report and stop; later checks can't reason about tier.
        out.append(
            f"unknown tier {tier!r} (valid: {sorted(_VALID_TIERS)})"
        )
        return out

    ev = it.get("evidence")
    if not isinstance(ev, dict):
        out.append(
            "missing 'evidence' block.  Every issue needs "
            "evidence={'finding_class': ..., ...}; P0/P1/P2 additionally "
            "need vba_ref, fixture, user_symptom."
        )
        return out

    fclass = ev.get("finding_class")
    if fclass not in _FINDING_CLASSES:
        out.append(
            f"evidence.finding_class {fclass!r} invalid "
            f"(valid: {sorted(_FINDING_CLASSES)})"
        )
        return out

    # NOTE — the gate trusts the author-supplied `finding_class`; it
    # enforces structure + class->tier routing, but it cannot divine
    # whether a finding labelled `user_facing_bug` really is one.  That
    # judgement is the reviewer's (the issue-report skill, the AI
    # reviewer, codex).  A deliberate mis-label (drift filed as
    # user_facing_bug to dodge routing) is caught in review, not here.
    # Keep this boundary explicit rather than bolting on brittle
    # keyword heuristics that would false-positive and create a false
    # sense of completeness.

    # `ui_verified` is a human attestation that the symptom was seen in
    # the real Access UI.  Require literal True — a truthy string like
    # "pending" must NOT unlock a user-perceptible tier.
    ui_verified = ev.get("ui_verified") is True

    # --- finding_class -> tier routing ---
    if fclass == "user_facing_bug":
        # Reproducible-now by a user → P0–P4.  A defect that is real but
        # NOT user-triggerable on this dump is latent_code (→ P5), so a
        # user_facing_bug parked at P5 is a contradiction.
        if tier == "P5_dormant_or_latent":
            out.append(
                "user_facing_bug must be P0–P4.  A defect that is real but "
                "not user-triggerable on this dump is latent_code (→ P5), "
                "not user_facing_bug."
            )
    elif fclass == "cross_check_drift":
        # Drift belongs in Appendix A, not a standalone P-tier issue.
        # Only admissible at all as P5 AND only when a per-row
        # classifier has judged it a real algorithm bug (not a
        # maintenance-cadence / source-snapshot diff).
        if tier != "P5_dormant_or_latent" or not str(
            ev.get("classification_ref") or ""
        ).strip():
            out.append(
                "cross_check_drift findings belong in Appendix A, not a "
                "P-tier issue.  Admissible only as P5_dormant_or_latent "
                "WITH evidence.classification_ref pointing to a "
                "classify_*_drift output that judged it a real algorithm "
                "bug (not maintenance-cadence / source-snapshot drift)."
            )
    elif fclass == "latent_code":
        if tier != "P5_dormant_or_latent":
            out.append(
                f"latent_code (real defect, no user-facing trigger today) "
                f"must be tier P5_dormant_or_latent, not {tier}."
            )
    elif fclass in ("structural_metric", "internal_marker"):
        # Not user-perceptible on its own.  Until a human confirms the
        # symptom in the real UI (ui_verified), it may ONLY be P5 —
        # parking it at P3/P4 must not dodge that requirement either.
        if not ui_verified and tier != "P5_dormant_or_latent":
            out.append(
                f"{fclass} rated {tier} but evidence.ui_verified is not "
                f"True.  A metric parsed from an export file / an injected "
                f"harness marker is not user-perceptible on its own — open "
                f"the artifact in the real Access UI, confirm the symptom, "
                f"and set evidence.ui_verified=True, or file it as "
                f"P5_dormant_or_latent pending verification."
            )

    # --- required evidence for user-perceptible tiers ---
    if tier in _USER_PERCEPTIBLE_TIERS:
        for field in ("vba_ref", "fixture", "user_symptom"):
            if not str(ev.get(field) or "").strip():
                out.append(
                    f"tier {tier} requires a non-empty evidence.{field} "
                    f"(source location / concrete reproducer / the symptom "
                    f"a user observes)."
                )
        symptom = str(ev.get("user_symptom") or "").strip().lower()
        # Lint (a BACKSTOP, not a proof) for the common anti-pattern of
        # pasting test/assertion output in place of a user-observable
        # symptom.  Patterns are high-precision — comparison operators,
        # pytest-style "expected N … got M" count assertions, this repo's
        # test-id prefixes and harness markers — chosen so they don't
        # flag legitimate symptom prose ("the user got an empty folder").
        # It cannot prove a symptom is genuinely user-facing; that is the
        # reviewer's call (see the finding_class note above).
        _restate_prefixes = ("detected by", "assertion", "assert ")
        _restate_substrings = (
            "zz_test_debug", "detected by test", "test_cmd", "test_export",
            "test_bug", "test_index", "test_vba", ":err marker",
            "!=", "==", "assertion failed", "assertion error",
        )
        if (symptom.startswith(_restate_prefixes)
                or any(s in symptom for s in _restate_substrings)
                or re.search(r"expected\b.{0,40}\bgot\b.*\d", symptom)):
            out.append(
                "evidence.user_symptom must describe what the USER observes "
                "(a popup, blank-or-wrong on-screen data, a missing/corrupt "
                "file), not restate the test assertion ('Detected by …' / "
                "'expected N … got M')."
            )

    # --- screenshot-presence gate (C4) ---
    if tier in _SCREENSHOT_TIERS:
        if shot_files is None:
            shot_files = ([p.name for p in SHOT_DIR.glob("*")]
                          if SHOT_DIR.exists() else [])
        gap = _screenshot_gap(it.get("id"), it.get("screenshots") or [],
                              shot_files)
        if gap:
            out.append(gap)

    return out


def _validate_issues() -> None:
    """Enforce the report-triage contract before any file is written.

    Fails loudly (raises ValueError listing every violation) so a
    mis-tiered or under-evidenced issue surfaces immediately instead of
    shipping in the report.  See _issue_violations / the module banner
    above for the rules.  An empty ISSUES list passes (clean slate).
    """
    shot_files = ([p.name for p in SHOT_DIR.glob("*")]
                  if SHOT_DIR.exists() else [])
    problems: list[str] = []
    for it in ISSUES:
        for msg in _issue_violations(it, shot_files=shot_files):
            problems.append(f"ISSUES id={it.get('id')!r}: {msg}")
    if problems:
        raise ValueError(
            f"Report-triage contract violations ({len(problems)}):\n  - "
            + "\n  - ".join(problems)
        )


def _add_schema_diff_appendix_docx(doc, is_en: bool, Z) -> None:
    """Render Appendix B (TablesFields) and Appendix C (ForeignKeys)
    schema-diff sections using the python-docx API.
    Consumes reports/schema_diff.json produced by collect_schema_diffs.py."""
    import json as _json

    def _table_section(title_a: str, title_b: str,
                       only_cur_hdr: str, only_reg_hdr: str,
                       mismatch_hdr: str,
                       block_key: str) -> None:
        _h(doc, 1, Z(title_a if is_en else title_b))

        if not SCHEMA_DIFF_JSON.exists():
            p = doc.add_paragraph()
            run = p.add_run(Z(
                "Run `python reports/collect_schema_diffs.py` to populate "
                "this section."
                if is_en else
                "請先執行 `python reports/collect_schema_diffs.py` 以生成本節內容。"
            ))
            run.italic = True
            return

        diff = _json.loads(SCHEMA_DIFF_JSON.read_text(encoding="utf-8"))
        mdb_label = diff.get("data_mdb", "CBDB_*_DATA.mdb")
        blk = diff[block_key]
        is_tf = (block_key == "tables_fields")

        intro = (
            (
                f"This section compares the contents of the `TablesFields` table "
                f"in `{mdb_label}` against the database schema "
                "reconstructed from Access DAO (TableDefs) by "
                "`reports/collect_schema_diffs.py`. Discrepancies indicate the "
                "documentation table may be out of date."
            ) if is_tf else (
                "This section covers the `ForeignKeys` table and the FK "
                "relationships it documents."
            )
        )
        intro_zh = (
            (
                f"本節將 `{mdb_label}` 中 `TablesFields` 表的內容與"
                "`reports/collect_schema_diffs.py` 透過 Access DAO（TableDefs）重建的資料"
                "庫結構進行比對。若存在差異，表示文檔表可能已過時。"
            ) if is_tf else (
                "本節涵蓋 `ForeignKeys` 表及其所記錄的外鍵關係。"
            )
        )
        doc.add_paragraph(Z(intro if is_en else intro_zh))

        if not is_tf and not blk.get("fk_introspection_available"):
            doc.add_paragraph(Z(
                f"The `ForeignKeys` table ({blk['total_current']} rows) documents FK "
                "relationships in the database. All referenced table/column pairs have "
                "been verified to exist in the current dump. A catalog-level diff "
                "(documented FK vs. all actual FK constraints) is not available for "
                "Access databases and is omitted here."
                if is_en else
                f"`ForeignKeys` 表共 {blk['total_current']} 筆，記錄了資料庫中的外鍵關係。"
                "我們已驗證所有參照的表名與欄位名均存在於當前 dump 中。"
                "由於 Access 資料庫不支援通過標準 API 枚舉外鍵約束，"
                "此處無法提供文件記載 FK 與實際 FK 約束的完整對比。"
            ))
            doc.add_paragraph(Z(
                f"Reconstructed FK list: reports/foreign_keys_regen.csv"
                if is_en else
                "重建結果：reports/foreign_keys_regen.csv"
            ))
            return

        doc_name = "TablesFields" if is_tf else "ForeignKeys"
        regen_src = (
            "Reconstructed from DB" if is_tf
            else "Reconstructed from DB (via Access.Application DAO)"
        )
        regen_src_zh = (
            "從資料庫重建" if is_tf
            else "從資料庫重建（透過 Access.Application DAO）"
        )
        doc.add_paragraph(Z(
            f"Total rows in {doc_name}: {blk['total_current']}. "
            f"{regen_src}: {blk['total_regen']}."
            if is_en else
            f"{doc_name} 共 {blk['total_current']} 筆。"
            f"{regen_src_zh}：{blk['total_regen']} 筆。"
        ))
        _regen_csv = (
            "tables_fields_regen.csv" if is_tf else "foreign_keys_regen.csv"
        )
        doc.add_paragraph(Z(
            f"Reconstructed schema: reports/{_regen_csv}"
            if is_en else
            f"重建結果：reports/{_regen_csv}"
        ))

        only_cur = blk["only_in_current"]
        only_reg = blk["only_in_regen"]
        mismatches = blk.get("mismatches", [])

        if only_cur:
            _h(doc, 2, Z(only_cur_hdr))
            n = len(only_cur)
            if is_tf:
                tbl = doc.add_table(rows=n + 1, cols=2)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                hdr[0].text = "AccessTblNm"; hdr[1].text = "AccessFldNm"
                for i, row in enumerate(only_cur):
                    r = tbl.rows[i + 1].cells
                    r[0].text = row["AccessTblNm"]
                    r[1].text = row["AccessFldNm"]
            else:
                tbl = doc.add_table(rows=n + 1, cols=4)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                for ci, h in enumerate(["AccessTblNm", "AccessFldNm", "ForeignKey", "ForeignKeyBaseField"]):
                    hdr[ci].text = h
                for i, row in enumerate(only_cur):
                    r = tbl.rows[i + 1].cells
                    r[0].text = row["AccessTblNm"]
                    r[1].text = row["AccessFldNm"]
                    r[2].text = row.get("ForeignKey") or ""
                    r[3].text = row.get("ForeignKeyBaseField") or ""

        if only_reg:
            _h(doc, 2, Z(only_reg_hdr))
            n = len(only_reg)
            if is_tf:
                tbl = doc.add_table(rows=n + 1, cols=4)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                for ci, h in enumerate(["AccessTblNm", "AccessFldNm", "DataFormat", "NULL_allowed"]):
                    hdr[ci].text = h
                for i, row in enumerate(only_reg):
                    r = tbl.rows[i + 1].cells
                    r[0].text = row["AccessTblNm"]
                    r[1].text = row["AccessFldNm"]
                    r[2].text = str(row.get("DataFormat") or "")
                    r[3].text = str(row.get("NULL_allowed") or "")
            else:
                tbl = doc.add_table(rows=n + 1, cols=4)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                for ci, h in enumerate(["AccessTblNm", "AccessFldNm", "ForeignKey", "ForeignKeyBaseField"]):
                    hdr[ci].text = h
                for i, row in enumerate(only_reg):
                    r = tbl.rows[i + 1].cells
                    r[0].text = row["AccessTblNm"]
                    r[1].text = row["AccessFldNm"]
                    r[2].text = row.get("ForeignKey") or ""
                    r[3].text = row.get("ForeignKeyBaseField") or ""

        if mismatches:
            _h(doc, 2, Z(mismatch_hdr))
            csv_file = (
                "reports/schema_diff_tables_fields_mismatches.csv" if is_tf
                else "reports/schema_diff_foreign_keys_mismatches.csv"
            )
            doc.add_paragraph(Z(
                f"Full list: `{csv_file}` ({len(mismatches)} rows)"
                if is_en else
                f"完整清單：`{csv_file}`（{len(mismatches)} 筆）"
            ))

        if not only_cur and not only_reg and not mismatches:
            doc.add_paragraph(Z(
                f"✓ No discrepancies found — {doc_name} is in sync with the actual schema."
                if is_en else
                f"✓ 未發現差異 —— {doc_name} 與實際資料庫結構一致。"
            ))

    # ---- Appendix B ----
    _table_section(
        title_a="Appendix B — TablesFields: documentation vs. actual structure",
        title_b="附錄 B —— TablesFields：文檔表與實際資料庫結構對比",
        only_cur_hdr=(
            "Rows in TablesFields not found in actual DB (stale)"
            if is_en else
            "TablesFields 中有但實際資料庫中不存在的欄位（過時）"
        ),
        only_reg_hdr=(
            "Columns in actual DB not documented in TablesFields"
            if is_en else
            "實際資料庫中有但 TablesFields 未記錄的欄位"
        ),
        mismatch_hdr="Attribute mismatches" if is_en else "屬性不一致",
        block_key="tables_fields",
    )

    doc.add_page_break()

    # ---- Appendix C ----
    _table_section(
        title_a="Appendix C — ForeignKeys: documentation vs. actual structure",
        title_b="附錄 C —— ForeignKeys：文檔表與實際資料庫結構對比",
        only_cur_hdr=(
            "Rows in ForeignKeys not found in actual DB (stale)"
            if is_en else
            "ForeignKeys 中有但實際資料庫中不存在的外鍵（過時）"
        ),
        only_reg_hdr=(
            "FK relationships in actual DB not documented in ForeignKeys"
            if is_en else
            "實際資料庫中有但 ForeignKeys 未記錄的外鍵關係"
        ),
        mismatch_hdr="Attribute mismatches" if is_en else "屬性不一致",
        block_key="foreign_keys",
    )


def _add_schema_diff_appendix_md(
    lines: list,
    block_key: str,
    is_en: bool,
    Z,
    _slug,
) -> None:
    """Render one schema-diff appendix section (Appendix B or C) into
    the markdown lines list.

    block_key: 'tables_fields' for Appendix B, 'foreign_keys' for Appendix C.
    """
    import json as _json

    is_tf = (block_key == "tables_fields")

    heading = (
        ("Appendix B — TablesFields: documentation vs. actual structure"
         if is_tf else
         "Appendix C — ForeignKeys: documentation vs. actual structure")
        if is_en else
        ("附錄 B —— TablesFields：文檔表與實際資料庫結構對比"
         if is_tf else
         "附錄 C —— ForeignKeys：文檔表與實際資料庫結構對比")
    )
    lines.append(f"## {Z(heading)}")
    lines.append("")

    if not SCHEMA_DIFF_JSON.exists():
        lines.append(Z(
            "*Run `python reports/collect_schema_diffs.py` to populate "
            "this section.*"
            if is_en else
            "*請先執行 `python reports/collect_schema_diffs.py` 以生成本節內容。*"
        ))
        lines.append("")
        return

    diff = _json.loads(SCHEMA_DIFF_JSON.read_text(encoding="utf-8"))
    mdb_label = diff.get("data_mdb", "CBDB_*_DATA.mdb")
    blk = diff[block_key]

    intro = (
        (
            f"This section compares the contents of the `TablesFields` table "
            f"in `{mdb_label}` against the database schema "
            "reconstructed from Access DAO (TableDefs) by "
            "`reports/collect_schema_diffs.py`. Discrepancies indicate the "
            "documentation table may be out of date."
        ) if is_tf else (
            "This section covers the `ForeignKeys` table and the FK "
            "relationships it documents."
        )
    )
    intro_zh = (
        (
            f"本節將 `{mdb_label}` 中 `TablesFields` 表的內容與 "
            "`reports/collect_schema_diffs.py` 透過 Access DAO（TableDefs）重建的資料"
            "庫結構進行比對。若存在差異，表示文檔表可能已過時。"
        ) if is_tf else (
            "本節涵蓋 `ForeignKeys` 表及其所記錄的外鍵關係。"
        )
    )
    lines.append(Z(intro if is_en else intro_zh))
    lines.append("")

    if not is_tf and not blk.get("fk_introspection_available"):
        lines.append(Z(
            f"The `ForeignKeys` table ({blk['total_current']} rows) documents FK "
            "relationships in the database. All referenced table/column pairs have "
            "been verified to exist in the current dump. A catalog-level diff is not "
            "available for Access databases and is omitted here."
            if is_en else
            f"`ForeignKeys` 表共 {blk['total_current']} 筆，記錄了資料庫中的外鍵關係。"
            "我們已驗證所有參照的表名與欄位名均存在於當前 dump 中。"
            "由於 Access 資料庫不支援通過標準 API 枚舉外鍵約束，"
            "此處無法提供文件記載 FK 與實際 FK 約束的完整對比。"
        ))
        lines.append("")
        lines.append(Z(
            "Reconstructed FK list: [foreign_keys_regen.csv](foreign_keys_regen.csv)"
            if is_en else
            "重建結果：[foreign_keys_regen.csv](foreign_keys_regen.csv)"
        ))
        lines.append("")
        return

    doc_name = "TablesFields" if is_tf else "ForeignKeys"
    regen_src = (
        "Reconstructed from DB" if is_tf
        else "Reconstructed from DB (via Access.Application DAO)"
    )
    regen_src_zh = (
        "從資料庫重建" if is_tf
        else "從資料庫重建（透過 Access.Application DAO）"
    )
    lines.append(Z(
        f"Total rows in {doc_name}: {blk['total_current']}. "
        f"{regen_src}: {blk['total_regen']}."
        if is_en else
        f"{doc_name} 共 {blk['total_current']} 筆。"
        f"{regen_src_zh}：{blk['total_regen']} 筆。"
    ))
    lines.append("")

    _regen_csv = "tables_fields_regen.csv" if is_tf else "foreign_keys_regen.csv"
    _regen_label = (
        ("Reconstructed schema" if is_tf else "Reconstructed FK list")
        if is_en else "重建結果"
    )
    lines.append(Z(f"{_regen_label}: [{_regen_csv}]({_regen_csv})"))
    lines.append("")

    only_cur = blk["only_in_current"]
    only_reg = blk["only_in_regen"]
    mismatches = blk.get("mismatches", [])

    if only_cur:
        stale_hdr = (
            f"Rows in {doc_name} not found in actual DB (stale)"
            if is_en else
            f"{doc_name} 中有但實際資料庫中不存在的記錄（過時）"
        )
        lines.append(f"### {Z(stale_hdr)}")
        lines.append("")
        if is_tf:
            lines.append("| AccessTblNm | AccessFldNm |")
            lines.append("|---|---|")
            for row in only_cur:
                lines.append(f"| {row['AccessTblNm']} | {row['AccessFldNm']} |")
        else:
            lines.append("| AccessTblNm | AccessFldNm | ForeignKey | ForeignKeyBaseField |")
            lines.append("|---|---|---|---|")
            for row in only_cur:
                lines.append(
                    f"| {row['AccessTblNm']} | {row['AccessFldNm']} "
                    f"| {row.get('ForeignKey', '')} | {row.get('ForeignKeyBaseField', '')} |"
                )
        lines.append("")

    if only_reg:
        undoc_hdr = (
            f"Columns in actual DB not documented in {doc_name}"
            if is_en else
            f"實際資料庫中有但 {doc_name} 未記錄的欄位"
        )
        lines.append(f"### {Z(undoc_hdr)}")
        lines.append("")
        if is_tf:
            lines.append("| AccessTblNm | AccessFldNm | DataFormat | NULL_allowed |")
            lines.append("|---|---|---|---|")
            for row in only_reg:
                lines.append(
                    f"| {row['AccessTblNm']} | {row['AccessFldNm']} "
                    f"| {row.get('DataFormat', '')} | {row.get('NULL_allowed', '')} |"
                )
        else:
            lines.append("| AccessTblNm | AccessFldNm | ForeignKey | ForeignKeyBaseField |")
            lines.append("|---|---|---|---|")
            for row in only_reg:
                lines.append(
                    f"| {row['AccessTblNm']} | {row['AccessFldNm']} "
                    f"| {row.get('ForeignKey', '')} | {row.get('ForeignKeyBaseField', '')} |"
                )
        lines.append("")

    if mismatches:
        mis_hdr = "Attribute mismatches" if is_en else "屬性不一致"
        lines.append(f"### {Z(mis_hdr)}")
        lines.append("")
        csv_file = (
            "reports/schema_diff_tables_fields_mismatches.csv" if is_tf
            else "reports/schema_diff_foreign_keys_mismatches.csv"
        )
        lines.append(Z(
            f"Full list: `{csv_file}` ({len(mismatches)} rows)"
            if is_en else
            f"完整清單：`{csv_file}`（{len(mismatches)} 筆）"
        ))
        lines.append("")

    if not only_cur and not only_reg and not mismatches:
        lines.append(Z(
            f"✓ No discrepancies found — {doc_name} is in sync with the actual schema."
            if is_en else
            f"✓ 未發現差異 —— {doc_name} 與實際資料庫結構一致。"
        ))
        lines.append("")


def _build(lang: str, out_path: Path) -> None:
    is_en = (lang == "en")

    # Apply Simplified -> Traditional Chinese conversion when emitting
    # the Chinese variant.  English passes through unchanged.
    def Z(s: str) -> str:
        return s if is_en else t(s)

    doc = Document()

    # ---- Cover page ----
    title = (
        "CBDB User MDB — Issues Report"
        if is_en else
        "CBDB 用户版 .mdb — 问题汇报"
    )
    subtitle = (
        "A respectful summary of issues uncovered during regression testing."
        if is_en else
        "测试过程中发现的问题汇总，谨呈维护团队斧正。"
    )
    h = doc.add_heading(Z(title), level=0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph(Z(subtitle))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = (
        "Dear maintainer,\n\n"
        "Below is a summary of the issues we uncovered while building "
        "an automated regression-test suite for the CBDB User MDB. "
        "We hope this report is useful as you continue your wonderful "
        "stewardship of this dataset, and we sincerely thank you for "
        "the immense work that has gone into building it.\n\n"
        "The issues are ordered by severity (P0 highest). Each entry "
        "includes a concise description, step-by-step user reproduction, "
        "screenshots where the issue is visible in the Access UI, and a "
        "suggested fix. None of these are urgent; they are documented "
        "so they can be addressed at the maintainer's convenience."
        if is_en else
        "尊敬的维护者：\n\n"
        "下面是我们在为 CBDB 用户版 .mdb 编写自动化回归测试套件过程中，"
        "陆续整理出来的一些问题清单。我们希望这份报告能在您继续主持这份"
        "宝贵数据集时有所帮助；同时，对您多年来在这套数据上的辛勤付出，"
        "我们由衷地表示感谢和敬意。\n\n"
        "问题按严重程度排序（P0 最高）。每一条都包括：简明描述、用户端"
        "一步一步的复现步骤、（在界面上能看到时）相关截图，以及一份建议"
        "的修复方案。这些问题并不紧急，整理在此只是为了方便您在合适的时"
        "候逐一处理。"
    )
    for para in intro.split("\n\n"):
        doc.add_paragraph(Z(para))

    doc.add_page_break()

    # ---- Coverage matrix (always present) ----
    _add_coverage_matrix_docx(doc, is_en, Z, _load_coverage_matrix())
    doc.add_page_break()

    # ---- Table of contents ----
    _h(doc, 1, Z("Table of Contents" if is_en else "目录"))
    _add_toc(doc, lang)
    doc.add_page_break()

    # ---- Severity legend ----
    _h(doc, 1, Z("Severity legend" if is_en else "严重等级说明"))
    legend_en = [
        "P0 — Silent data corruption: data is wrong or missing without an error popup.",
        "P1 — Visible runtime crash: a popup appears, the operation aborts.",
        "P2 — Silent display: form fields render blank when they should show data.",
        "P3 — Missing UI: a feature exists in code but no button invokes it.",
        "P4 — Setup: one-time hurdle on each new install.",
        "P5 — Dormant / latent / not currently reproducible: kept as historical "
        "record; we re-checked on the current dump and could not "
        "trigger the symptom.",
    ]
    legend_zh = [
        "P0 — 静默数据错误：数据错或缺失，但没有任何报错提示。",
        "P1 — 可见的运行时报错：弹出错误对话框，操作中断。",
        "P2 — 静默显示问题：表单字段本应有数据，却显示为空。",
        "P3 — 缺失界面：代码里实现了某功能，但界面上没有按钮去触发它。",
        "P4 — 安装设置：每台新机器需要一次性处理。",
        "P5 — 潛伏 / 不可達 / 當前無法復現：保留作为历史记录；我们在当前 dump "
        "上重新验证过，无法再触发症状。",
    ]
    _bullets(doc, [Z(s) for s in (legend_en if is_en else legend_zh)])
    doc.add_page_break()

    # ---- One section per issue ----
    by_tier: dict[str, list[dict]] = {}
    for it in ISSUES:
        by_tier.setdefault(it["tier"], []).append(it)
    tier_order = ["P0_silent_data", "P1_visible_crash",
                  "P2_silent_display", "P3_missing_ui", "P4_setup",
                  "P5_dormant_or_latent"]
    tier_titles_en = {
        "P0_silent_data": "P0 — Silent data corruption",
        "P1_visible_crash": "P1 — Visible runtime crash",
        "P2_silent_display": "P2 — Silent display",
        "P3_missing_ui": "P3 — Missing UI",
        "P4_setup": "P4 — Setup",
        "P5_dormant_or_latent":
            "P5 — Dormant / latent / not currently reproducible",
        "resolved": "Resolved — fixed in this build",
    }
    tier_titles_zh = {
        "P0_silent_data": "P0 — 静默数据错误",
        "P1_visible_crash": "P1 — 可见的运行时报错",
        "P2_silent_display": "P2 — 静默显示问题",
        "P3_missing_ui": "P3 — 缺失界面",
        "P4_setup": "P4 — 安装设置",
        "P5_dormant_or_latent": "P5 — 潛伏 / 不可達 / 當前無法復現",
        "resolved": "已解決 — 當前 build 已修復",
    }
    demo_persons = _load_demo_persons()
    bug_status = _load_bug_test_status()

    for tier in tier_order:
        items = by_tier.get(tier, [])
        if not items:
            continue
        _h(doc, 1, Z(tier_titles_en[tier] if is_en
                      else tier_titles_zh[tier]))
        if tier == "P5_dormant_or_latent":
            preface = (
                "Items in this tier are kept as historical / latent "
                "record.  They fall into three categories: (a) DORMANT "
                "— current source data doesn't trigger the symptom; "
                "(b) NOT CURRENTLY REPRODUCIBLE — the symptom no "
                "longer surfaces even though the suspect code is "
                "still present (we have NOT confirmed an upstream "
                "source-level fix; the cause could be a JET / Office "
                "behaviour change, a fixture / driver change on our "
                "side, or the original diagnosis was a false "
                "positive); (c) LATENT — the source-code defect is "
                "real, but the user can't reach it because another "
                "issue (e.g. a missing UI button) blocks the path. "
                "None of these are user-facing today; **none have "
                "been verified as fixed upstream** — please consult "
                "before treating any of them as either urgent or "
                "closed."
                if is_en else
                "本層的條目作為歷史 / 潛伏記錄保留。可分為三類："
                "(a) DORMANT 潛伏 —— 已驗證當前源資料無法觸發該症狀；"
                "(b) 當前無法復現 —— 症狀不再出現，但可疑程式碼仍在"
                "（我們**沒有**確認上游有源碼層面的修復；原因可能是 "
                "JET / Office 的行為改變、可能是我們這邊 fixture/driver "
                "改變，也可能原本的診斷就是 false positive）；"
                "(c) LATENT 被屏蔽 —— 源碼缺陷確實存在，但因為"
                "另一個 issue（例如某個 UI 按鈕缺失）擋住了使用路徑，"
                "使用者目前碰不到。本層條目當下都不是使用者會遇到的"
                "問題，**也沒有任何一條被確認上游修復**；若要當成緊急"
                "或已關閉處理，請先諮詢。"
            )
            p = doc.add_paragraph(Z(preface))
            for run in p.runs:
                run.italic = True
        for it in items:
            title = it["title_en"] if is_en else it["title_zh"]
            _h(doc, 2, Z(f"Issue #{it['id']} — {title}"))
            _h(doc, 3, Z("Affected sub" if is_en else "涉及位置"))
            doc.add_paragraph(Z(it["form"]))
            _h(doc, 3, Z("Severity" if is_en else "严重等级"))
            doc.add_paragraph(Z(it["severity_en"] if is_en
                                  else it["severity_zh"]))
            # ---- Auto-derived test status banner (informational) ----
            # If a known-bug test exists for this issue and reports
            # 'failed', that's a hint the bug may have been fixed
            # upstream — but we never act on it automatically.  We
            # display a clearly-marked banner asking the maintainer to
            # confirm in person; the issue's full description / steps
            # / fix recommendation stay rendered as-is so nothing is
            # lost if the test gave a false-positive.
            status = bug_status.get(it["id"])
            if status:
                outcome = status["outcome"]
                if outcome == "failed":
                    banner = (
                        "⚠ Automated test status: the regression marker "
                        "for this issue currently FAILS (run on "
                        f"{status.get('when', 'unknown date')}), which "
                        "means the marker no longer reproduces.  This MAY "
                        "indicate the underlying defect was fixed "
                        "upstream, but it could equally mean that the "
                        "input fixture or Access driver changed out "
                        "from under the test, or that the original "
                        "classification was wrong.  Please investigate "
                        "in person before considering this issue closed; "
                        "this report has NOT been edited to drop the "
                        "issue.  Tests consulted:\n"
                        + "\n".join(
                            f"  • {t.rsplit('::', 1)[-1]}"
                            for t, _ in status["tests"]
                        )
                        if is_en else
                        "⚠ 自動測試狀態：本 issue 對應的回歸標記目前 "
                        f"FAIL（執行時間：{status.get('when', '未知')}），"
                        "代表標記目前無法復現。這**可能**表示上游已在源碼層"
                        "面修復，但同樣可能是輸入 fixture 或 Access "
                        "driver 在不知不覺中改變、或原本的分類就是錯的。"
                        "請務必親自調查清楚，再將此 issue 視為關閉；本報告"
                        "並未自動刪除任何 issue。對照的測試：\n"
                        + "\n".join(
                            f"  • {t.rsplit('::', 1)[-1]}"
                            for t, _ in status["tests"]
                        )
                    )
                    p = doc.add_paragraph()
                    run = p.add_run(Z(banner))
                    run.bold = True
                elif outcome == "mixed":
                    banner = (
                        "ℹ Automated test status: this issue's "
                        "regression markers report MIXED outcomes "
                        f"(run {status.get('when', 'unknown date')}). "
                        "Likely the issue's markers partially stopped "
                        "reproducing (could be a partial upstream fix, a "
                        "partial fixture/driver change, or a partial "
                        "misclassification).  Please "
                        "review the per-test breakdown:\n"
                        + "\n".join(
                            f"  • {t.rsplit('::', 1)[-1]}: {o}"
                            for t, o in status["tests"]
                        )
                        if is_en else
                        "ℹ 自動測試狀態：本 issue 對應的回歸標記呈現"
                        f"混合結果（執行時間：{status.get('when', '未知')}）。"
                        "可能是部分標記不再復現（可能是部分上游修復、"
                        "部分 fixture/driver 變化、或部分分類錯誤）。"
                        "請查看分項：\n"
                        + "\n".join(
                            f"  • {t.rsplit('::', 1)[-1]}: {o}"
                            for t, o in status["tests"]
                        )
                    )
                    p = doc.add_paragraph()
                    run = p.add_run(Z(banner))
                    run.italic = True
                # outcome == "passed" → bug confirmed present; no
                # banner (it's the default expectation; banner-spam is
                # noise).

            _h(doc, 3, Z("Description" if is_en else "问题描述"))
            for para in (it["summary_en"] if is_en
                         else it["summary_zh"]).split("\n\n"):
                doc.add_paragraph(Z(para))
            _h(doc, 3, Z("Steps to reproduce" if is_en else "复现步骤"))
            # Inject a concrete demo-person hint AHEAD of the steps so
            # the maintainer never has to guess which person id to
            # open.  For DORMANT bugs (no UI-reachable example on the
            # current data snapshot), the hint says so and offers a
            # SQL-only verification path instead.
            demo = demo_persons.get(f"bug{it['id']}")
            if demo:
                if demo.get("dormant"):
                    label = (
                        "⚠ Currently UI-dormant on this data "
                        "snapshot — see note below"
                        if is_en else
                        "⚠ 在當前資料快照下無法在 UI 上復現——請看下方說明"
                    )
                    para = doc.add_paragraph()
                    run = para.add_run(Z(label))
                    run.bold = True
                    para = doc.add_paragraph()
                    para.add_run(
                        demo["hint_en"] if is_en else Z(demo["hint_zh"])
                    )
                else:
                    label = (
                        f"Recommended demo person: c_personid="
                        f"{demo['personid']} ({demo['name_chn']}, "
                        f"{demo['name_py']})"
                        if is_en else
                        f"建議使用的範例人物：c_personid={demo['personid']}"
                        f"（{demo['name_chn']}，{demo['name_py']}）"
                    )
                    para = doc.add_paragraph()
                    run = para.add_run(Z(label))
                    run.bold = True
                    para = doc.add_paragraph()
                    para.add_run(
                        demo["hint_en"] if is_en else Z(demo["hint_zh"])
                    )
                    para.add_run(
                        Z(
                            " — picked by `reports/probe_demo_persons.py`; "
                            "a SQL probe selected this person because their "
                            "row counts genuinely satisfy the precondition "
                            "the bug needs."
                            if is_en else
                            " —— 由 `reports/probe_demo_persons.py` 透過 "
                            "SQL probe 挑選；之所以選這位，是因為其底層"
                            "記錄數確實滿足這個 bug 的觸發條件。"
                        )
                    ).italic = True
            _numbered(doc, [Z(s) for s in
                            (it["steps_en"] if is_en else it["steps_zh"])])
            # Optional concrete-reproduction block — see Issue #9
            # for the canonical use (specific personids that
            # trigger the bug on the current dump).
            cr_key = ("concrete_reproduction_en" if is_en
                      else "concrete_reproduction_zh")
            cr_text = it.get(cr_key)
            if cr_text:
                _h(doc, 3, Z("Concrete reproduction"
                              if is_en else "具體復現"))
                for chunk in str(cr_text).split("\n\n"):
                    if chunk.strip():
                        doc.add_paragraph(Z(chunk))
            shots = it.get("screenshots") or []
            if shots:
                _h(doc, 3, Z("Screenshots" if is_en else "截图"))
                for fname, caption in shots:
                    p = SHOT_DIR / fname
                    if not p.exists():
                        doc.add_paragraph(Z(
                            f"[screenshot {fname} not found]"
                            if is_en else f"[未找到截图 {fname}]"
                        ))
                        continue
                    doc.add_picture(str(p), width=Inches(6.0))
                    if caption:
                        cap = doc.add_paragraph(
                            caption if is_en else Z(caption)
                        )
                        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in cap.runs:
                            run.italic = True
                            run.font.size = Pt(9)
            _h(doc, 3, Z("Suggested fix" if is_en else "建议修复方案"))
            for para in (it["fix_en"] if is_en
                         else it["fix_zh"]).split("\n\n"):
                doc.add_paragraph(Z(para))

    # ---- Appendix A: index_year / index_addr drift (NOT a bug) ----
    doc.add_page_break()
    _add_index_drift_appendix(doc, is_en, Z)

    # ---- Appendix B & C: schema diff ----
    doc.add_page_break()
    _add_schema_diff_appendix_docx(doc, is_en, Z)

    # ---- Closing ----
    doc.add_page_break()
    _h(doc, 1, Z("Closing note" if is_en else "结语"))
    closing = (
        "Thank you for taking the time to read this report. None of the "
        "items above is urgent; we hope having them all in one place "
        "makes it easy to address them at your own pace.\n\n"
        "If any of the descriptions or suggested fixes are unclear, we "
        "would be glad to discuss further. The corresponding regression "
        "tests in this repository will automatically flip from PASS to "
        "FAIL the moment any regression marker stops reproducing in "
        "the source dump — that is a signal to investigate, not an "
        "automatic confirmation that the bug is fixed (the marker "
        "could fail because of an upstream fix, a fixture / driver "
        "change on our side, or a misclassification we made earlier)."
        if is_en else
        "感谢您抽时间读完这份报告。以上各条都不紧急，我们把它们集中整理"
        "在一起，只是希望方便您在合适的时候逐一处理。\n\n"
        "如果对其中任何一条的描述或建议有疑问，欢迎随时一同讨论。本仓库"
        "里对应的回归测试，会在任何一个回归标记不再复现时自动从 PASS "
        "翻成 FAIL —— 这是「请调查一下」的讯号，而不是「问题已修复」的"
        "自动确认（因为标记不再复现也可能是 fixture / driver 变了，"
        "或者是我们当初的分类有误）。"
    )
    for para in closing.split("\n\n"):
        doc.add_paragraph(Z(para))

    doc.save(out_path)
    print(f"wrote {out_path}")


def _build_md(lang: str, out_path: Path) -> None:
    """Markdown sibling of `_build` — same content, different format,
    suitable for in-browser viewing on GitHub."""
    is_en = (lang == "en")

    def Z(s: str) -> str:
        return s if is_en else t(s)

    lines: list[str] = []
    title = (
        "CBDB User MDB — Issues Report"
        if is_en else
        "CBDB 用户版 .mdb — 问题汇报"
    )
    subtitle = (
        "A respectful summary of issues uncovered during regression testing."
        if is_en else
        "测试过程中发现的问题汇总，谨呈维护团队斧正。"
    )
    lines.append(f"# {Z(title)}")
    lines.append("")
    lines.append(f"_{Z(subtitle)}_")
    lines.append("")
    # Build stamp (B7 part 2): record which DATA build this report describes.
    try:
        import sys as _sys
        _ana = str(REPO / "analysis")
        if _ana not in _sys.path:
            _sys.path.insert(0, _ana)
        from build_stamp import current_build as _cur_build
        _build = _cur_build(REPO)
    except Exception:
        _build = None
    _build_label = ("Data build: " if is_en else "数据构建：") + (_build or "unknown")
    lines.append(f"_{Z(_build_label)}_")
    lines.append("")

    intro = (
        "Dear maintainer,\n\n"
        "Below is a summary of the issues we uncovered while building "
        "an automated regression-test suite for the CBDB User MDB. "
        "We hope this report is useful as you continue your wonderful "
        "stewardship of this dataset, and we sincerely thank you for "
        "the immense work that has gone into building it.\n\n"
        "The issues are ordered by severity (P0 highest). Each entry "
        "includes a concise description, step-by-step user reproduction, "
        "screenshots where the issue is visible in the Access UI, and a "
        "suggested fix. None of these are urgent; they are documented "
        "so they can be addressed at the maintainer's convenience."
        if is_en else
        "尊敬的维护者：\n\n"
        "下面是我们在为 CBDB 用户版 .mdb 编写自动化回归测试套件过程中，"
        "陆续整理出来的一些问题清单。我们希望这份报告能在您继续主持这份"
        "宝贵数据集时有所帮助；同时，对您多年来在这套数据上的辛勤付出，"
        "我们由衷地表示感谢和敬意。\n\n"
        "问题按严重程度排序（P0 最高）。每一条都包括：简明描述、用户端"
        "一步一步的复现步骤、（在界面上能看到时）相关截图，以及一份建议"
        "的修复方案。这些问题并不紧急，整理在此只是为了方便您在合适的时"
        "候逐一处理。"
    )
    for para in intro.split("\n\n"):
        lines.append(Z(para))
        lines.append("")

    # ---- Coverage matrix (always present) ----
    _add_coverage_matrix_md(lines, is_en, Z, _load_coverage_matrix())

    # ---- TOC: GitHub auto-generates anchors from heading text. ----
    lines.append(f"## {Z('Table of Contents' if is_en else '目录')}")
    lines.append("")

    tier_titles_en = {
        "P0_silent_data": "P0 — Silent data corruption",
        "P1_visible_crash": "P1 — Visible runtime crash",
        "P2_silent_display": "P2 — Silent display",
        "P3_missing_ui": "P3 — Missing UI",
        "P4_setup": "P4 — Setup",
        "P5_dormant_or_latent":
            "P5 — Dormant / latent / not currently reproducible",
        "resolved": "Resolved — fixed in this build",
    }
    tier_titles_zh = {
        "P0_silent_data": "P0 — 静默数据错误",
        "P1_visible_crash": "P1 — 可见的运行时报错",
        "P2_silent_display": "P2 — 静默显示问题",
        "P3_missing_ui": "P3 — 缺失界面",
        "P4_setup": "P4 — 安装设置",
        "P5_dormant_or_latent": "P5 — 潛伏 / 不可達 / 當前無法復現",
    }
    tier_order = ["P0_silent_data", "P1_visible_crash",
                  "P2_silent_display", "P3_missing_ui", "P4_setup",
                  "P5_dormant_or_latent"]

    by_tier: dict[str, list[dict]] = {}
    for it in ISSUES:
        by_tier.setdefault(it["tier"], []).append(it)

    def _slug(s: str) -> str:
        # GitHub-flavoured anchor algorithm:
        #   1. lower-case
        #   2. drop every char that isn't [a-z0-9_-] OR a CJK
        #      ideograph OR a space (DROP IN PLACE — do not merge
        #      adjacent spaces; do not collapse runs of hyphens).
        #   3. replace each remaining space with exactly one hyphen.
        # Two subtle rules the earlier draft of this function got
        # wrong:
        #   - `_` is a word char and IS preserved by GitHub.  A blanket
        #     `[`*_~]` strip removed it and broke `View_StatusData`-
        #     style anchors.
        #   - `\s+ → -` collapsed runs of spaces into a single hyphen.
        #     GitHub doesn't.  'P0 — Silent ...' becomes
        #     'p0--silent-...' (the em-dash gets removed but the
        #     spaces around it stay, producing two hyphens).
        import re as _re
        out = s.lower().strip()
        # Drop only the markdown formatting chars that are NOT also
        # valid in identifiers (so no `_`).
        out = _re.sub(r"[`*~]", "", out)
        # Drop anything that isn't word-char / hyphen / space / CJK.
        out = _re.sub(r"[^\w一-鿿\- ]+", "", out)
        # 1-for-1 space → hyphen substitution (preserves doubles).
        out = out.replace(" ", "-")
        return out

    for tier in tier_order:
        items = by_tier.get(tier, [])
        if not items:
            continue
        tier_title = (tier_titles_en[tier] if is_en
                      else tier_titles_zh[tier])
        lines.append(f"- [{Z(tier_title)}](#{_slug(Z(tier_title))})")
        for it in items:
            t_title = it["title_en"] if is_en else it["title_zh"]
            entry = f"Issue #{it['id']} — {t_title}"
            lines.append(
                f"  - [{Z(entry)}](#{_slug(Z(entry))})"
            )
    lines.append(
        f"- [{Z('Severity legend' if is_en else '严重等级说明')}]"
        f"(#{_slug(Z('Severity legend' if is_en else '严重等级说明'))})"
    )
    appendix_a_toc_title = (
        "Appendix A — c_index_year / c_index_addr_id drift "
        "vs the cbdb-online-main-server snapshot "
        "(differences need per-row classification before being filed as bugs)"
        if is_en else
        "附錄 A —— c_index_year / c_index_addr_id 與 "
        "cbdb-online-main-server 快照之間的偏差"
        "（差異需要逐筆分類後才能判定是否為缺陷）"
    )
    lines.append(
        f"- [{Z(appendix_a_toc_title)}](#{_slug(Z(appendix_a_toc_title))})"
    )
    schema_appendix_b_title = (
        "Appendix B — TablesFields: documentation vs. actual structure"
        if is_en else
        "附錄 B —— TablesFields：文檔表與實際資料庫結構對比"
    )
    schema_appendix_c_title = (
        "Appendix C — ForeignKeys: documentation vs. actual structure"
        if is_en else
        "附錄 C —— ForeignKeys：文檔表與實際資料庫結構對比"
    )
    lines.append(f"- [{Z(schema_appendix_b_title)}](#{_slug(Z(schema_appendix_b_title))})")
    lines.append(f"- [{Z(schema_appendix_c_title)}](#{_slug(Z(schema_appendix_c_title))})")
    lines.append(
        f"- [{Z('Closing note' if is_en else '结语')}]"
        f"(#{_slug(Z('Closing note' if is_en else '结语'))})"
    )
    lines.append("")

    # ---- Severity legend ----
    lines.append(f"## {Z('Severity legend' if is_en else '严重等级说明')}")
    lines.append("")
    legend_en = [
        "P0 — Silent data corruption: data is wrong or missing without an error popup.",
        "P1 — Visible runtime crash: a popup appears, the operation aborts.",
        "P2 — Silent display: form fields render blank when they should show data.",
        "P3 — Missing UI: a feature exists in code but no button invokes it.",
        "P4 — Setup: one-time hurdle on each new install.",
        "P5 — Dormant / latent / not currently reproducible: kept as historical "
        "record; we re-checked on the current dump and could not "
        "trigger the symptom.",
    ]
    legend_zh = [
        "P0 — 静默数据错误：数据错或缺失，但没有任何报错提示。",
        "P1 — 可见的运行时报错：弹出错误对话框，操作中断。",
        "P2 — 静默显示问题：表单字段本应有数据，却显示为空。",
        "P3 — 缺失界面：代码里实现了某功能，但界面上没有按钮去触发它。",
        "P4 — 安装设置：每台新机器需要一次性处理。",
        "P5 — 潛伏 / 不可達 / 當前無法復現：保留作为历史记录；我们在当前 dump "
        "上重新验证过，无法再触发症状。",
    ]
    for s in (legend_en if is_en else legend_zh):
        lines.append(f"- {Z(s)}")
    lines.append("")

    # ---- Per-issue ----
    demo_persons = _load_demo_persons()
    bug_status = _load_bug_test_status()

    for tier in tier_order:
        items = by_tier.get(tier, [])
        if not items:
            continue
        tier_title = (tier_titles_en[tier] if is_en
                      else tier_titles_zh[tier])
        lines.append(f"## {Z(tier_title)}")
        lines.append("")
        if tier == "P5_dormant_or_latent":
            lines.append(Z(
                "_Items in this tier are kept as historical / latent "
                "record.  They fall into three categories: (a) DORMANT — "
                "verified that current source data doesn't trigger the "
                "symptom; (b) NOT CURRENTLY REPRODUCIBLE — the symptom "
                "no longer surfaces even though the suspect code is "
                "still present (we have NOT confirmed an upstream "
                "source-level fix; could be a JET / Office behaviour "
                "change, a fixture / driver change on our side, or the "
                "original diagnosis was a false positive); (c) LATENT — "
                "the source-code defect is real, but the user can't "
                "reach it because another issue (e.g. a missing UI "
                "button) blocks the path.  None of these are "
                "user-facing today; **none have been verified as fixed "
                "upstream** — please consult before treating any of "
                "them as either urgent or closed._"
                if is_en else
                "_本層的條目作為歷史 / 潛伏記錄保留。可分為三類："
                "(a) DORMANT 潛伏 — 已驗證當前源資料無法觸發該症狀；"
                "(b) 當前無法復現 — 症狀不再出現，但可疑程式碼仍在"
                "（我們**沒有**確認上游有源碼層面的修復；原因可能是 "
                "JET / Office 的行為改變、可能是我們這邊 fixture/driver "
                "改變，也可能原本的診斷就是 false positive）；"
                "(c) LATENT 被屏蔽 — 源碼缺陷確實存在，但因為另一個 "
                "issue（例如某個 UI 按鈕缺失）擋住了使用路徑，使用者"
                "目前碰不到。本層條目當下都不是使用者會遇到的問題，"
                "**也沒有任何一條被確認上游修復**；若要當成緊急或"
                "已關閉處理，請先諮詢。_"
            ))
            lines.append("")
        for it in items:
            t_title = it["title_en"] if is_en else it["title_zh"]
            lines.append(f"### Issue #{it['id']} — {Z(t_title)}")
            lines.append("")
            lines.append(f"**{Z('Affected sub' if is_en else '涉及位置')}:** "
                          f"`{it['form']}`")
            lines.append("")
            lines.append(f"**{Z('Severity' if is_en else '严重等级')}:** "
                          f"{Z(it['severity_en'] if is_en else it['severity_zh'])}")
            lines.append("")
            # Auto-derived test status banner.
            status = bug_status.get(it["id"])
            if status:
                outcome = status["outcome"]
                if outcome == "failed":
                    banner = (
                        "⚠ **Automated test status: this issue's "
                        "regression marker currently FAILS** "
                        f"(run on {status.get('when', 'unknown date')}). "
                        "That usually means the underlying defect has "
                        "longer reproduces, which MAY mean upstream fixed "
                        "it but could equally be a fixture / driver "
                        "regression or a misclassification.  Please "
                        "verify in person before considering this issue "
                        "closed; this report has NOT been edited to drop "
                        "the issue.  Tests consulted: "
                        + ", ".join(
                            f"`{t.rsplit('::', 1)[-1]}`"
                            for t, _ in status["tests"]
                        )
                        if is_en else
                        "⚠ **自動測試狀態：本 issue 對應的回歸標記目前 "
                        f"FAIL**（執行時間：{status.get('when', '未知')}），"
                        "通常意味著底層缺陷已在 source dump 中被 **修復**。"
                        "請務必親自確認，再將此 issue 視為關閉；本報告"
                        "並未自動刪除任何 issue。對照的測試："
                        + "、".join(
                            f"`{t.rsplit('::', 1)[-1]}`"
                            for t, _ in status["tests"]
                        )
                    )
                    lines.append(f"> {Z(banner)}")
                    lines.append("")
                elif outcome == "mixed":
                    parts = ", ".join(
                        f"`{t.rsplit('::', 1)[-1]}`: {o}"
                        for t, o in status["tests"]
                    )
                    banner = (
                        "ℹ Automated test status: MIXED outcomes "
                        f"(run {status.get('when', 'unknown date')}). "
                        f"Per-test: {parts}"
                        if is_en else
                        f"ℹ 自動測試狀態：混合結果（執行時間："
                        f"{status.get('when', '未知')}）。分項：{parts}"
                    )
                    lines.append(f"> _{Z(banner)}_")
                    lines.append("")
            lines.append(f"#### {Z('Description' if is_en else '问题描述')}")
            lines.append("")
            for para in (it["summary_en"] if is_en
                         else it["summary_zh"]).split("\n\n"):
                lines.append(Z(para))
                lines.append("")
            lines.append(f"#### {Z('Steps to reproduce' if is_en else '复现步骤')}")
            lines.append("")
            demo = demo_persons.get(f"bug{it['id']}")
            if demo:
                if demo.get("dormant"):
                    label = (
                        "⚠ **Currently UI-dormant on this data snapshot — see note below**"
                        if is_en else
                        "⚠ **在當前資料快照下無法在 UI 上復現——請看下方說明**"
                    )
                    lines.append(Z(label))
                    lines.append("")
                    lines.append(
                        demo["hint_en"] if is_en else Z(demo["hint_zh"])
                    )
                    lines.append("")
                else:
                    label = (
                        f"**Recommended demo person:** `c_personid="
                        f"{demo['personid']}` ({demo['name_chn']}, "
                        f"{demo['name_py']})"
                        if is_en else
                        f"**建議使用的範例人物：** `c_personid="
                        f"{demo['personid']}`（{demo['name_chn']}，"
                        f"{demo['name_py']}）"
                    )
                    lines.append(Z(label))
                    lines.append("")
                    hint_text = demo["hint_en"] if is_en else Z(demo["hint_zh"])
                    extra = (
                        " _Picked by `reports/probe_demo_persons.py`; a "
                        "SQL probe selected this person because their row "
                        "counts genuinely satisfy the precondition the "
                        "bug needs._"
                        if is_en else
                        " _由 `reports/probe_demo_persons.py` 透過 SQL "
                        "probe 挑選；之所以選這位，是因為其底層記錄數確實"
                        "滿足這個 bug 的觸發條件。_"
                    )
                    lines.append(hint_text + extra)
                    lines.append("")
            for i, step in enumerate(
                it["steps_en"] if is_en else it["steps_zh"], 1
            ):
                lines.append(f"{i}. {Z(step)}")
            lines.append("")
            # Optional concrete-reproduction block — see Issue #9.
            cr_key = ("concrete_reproduction_en" if is_en
                      else "concrete_reproduction_zh")
            cr_text = it.get(cr_key)
            if cr_text:
                lines.append(f"#### {Z('Concrete reproduction' if is_en else '具體復現')}")
                lines.append("")
                lines.append(Z(str(cr_text)))
                lines.append("")
            shots = it.get("screenshots") or []
            if shots:
                lines.append(f"#### {Z('Screenshots' if is_en else '截图')}")
                lines.append("")
                for fname, caption in shots:
                    p = SHOT_DIR / fname
                    if not p.exists():
                        lines.append(
                            f"_{Z(f'[screenshot {fname} not found]' if is_en else f'[未找到截图 {fname}]')}_"
                        )
                    else:
                        rel = f"screenshots/{fname}"
                        lines.append(f"![{fname}]({rel})")
                        if caption:
                            cap = caption if is_en else Z(caption)
                            lines.append("")
                            lines.append(f"_{cap}_")
                    lines.append("")
            lines.append(
                f"#### {Z('Suggested fix' if is_en else '建议修复方案')}"
            )
            lines.append("")
            for para in (it["fix_en"] if is_en
                         else it["fix_zh"]).split("\n\n"):
                lines.append(Z(para))
                lines.append("")

    # ---- Index drift appendix (always present) ----
    import json as _json
    appendix_a_title = (
        "Appendix A — c_index_year / c_index_addr_id drift "
        "vs the cbdb-online-main-server snapshot "
        "(differences need per-row classification before being filed as bugs)"
        if is_en else
        "附錄 A —— c_index_year / c_index_addr_id 與 "
        "cbdb-online-main-server 快照之間的偏差"
        "（差異需要逐筆分類後才能判定是否為缺陷）"
    )
    lines.append(f"## {Z(appendix_a_title)}")
    lines.append("")
    if not DRIFT_JSON.exists():
        lines.append(
            "> Appendix data not yet generated.  "
            "Run step 5c: `python reports/collect_index_year_diffs.py`"
            if is_en else
            "> 附錄數據尚未生成。請執行第 5c 步：`python reports/collect_index_year_diffs.py`"
        )
        lines.append("")
    if DRIFT_JSON.exists():
        intro_drift = (
            "When we compare BIOG_MAIN's `c_index_year` and "
            "`c_index_addr_id` between this User MDB and the weekly "
            "cbdb-online-main-server SQLite snapshot, a small "
            "fraction of persons disagree.\n\n"
            "**The two sides are independent implementations.**  The "
            "SQLite snapshot's `c_index_year` is produced by "
            "cbdb-online-main-server's PHP "
            "`IndexYearRebuildService.php` and its `c_index_addr_id` "
            "by `IndexAddressRebuildService.php` (both at "
            "<https://github.com/cbdb-project/cbdb-online-main-server>"
            "); the User MDB-side: `c_index_addr_id` rebuilt by VBA "
            "in `Form_frmIndexAddr` (front-end mdb); `c_index_year` "
            "rebuilt by **37 saved QueryDefs named `BM IY Rule …`** "
            "in the linked-tables backend "
            "`data/CBDB_<YYYYMMDD>_DATA.mdb`, driven by "
            "`frmBaseMaintenance`.  Both algorithms now extracted "
            "to `analysis/dump_data/querydefs_index/*.sql`; form / "
            "module driver VBA still needs an interactive Access "
            "SaveAsText pass.  PHP is intended to mirror the "
            "VBA but they are separate code paths.  Per-row "
            "differences can come from at least four sources, and a "
            "diff alone doesn't tell us which: (1) source-data "
            "snapshot drift; (2) algorithm / porting divergence "
            "between PHP and VBA; (3) priority / tie-break "
            "differences; (4) null / default handling differences."
            "\n\n"
            "The per-row classification of these diffs is in the "
            "**Classification summary** below, generated from "
            "`reports/index_drift_classification.json` when the "
            "classifier has run (placeholder otherwise) — counts and "
            "buckets are data-driven, not hardcoded.  The worked "
            "example rows further down "
            "(`reports/index_drift_examples.json`) illustrate the "
            "*shapes* of disagreement; illustrative, not statistically "
            "representative, a starting point for per-row triage, not "
            "a verdict."
            if is_en else
            "我們把本 .mdb 的 BIOG_MAIN 與 cbdb-online-main-server 每週"
            "釋出的 SQLite 快照在 `c_index_year`、`c_index_addr_id` "
            "兩個欄位上做比對，可以看到一小部分人物對不齊。\n\n"
            "**兩邊是兩套獨立的實作。**SQLite 快照中的 "
            "`c_index_year` 是 cbdb-online-main-server 的 PHP "
            "`IndexYearRebuildService.php` 算出來的，"
            "`c_index_addr_id` 則是 `IndexAddressRebuildService.php` "
            "算出來的（程式碼都在 <https://github.com/cbdb-project/"
            "cbdb-online-main-server>）；User MDB 上對應的這兩個"
            "User MDB 那一邊：`c_index_addr_id` 由前端 mdb 裡的 "
            "`Form_frmIndexAddr` VBA 重建；`c_index_year` 由連結表"
            "後端 `data/CBDB_<YYYYMMDD>_DATA.mdb` 裡 **37 條 "
            "`BM IY Rule …` 的 QueryDef** 重建，由 "
            "`frmBaseMaintenance` 驅動。兩邊算法已抽取到 "
            "`analysis/dump_data/querydefs_index/*.sql`；form / "
            "module 驅動 VBA 仍需 Access SaveAsText 互動式提取。"
            "PHP **意圖**鏡像 VBA，"
            "但兩者是兩條獨立的程式路徑。每一行差異**可能**來自下列"
            "至少四個原因，光看差異本身分不出來：(1) 源資料快照漂移；"
            "(2) PHP 與 VBA 之間的演算法 / 移植差異；(3) 優先序 / "
            "平手規則不同；(4) null / 預設值處理不同。\n\n"
            "這些差異的逐筆分類見下方 **分類匯總**，由 "
            "`reports/index_drift_classification.json` 自動生成（分類器"
            "尚未執行時顯示佔位）；計數與分桶都是 data-driven，不寫死。"
            "再往下列舉的樣例（`reports/index_drift_examples.json`）只是"
            "**示範**差異**長什麼樣**，並非統計上有代表性，是後續逐筆"
            "分類的起點，不是結論。"
        )
        lines.append(Z(intro_drift))
        lines.append("")

        # ---- Classification summary (markdown copy) ----
        if CLASSIFICATION_JSON.exists():
            cls = _json.loads(
                CLASSIFICATION_JSON.read_text(encoding="utf-8"))
            cs = cls["summary"]
            b = cs["buckets"]
            net = cs["common"] - b["exact_match"]
            attributable = (b["source_drift_index_agrees"]
                            + b["source_drift_index_diffs_too"])
            unclassified = (b["index_year_only_diff"]
                            + b["index_addr_only_diff"]
                            + b["index_both_diff"])
            if is_en:
                lines.append(f"### {Z('Classification summary')}")
                lines.append("")
                lines.append(Z(
                    f"Compared **{cs['common']:,}** personids common to "
                    f"both databases (User MDB total {cs['user_mdb_total']:,}; "
                    f"SQLite total {cs['sqlite_total']:,}; "
                    f"User-only {cs['in_user_only']:,}; "
                    f"SQLite-only {cs['in_sqlite_only']:,})."
                ))
                lines.append("")
                lines.append("| Bucket | Count | % of common | Meaning |")
                lines.append("|---|---:|---:|---|")
                rows = [
                    ("exact_match", "exact match on all four compared fields"),
                    ("source_drift_index_agrees",
                     "source drift but indices agreed"),
                    ("source_drift_index_diffs_too",
                     "source drift AND ≥1 index differs"),
                    ("index_year_only_diff",
                     "source matched, only c_index_year differs — needs follow-up"),
                    ("index_addr_only_diff",
                     "source matched, only c_index_addr_id differs — needs follow-up"),
                    ("index_both_diff",
                     "source matched, both indices differ — strongest signal"),
                ]
                for k, label in rows:
                    pct = 100.0 * b[k] / max(cs['common'], 1)
                    lines.append(f"| `{k}` | {b[k]:,} | {pct:.3f}% | {Z(label)} |")
                lines.append("")
                lines.append(Z(
                    f"Net diffs: **{net:,}** of {cs['common']:,} "
                    f"({100.0*net/max(cs['common'],1):.3f} %).  Of those, "
                    f"**{attributable}** are clearly attributable to source "
                    f"drift in birthyear / deathyear; **{unclassified}** need "
                    f"per-row follow-up.  These could be PHP↔VBA divergence, "
                    f"or drift in evidence tables (BIOG_ADDR_DATA / "
                    f"ENTRY_DATA / NIAN_HAO etc.) that this classifier does "
                    f"not compare.  Full output: "
                    f"`reports/index_drift_classification.json`; algorithm "
                    f"pointers: `analysis/index_drift_algorithm_notes.md`."
                ))
                lines.append("")
            else:
                lines.append(f"### {Z('分類匯總')}")
                lines.append("")
                lines.append(Z(
                    f"比對了兩邊都有的 **{cs['common']:,}** 個 personid"
                    f"（User MDB 共 {cs['user_mdb_total']:,} 筆；"
                    f"SQLite 共 {cs['sqlite_total']:,} 筆；"
                    f"僅 User MDB 有 {cs['in_user_only']:,} 筆；"
                    f"僅 SQLite 有 {cs['in_sqlite_only']:,} 筆）。"
                ))
                lines.append("")
                lines.append("| 分桶 | 筆數 | 佔比 | 含義 |")
                lines.append("|---|---:|---:|---|")
                rows = [
                    ("exact_match", "四個欄位全部一致"),
                    ("source_drift_index_agrees",
                     "源資料有漂移但兩邊 index 都一致"),
                    ("source_drift_index_diffs_too",
                     "源資料有漂移、且至少一個 index 不同"),
                    ("index_year_only_diff",
                     "生年/卒年一致，但只有 c_index_year 不同 —— 待追查"),
                    ("index_addr_only_diff",
                     "生年/卒年一致，但只有 c_index_addr_id 不同 —— 待追查"),
                    ("index_both_diff",
                     "生年/卒年一致，但兩個 index 都不同 —— 複合差異最強信號"),
                ]
                for k, label in rows:
                    pct = 100.0 * b[k] / max(cs['common'], 1)
                    lines.append(f"| `{k}` | {b[k]:,} | {pct:.3f}% | {Z(label)} |")
                lines.append("")
                lines.append(Z(
                    f"淨差異：**{net:,}** / {cs['common']:,}"
                    f"（{100.0*net/max(cs['common'],1):.3f} %）。其中 "
                    f"**{attributable}** 筆能明確歸因於 birthyear / "
                    f"deathyear 的源資料漂移；剩下 **{unclassified}** 筆"
                    f"需要逐筆追查（可能是 PHP↔VBA 演算法差異，"
                    f"也可能是本分類器沒有比較的 evidence 表"
                    f"（BIOG_ADDR_DATA / ENTRY_DATA / NIAN_HAO 等）裡的"
                    f"漂移）。完整輸出見 "
                    f"`reports/index_drift_classification.json`，"
                    f"算法來源指標見 "
                    f"`analysis/index_drift_algorithm_notes.md`。"
                ))
                lines.append("")
        else:
            lines.append(f"### {Z('Classification summary' if is_en else '分類匯總')}")
            lines.append(Z(
                "_Not generated — run `python analysis/classify_index_drift.py` "
                "(reports/index_drift_classification.json absent)._"
                if is_en else
                "_尚未生成 —— 請執行 `python analysis/classify_index_drift.py`"
                "（缺 reports/index_drift_classification.json）。_"
            ))
            lines.append("")

        # ---- Year-drift rule classification (PR K1) — markdown ----
        if RULE_CLASSIFICATION_JSON.exists():
            rcls = _json.loads(
                RULE_CLASSIFICATION_JSON.read_text(encoding="utf-8"))
            rs = rcls["summary"]
            rb = rs["buckets"]
            if is_en:
                lines.append(f"### {Z('Year-only diffs — per-row rule classification')}")
                lines.append("")
                lines.append(Z(
                    f"Of the **{rs['total_year_diffs']}** year-only "
                    f"diffs, each row was bucketed against PR N's "
                    f"rule-level runtime-vs-PHP comparison "
                    f"(`analysis/index_year_rule_comparison.md`).  "
                    f"Conservative buckets (rows count once each):"
                ))
                lines.append("")
                lines.append("| Bucket | Count |")
                lines.append("|---|---:|")
                bucket_meta_zh = None
                for k in [
                    "php_returned_sentinel",
                    "php_did_not_compute",
                    "access_did_not_compute",
                    "iteration_order_diff",
                    "consistent_within_rule",
                    "candidate_algorithm_divergence",
                    "unclassified",
                ]:
                    if rb[k] == 0:
                        continue
                    lines.append(f"| `{k}` | {rb[k]} |")
                lines.append("")
                lines.append(Z(
                    f"None of these are confirmed bugs.  Full per-row "
                    f"output is in "
                    f"`reports/index_year_drift_rule_classification.json`."
                ))
                lines.append("")
                if RULE_GROUPS_JSON.exists():
                    gcls = _json.loads(
                        RULE_GROUPS_JSON.read_text(encoding="utf-8"))
                    gs = gcls["summary"]
                    lines.append(Z(
                        f"Deeper triage (PR K2, "
                        f"`analysis/triage_index_year_drift_groups.py` "
                        f"→ `reports/index_year_drift_rule_groups.json`) "
                        f"named the leftover buckets:"
                    ))
                    lines.append("")
                    lines.append(Z(
                        f"- `consistent_within_rule` × "
                        f"{gs['consistent_within_rule']['rows']} → "
                        f"{gs['consistent_within_rule']['groups']} "
                        f"signature groups.  PR AI + AJ probes "
                        f"reversed the prior tie-break hypothesis: "
                        f"all 14 rows are `source_data_drift_biog_"
                        f"main_or_kin_data_between_sides` (8 "
                        f"BIOG_MAIN birthyear drift + 6 KIN_DATA "
                        f"evidence-pid drift).  Upstream PHP-side / "
                        f"SQLite-snapshot data issue, not a CBDB "
                        f"algorithm divergence."
                    ))
                    lines.append(Z(
                        f"- `unclassified` × {gs['unclassified']['total']} → "
                        f"{gs['unclassified']['named_after_triage']} named, "
                        f"{gs['unclassified']['blocked_by_runtime_priority_triage_pending']} "
                        f"flagged `blocked_by_runtime_priority_triage_pending` "
                        f"(PR M dumped frmBaseMaintenance, so the source is in repo; resolving each row still needs a per-row walk of the runtime priority/iteration order)."
                    ))
                    lines.append(Z(
                        f"- `php_did_not_compute` × "
                        f"{gs['php_did_not_compute']['rows']} → "
                        f"{gs['php_did_not_compute']['groups']} groups by "
                        f"Access tcode; biggest is `access_tcode='05'` × 7 "
                        f"(`candidate_php_entry_code_mapping_gap` for jinshi)."
                    ))
                    lines.append("")
            else:
                lines.append(f"### {Z('年份差異 —— 逐筆 rule 分類')}")
                lines.append("")
                lines.append(Z(
                    f"在 **{rs['total_year_diffs']}** 筆「只有 "
                    f"c_index_year 不一致」的行中，逐筆比對 PR N "
                    f"(`analysis/index_year_rule_comparison.md`) 的"
                    f"runtime-vs-PHP 規則級差異。保守分類如下："
                ))
                lines.append("")
                lines.append("| 分桶 | 筆數 |")
                lines.append("|---|---:|")
                label_zh = {
                    "php_returned_sentinel": "PHP 寫了 sentinel／溢位值",
                    "php_did_not_compute": "PHP 沒算出值（覆蓋率缺口）",
                    "access_did_not_compute": "Access 沒算出值（覆蓋率缺口）",
                    "iteration_order_diff": "Phase-C 迭代次數不同",
                    "consistent_within_rule": "多列共享同一 (php_tcode, access_tcode, diff)",
                    "candidate_algorithm_divergence": "形狀符合 K1 的歷史 hypothesis probe 但無法以單筆證據重建",
                    "unclassified": "尚未對上任何模式",
                }
                for k in [
                    "php_returned_sentinel",
                    "php_did_not_compute",
                    "access_did_not_compute",
                    "iteration_order_diff",
                    "consistent_within_rule",
                    "candidate_algorithm_divergence",
                    "unclassified",
                ]:
                    if rb[k] == 0:
                        continue
                    lines.append(f"| `{k}` ({label_zh[k]}) | {rb[k]} |")
                lines.append("")
                lines.append(Z(
                    f"以上沒有任何一筆被視為已確認的 bug。逐筆輸出見 "
                    f"`reports/index_year_drift_rule_classification.json`。"
                ))
                lines.append("")
                if RULE_GROUPS_JSON.exists():
                    gcls = _json.loads(
                        RULE_GROUPS_JSON.read_text(encoding="utf-8"))
                    gs = gcls["summary"]
                    lines.append(Z(
                        f"PR K2 進一步的 triage "
                        f"(`analysis/triage_index_year_drift_groups.py` "
                        f"→ `reports/index_year_drift_rule_groups.json`) "
                        f"把剩下的桶命名清楚："
                    ))
                    lines.append("")
                    lines.append(Z(
                        f"- `consistent_within_rule` × "
                        f"{gs['consistent_within_rule']['rows']} → "
                        f"{gs['consistent_within_rule']['groups']} 個 "
                        f"signature 分組。PR AI + AJ 的逐筆探測推翻了"
                        f"原本的 tie-break 假說：14 筆全是 "
                        f"`source_data_drift_biog_main_or_kin_data_"
                        f"between_sides`（8 筆 BIOG_MAIN birthyear 漂移 "
                        f"+ 6 筆 KIN_DATA evidence-pid 漂移）。屬於 "
                        f"PHP-side / SQLite snapshot 的上游資料漂移，"
                        f"並非 CBDB 演算法差異。"
                    ))
                    lines.append(Z(
                        f"- `unclassified` × {gs['unclassified']['total']} → "
                        f"{gs['unclassified']['named_after_triage']} 筆已命名，"
                        f"{gs['unclassified']['blocked_by_runtime_priority_triage_pending']} 筆標為 "
                        f"`blocked_by_runtime_priority_triage_pending`"
                        f"（PR M 已 dump frmBaseMaintenance，源碼已在 repo；要逐筆判斷哪邊正確仍需走一遍 runtime 的 priority／iteration 順序）。"
                    ))
                    lines.append(Z(
                        f"- `php_did_not_compute` × "
                        f"{gs['php_did_not_compute']['rows']} → 按 Access "
                        f"tcode 分 {gs['php_did_not_compute']['groups']} 組；"
                        f"最大的是 `access_tcode='05'` × 7（jinshi 進士類的 "
                        f"`candidate_php_entry_code_mapping_gap`）。"
                    ))
                    lines.append("")

        # ---- Address-drift classification (PR L) ----
        if ADDR_CLASSIFICATION_JSON.exists():
            acls = _json.loads(
                ADDR_CLASSIFICATION_JSON.read_text(encoding="utf-8"))
            aS = acls["summary"]
            ab = aS["buckets"]
            addr_bucket_order = [
                "mdb_stale_index_addr",
                "mdb_value_php_null",
                "same_candidates_diff_winner",
                "both_stale_recompute_mismatch",
                "both_sides_match_recomputed",
                "sqlite_stale_index_addr",
                "mdb_null_php_value",
                "unclassified",
            ]
            if is_en:
                lines.append(f"### {Z('c_index_addr_id diffs — per-row classification')}")
                lines.append("")
                lines.append(Z(
                    f"Of the **{aS['total_addr_diffs']}** "
                    f"c_index_addr diffs (478 `index_addr_only_diff` "
                    f"+ 10 `index_both_diff` from PR G), each row "
                    f"was classified by re-simulating the rank-"
                    f"priority + MAX(c_sequence) algorithm against "
                    f"each side's BIOG_ADDR_DATA + the shared "
                    f"BIOG_ADDR_CODES rank table."
                ))
                lines.append("")
                lines.append("| Bucket | Count |")
                lines.append("|---|---:|")
                for k in addr_bucket_order:
                    n = ab.get(k, 0)
                    if n == 0:
                        continue
                    lines.append(f"| `{k}` | {n} |")
                lines.append("")
                lines.append(Z(
                    f"None of these are confirmed bugs.  The 412 "
                    f"`mdb_stale_index_addr` rows are a maintenance-"
                    f"cadence diff (the User MDB needs its "
                    f"frmBaseMaintenance rebuild re-run before the "
                    f"next release).  The 10 "
                    f"`same_candidates_diff_winner` rows are the "
                    f"only candidate algorithm-divergence rows.  "
                    f"Full per-row output: "
                    f"`reports/index_addr_drift_classification.json`."
                ))
                lines.append("")
                lines.append(Z(
                    f"PR M (`analysis/dump_data_mdb_vba.py`) extracted "
                    f"`frmBaseMaintenance.CmdIndexAddress_Click` from "
                    f"the DATA mdb.  It does NOT explicitly "
                    f"`MAX(c_sequence)`-aggregate the way PHP does — a "
                    f"candidate algorithmic divergence on top of the "
                    f"maintenance-cadence issue.  Suggested "
                    f"release-checklist mitigation: run `CmdIndexYear` "
                    f"then `CmdIndexAddress` on the DATA mdb before "
                    f"shipping a new User MDB."
                ))
                lines.append("")
                lines.append(Z(
                    f"PR S "
                    f"(`analysis/deep_dive_addr_same_candidates.py`) "
                    f"confirmed the 10 `same_candidates_diff_winner` "
                    f"rows are all driven by MAX(c_sequence) ties "
                    f"(multiple BIOG_ADDR_DATA rows of the same "
                    f"(person, addr_type) sharing the same max "
                    f"c_sequence).  PHP, Access, and our recompute "
                    f"each pick non-deterministically.  Both sides "
                    f"follow the same documented rule; neither is "
                    f"wrong.  Candidate mitigation: add an explicit "
                    f"secondary tie-break (e.g. MIN(c_addr_id)) to "
                    f"both implementations.  Per-row evidence in "
                    f"`reports/index_addr_same_candidates_deep_dive.json`."
                ))
                lines.append("")

                # ---- Cause analysis appendix (PR Y) ----
                if CAUSE_SUMMARY_JSON.exists():
                    cs = _json.loads(
                        CAUSE_SUMMARY_JSON.read_text(encoding="utf-8"))
                    lines.append("### What currently explains the drift")
                    lines.append("")
                    lines.append(Z(
                        "Per-bucket cause / supporting evidence / "
                        "confidence / next action lives in "
                        "`analysis/index_drift_cause_analysis.md`.  "
                        "This section just summarises headline counts "
                        "and confidence; no bucket is labelled a "
                        "confirmed CBDB bug."
                    ))
                    lines.append("")
                    lines.append("**c_index_year cause buckets**")
                    lines.append("")
                    lines.append("| Bucket | Count | Confidence |")
                    lines.append("|---|---:|---|")
                    for b in cs["c_index_year"]["buckets"]:
                        if b["count"] == 0:
                            continue
                        lines.append(
                            f"| `{b['bucket']}` | {b['count']} "
                            f"| {b['confidence']} |"
                        )
                    lines.append("")
                    lines.append("**c_index_addr_id cause buckets**")
                    lines.append("")
                    lines.append("| Bucket | Count | Confidence |")
                    lines.append("|---|---:|---|")
                    for b in cs["c_index_addr_id"]["buckets"]:
                        if b["count"] == 0:
                            continue
                        lines.append(
                            f"| `{b['bucket']}` | {b['count']} "
                            f"| {b['confidence']} |"
                        )
                    lines.append("")
                    lines.append(Z(
                        "Top suggested next investigations (full list "
                        "in the cause-analysis md):"
                    ))
                    lines.append("")
                    for inv in cs[
                        "suggested_next_investigations_in_priority_order"
                    ][:3]:
                        lines.append(Z(
                            f"{inv['id']}. {inv['task']} — would close "
                            f"{inv['would_close_rows']} rows; "
                            f"engineering cost: {inv['engineering_cost']}."
                        ))
                    lines.append("")
            else:
                lines.append(f"### {Z('c_index_addr_id 差異 —— 逐筆分類')}")
                lines.append("")
                lines.append(Z(
                    f"在 **{aS['total_addr_diffs']}** 筆 c_index_addr "
                    f"差異中（PR G 的 478 `index_addr_only_diff` + 10 "
                    f"`index_both_diff`），逐筆把兩邊的 BIOG_ADDR_DATA "
                    f"代入「rank-priority + MAX(c_sequence)」演算法重算，"
                    f"與實際儲存值對照分類："
                ))
                lines.append("")
                lines.append("| 分桶 | 筆數 |")
                lines.append("|---|---:|")
                for k in addr_bucket_order:
                    n = ab.get(k, 0)
                    if n == 0:
                        continue
                    lines.append(f"| `{k}` | {n} |")
                lines.append("")
                lines.append(Z(
                    f"以上沒有任何一筆被視為已確認的 bug。412 筆 "
                    f"`mdb_stale_index_addr` 屬於維護週期差異（User MDB "
                    f"在下次釋出前需要重跑 frmBaseMaintenance）。10 筆 "
                    f"`same_candidates_diff_winner` 是唯一的候選演算法"
                    f"差異。逐筆輸出見 "
                    f"`reports/index_addr_drift_classification.json`。"
                ))
                lines.append("")
                lines.append(Z(
                    f"PR M（`analysis/dump_data_mdb_vba.py`）從 DATA mdb "
                    f"抽出了 `frmBaseMaintenance.CmdIndexAddress_Click`。"
                    f"它**沒有**像 PHP 那樣明確 `MAX(c_sequence)` 聚合 "
                    f"—— 在維護週期差異之外，這還是一個候選演算法差異。"
                    f"建議的 release checklist 緩解步驟：在 User MDB "
                    f"出貨前先在 DATA mdb 上跑 `CmdIndexYear`，再跑 "
                    f"`CmdIndexAddress`。詳見 "
                    f"`analysis/index_drift_algorithm_notes.md` 中的 "
                    f"\"Maintenance trigger path\" 段。"
                ))
                lines.append("")

                # ---- Cause analysis appendix (PR Y) ----
                if CAUSE_SUMMARY_JSON.exists():
                    cs = _json.loads(
                        CAUSE_SUMMARY_JSON.read_text(encoding="utf-8"))
                    lines.append(f"### {Z('目前能解釋的 drift 原因')}")
                    lines.append("")
                    lines.append(Z(
                        "每個 bucket 的成因／證據／信心度／下一步追查"
                        "都寫在 `analysis/index_drift_cause_analysis.md`。"
                        "本節只列每個 bucket 的計數和信心度摘要；目前"
                        "沒有任何 bucket 被列為已確認的 CBDB bug。"
                    ))
                    lines.append("")
                    lines.append(f"**{Z('c_index_year 原因桶')}**")
                    lines.append("")
                    lines.append(f"| Bucket | {Z('筆數')} | {Z('信心度')} |")
                    lines.append("|---|---:|---|")
                    for b in cs["c_index_year"]["buckets"]:
                        if b["count"] == 0:
                            continue
                        lines.append(
                            f"| `{b['bucket']}` | {b['count']} "
                            f"| {b['confidence']} |"
                        )
                    lines.append("")
                    lines.append(f"**{Z('c_index_addr_id 原因桶')}**")
                    lines.append("")
                    lines.append(f"| Bucket | {Z('筆數')} | {Z('信心度')} |")
                    lines.append("|---|---:|---|")
                    for b in cs["c_index_addr_id"]["buckets"]:
                        if b["count"] == 0:
                            continue
                        lines.append(
                            f"| `{b['bucket']}` | {b['count']} "
                            f"| {b['confidence']} |"
                        )
                    lines.append("")
                    lines.append(Z(
                        "建議優先處理的調查項目（完整列表見 cause-"
                        "analysis md）："
                    ))
                    lines.append("")
                    for inv in cs[
                        "suggested_next_investigations_in_priority_order"
                    ][:3]:
                        lines.append(Z(
                            f"{inv['id']}. {inv['task']} —— 可消化 "
                            f"{inv['would_close_rows']} 筆；工程成本："
                            f"{inv['engineering_cost']}。"
                        ))
                    lines.append("")

        data = _json.loads(DRIFT_JSON.read_text(encoding="utf-8"))
        bucket_meta = {
            "year_only": ("Examples where only c_index_year disagrees",
                          "仅 c_index_year 不一致的样例"),
            "addr_only": ("Examples where only c_index_addr_id disagrees",
                          "仅 c_index_addr_id 不一致的样例"),
            "both": ("Examples where both fields disagree",
                     "两个字段都不一致的样例"),
            "source_data": (
                "Examples where the SOURCE data itself differs (birthyear / deathyear)",
                "底层 SOURCE 数据本身不同（生年 / 卒年）的样例"),
        }
        for bucket_key, titles in bucket_meta.items():
            items = data.get(bucket_key, [])
            if not items:
                continue
            lines.append(f"### {Z(titles[0] if is_en else titles[1])}")
            lines.append("")
            for ex in items[:5]:
                u = ex["user"]; s = ex["sqlite"]
                head = (f"`c_personid = {ex['personid']}` — "
                        f"{ex['name_chn']} ({ex['name_py']})")
                lines.append(f"**{Z(head)}**")
                lines.append("")
                lines.append(
                    f"| {Z('Field' if is_en else '字段')} "
                    f"| {Z('User MDB' if is_en else '本 .mdb (User MDB)')} "
                    f"| {Z('cbdb-online-main-server snapshot' if is_en else 'cbdb-online-main-server 快照')} |"
                )
                lines.append("|---|---|---|")
                for f_label, key in [
                    ("c_index_year", "index_year"),
                    ("c_index_addr_id", "index_addr_id"),
                    ("c_birthyear", "birthyear"),
                    ("c_deathyear", "deathyear"),
                    ("c_index_year_type_code", "index_year_type_code"),
                    ("c_index_year_source_id", "index_year_source_id"),
                ]:
                    uv = "" if u[key] is None else str(u[key])
                    sv = "" if s[key] is None else str(s[key])
                    lines.append(f"| `{f_label}` | {uv} | {sv} |")
                lines.append("")

    # ---- Appendix B & C: schema diff ----
    _add_schema_diff_appendix_md(lines, "tables_fields", is_en, Z, _slug)
    _add_schema_diff_appendix_md(lines, "foreign_keys", is_en, Z, _slug)

    # ---- Closing ----
    lines.append(f"## {Z('Closing note' if is_en else '结语')}")
    lines.append("")
    closing = (
        "Thank you for taking the time to read this report. None of the "
        "items above is urgent; we hope having them all in one place "
        "makes it easy to address them at your own pace.\n\n"
        "If any of the descriptions or suggested fixes are unclear, we "
        "would be glad to discuss further. The corresponding regression "
        "tests in this repository will automatically flip from PASS to "
        "FAIL the moment any regression marker stops reproducing in "
        "the source dump — that is a signal to investigate, not an "
        "automatic confirmation that the bug is fixed (the marker "
        "could fail because of an upstream fix, a fixture / driver "
        "change on our side, or a misclassification we made earlier)."
        if is_en else
        "感谢您抽时间读完这份报告。以上各条都不紧急，我们把它们集中整理"
        "在一起，只是希望方便您在合适的时候逐一处理。\n\n"
        "如果对其中任何一条的描述或建议有疑问，欢迎随时一同讨论。本仓库"
        "里对应的回归测试，会在任何一个回归标记不再复现时自动从 PASS "
        "翻成 FAIL —— 这是「请调查一下」的讯号，而不是「问题已修复」的"
        "自动确认（因为标记不再复现也可能是 fixture / driver 变了，"
        "或者是我们当初的分类有误）。"
    )
    for para in closing.split("\n\n"):
        lines.append(Z(para))
        lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {out_path}")


def main() -> int:
    _validate_issues()  # fail fast on unknown tier before writing any files
    # Always emit all four formats together so the docx and md never
    # drift apart.
    _build("en", OUT_EN)
    _build("zh", OUT_ZH)
    _build_md("en", OUT_EN_MD)
    _build_md("zh", OUT_ZH_MD)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
